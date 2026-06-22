"""V10.5 — Proiect Bransament Complet (replica documentului user) cu placeholders.

Replică structura EXACTĂ a "Proiect bransament.docx" încărcat de utilizator,
dar înlocuiește toate valorile variabile cu placeholdere care vin din data.
Materialele se generează automat din materials_db (ANEXA 13 — 554 SAP codes).

Secțiuni:
1. REFERAT (verificare proiect VGD)
2. FOAIE DE CAPAT
3. BORDEROU
4. MEMORIU TEHNIC (4 capitole)
5. ANEXA 14 — LISTA MATERIALE (auto-generată)
6. FISA ASPECTE DE MEDIU
"""
from __future__ import annotations

import io
from datetime import datetime
from typing import Any, Dict, List

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

import materials_db
from gas_doc_templates import (
    COMPANY, _today_ro, _get, _add_heading, _add_para, _add_kv_table,
)


def _add_centered(doc: Document, text: str, bold: bool = True, size: int = 14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)


def proiect_bransament_complet(proj: Dict[str, Any]) -> bytes:
    """Generate the FULL bransament project document.

    Replică structura exactă a documentului user-ului uploadat ca artifact.
    """
    data = proj.get("data") or {}
    doc = Document()

    # Set base margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # ====================================================
    # SECȚIUNEA 1 — REFERAT (verificare proiect VGD)
    # ====================================================
    p = doc.add_paragraph()
    p.add_run("Numele si prenumele verificatorului atestat: ").italic = True
    p.add_run(_get(data, "vgd_verificator_nume", "ing. ____________________")).bold = True

    p = doc.add_paragraph()
    p.add_run("Legitimatie nr. ").italic = True
    p.add_run(_get(data, "vgd_legitimatie_nr", "_______________")).bold = True
    p.add_run("\t\tNr. ").italic = True
    p.add_run(_get(data, "vgd_referat_nr", "________")).bold = True
    p.add_run(" / Data: ").italic = True
    p.add_run(_get(data, "vgd_referat_data", _today_ro())).bold = True

    doc.add_paragraph()
    _add_centered(doc, "REFERAT", bold=True, size=18)
    _add_centered(doc, "privind verificarea de calitate la cerinta V.G.D. a proiectului:", bold=False, size=11)
    _add_centered(doc, f"{_get(data, 'tipul_lucrarii', 'Bransament gaze naturale (racord) nou proiectat')}", bold=True, size=12)
    doc.add_paragraph()

    _add_heading(doc, "1. Date de identificare", level=2)
    _add_para(doc, f"• Proiectant general: {_get(data, 'proiectant_general_firma', COMPANY['name'])}")
    _add_para(doc, f"• Proiectant instalator autorizat: {_get(data, 'fact_proiectanta_nume')}, "
                   f"Nr. aut. {_get(data, 'fact_proiectanta_legitimatie')}, "
                   f"grad {_get(data, 'fact_proiectanta_autorizatie')}.")
    _add_para(doc, f"• Client: {_get(data, 'nume_client') or _get(data, 'beneficiar_nume')}")
    _add_para(doc, f"• Amplasamentul lucrarilor: {_get(data, 'amplasament_strada')}, "
                   f"{_get(data, 'amplasament_localitate')}, {_get(data, 'amplasament_judet')}")
    _add_para(doc, f"• Amplasamentul imobilului: {_get(data, 'adresa_imobil') or _get(data, 'amplasament_strada')}, "
                   f"Nr. {_get(data, 'loc_consum_numar', '_')}, "
                   f"{_get(data, 'amplasament_localitate')}, {_get(data, 'amplasament_judet')}")
    _add_para(doc, f"• Ordin de lucru Nr.: {_get(data, 'osd_ordin_nr')}, emis de {_get(data, 'osd_operator', 'S.C. Distrgaz Sud Retele S.A.')}")
    _add_para(doc, f"• Acord de acces Nr.: {_get(data, 'atr_numar') or _get(data, 'osd_atr_nr')}, "
                   f"emis de {_get(data, 'osd_operator', 'S.C. Distrgaz Sud Retele S.A.')}")

    _add_heading(doc, "2. Caracteristicile principale ale proiectului", level=2)
    cu_str = (f"C.U. Nr.: {_get(data, 'cu_numar', '____')}" if data.get("are_cu_existent") in {"Da", True}
              else "C.U. în curs de obținere")
    avize_list = []
    for aviz_key, aviz_label in [
        ("aviz_apa_nova", "Apa Nova"), ("aviz_e_distributie", "E-Distributie"),
        ("aviz_telekom", "Telekom"), ("aviz_netcity", "Netcity Telecom"),
        ("aviz_stb_ratb", "S.T.B."), ("aviz_mediu_apm", "Mediu (A.P.M.)"),
        ("aviz_administr_strazi", "Administratia Strazilor"),
    ]:
        v = data.get(aviz_key)
        if v:
            avize_list.append(f"{aviz_label}: {v}")
    _add_para(doc, f"• Certificat de urbanism, avize si autorizatii: {cu_str}")
    if avize_list:
        for a in avize_list:
            _add_para(doc, f"   – {a}")
    _add_para(doc, f"• Cladire cu destinatia: {_get(data, 'cladire_destinatie', 'locuinta')}, "
                   f"categoria de importanta: {_get(data, 'cladire_cat_importanta', 'C - normala')}.")
    _add_para(doc, f"• Obiectul lucrarii: {_get(data, 'tipul_lucrarii', 'Bransament gaze naturale (racord) nou proiectat')}")
    _add_para(doc, f"• Debit aprobat: {_get(data, 'debit_instalat_mc_h', '25')} Nmc/h")
    _add_para(doc, "• Combustibil utilizat: gaze naturale")
    _add_para(doc, f"• Instalatie: conducte din "
                   f"{'PE100 SDR11' if (data.get('br_material') or '').lower().startswith('poliet') else 'OL'}, "
                   f"Dn {_get(data, 'br_diametru_dn', '32 mm')}, "
                   f"L = {_get(data, 'br_lungime_m', '4')} m"
                   f"{' (' + _get(data, 'br_lungime_breakdown', '3m + 1m raiser') + ')' if data.get('br_lungime_breakdown') else ''}.")
    tub_protectie = _get(data, "br_tub_protectie_necesar", "nu este necesar")
    _add_para(doc, f"• Tub de protectie: {tub_protectie}. "
                   f"Pat de caramizi: L = {_get(data, 'br_pat_caramizi_l_m', '1')} m × "
                   f"{_get(data, 'br_pat_caramizi_lat_m', '0,4')} m")

    _add_para(doc, "")
    _add_para(doc, "Prezentul proiect s-a intocmit in conformitate cu normele ORD. A.N.R.E. 89/2018, "
                   "Legea 10/1995, Legea energiei electrice si a gazelor naturale 123/2012, si a "
                   "legislatiei tehnice in vigoare.", italic=True)

    _add_heading(doc, "3. Documente prezentate la avizare", level=2)
    _add_para(doc, "• Memoriul elaborat de proiectant precum si breviarul de calcul;")
    _add_para(doc, "• Instructiuni de protectia muncii (S.S.M.) si de protectia mediului;")
    _add_para(doc, "• Piese desenate conform borderou (schema izometrica, sectiune transversala si detalii piese);")

    _add_heading(doc, "4. Concluzii asupra verificarii", level=2)
    _add_para(doc, "In urma verificarii proiectul se considera corespunzator, semnandu-se si stampilandu-se conform indrumatorului.", italic=True)

    doc.add_paragraph()
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"
    table.cell(0, 0).text = f"Am primit {_get(data, 'nr_exemplare', '4')} exemplare"
    table.cell(0, 1).text = f"Am predat {_get(data, 'nr_exemplare', '4')} exemplare"
    table.cell(1, 0).text = f"Proiectant\n{_get(data, 'fact_proiectanta_nume')}"
    table.cell(1, 1).text = f"Verificator tehnic atestat,\n{_get(data, 'vgd_verificator_nume')}"

    doc.add_page_break()

    # ====================================================
    # SECȚIUNEA 2 — FOAIE DE CAPAT
    # ====================================================
    _add_centered(doc, "PROIECT", bold=True, size=22)
    _add_centered(doc, f"NR: {_get(data, 'proiect_nr', '____')}/{datetime.now().year}", bold=True, size=14)
    doc.add_paragraph()
    _add_centered(doc, "DENUMIRE:", bold=True, size=12)
    _add_centered(doc, _get(data, "tipul_lucrarii", "Bransament gaze naturale (racord) nou proiectat"), bold=False, size=12)
    doc.add_paragraph()
    _add_centered(doc, f"BENEFICIAR: {_get(data, 'beneficiar_investitiei', 'S.C. ENGIE Romania S.A.')}", bold=True)
    _add_centered(doc, f"Sediu social: {_get(data, 'beneficiar_sediu', 'Str. Marasesti, Nr. 4-6, CORP B, Bucuresti, Sector 4')}")
    doc.add_paragraph()
    _add_centered(doc, f"CLIENT: {_get(data, 'nume_client')}", bold=True)
    doc.add_paragraph()
    _add_centered(doc, "AMPLASAMENTUL LUCRARII:", bold=True)
    _add_centered(doc, f"– {_get(data, 'amplasament_strada')}, {_get(data, 'amplasament_localitate')}, {_get(data, 'amplasament_judet')}")
    doc.add_paragraph()
    _add_centered(doc, "AMPLASAMENTUL IMOBILULUI DE CONSUM:", bold=True)
    _add_centered(doc, f"– {_get(data, 'adresa_imobil') or _get(data, 'amplasament_strada')}, "
                       f"Nr. {_get(data, 'loc_consum_numar', '____')}, "
                       f"{_get(data, 'amplasament_localitate')}, {_get(data, 'amplasament_judet')}")
    doc.add_paragraph()
    _add_centered(doc, f"PROIECTANT GENERAL: {_get(data, 'proiectant_general_firma', COMPANY['name'])}", bold=True)
    _add_centered(doc, f"Instalator aut.: {_get(data, 'fact_proiectanta_nume')}")
    _add_centered(doc, f"Grad {_get(data, 'fact_proiectanta_autorizatie', 'P.G.D.')}, "
                       f"nr. aut. {_get(data, 'fact_proiectanta_legitimatie')}")
    doc.add_paragraph()
    _add_centered(doc, f"CONSTRUCTOR/EXECUTANT: {_get(data, 'exec_firma', _get(data, 'fact_executanta_societate', COMPANY['name']))}", bold=True)
    _add_centered(doc, f"Instalator aut.: {_get(data, 'fact_executanta_nume')}")
    _add_centered(doc, f"Grad {_get(data, 'fact_executanta_autorizatie', 'E.G.D.')}, "
                       f"nr. aut. {_get(data, 'fact_executanta_legitimatie')}")
    doc.add_paragraph()
    _add_centered(doc, f"FAZA: {_get(data, 'faza_proiect', 'P.A.C. + P.T.H. + D.E.')}", bold=True)

    doc.add_page_break()

    # ====================================================
    # SECȚIUNEA 3 — BORDEROU
    # ====================================================
    _add_centered(doc, "BORDEROU", bold=True, size=18)
    doc.add_paragraph()
    _add_heading(doc, "PIESE SCRISE", level=2)
    piese_scrise = [
        "Referat;", "Foaie de capat;", "Borderou documente cu paginile numerotate;",
        f"Ordin de lucru, Nr.: {_get(data, 'osd_ordin_nr')};",
        f"Acord acces la sistemul de distributie gaze naturale, Nr.: {_get(data, 'atr_numar')};",
        "Specificatii tehnice necesare stabilirii solutiei de alimentare;",
        "Schita solutie de alimentare;", "Specificatii pozitionare firida;",
    ]
    if data.get("are_cu_existent") in {"Da", True}:
        piese_scrise.append(f"Certificat de urbanism, Nr.: {_get(data, 'cu_numar')};")
        piese_scrise.append("Avize cerute prin Certificatul de urbanism;")
    piese_scrise.extend([
        "Devizul general al lucrarilor;", "Memoriu Tehnic;", "Breviar de calcul;",
        "Programul de control al calitatii lucrarilor pe faze determinante;",
        "Lista cantitatilor de lucrari (ANEXA 14);", "Caiet de sarcini;",
    ])
    for item in piese_scrise:
        p = doc.add_paragraph(item, style="List Bullet")
        p.runs[0].font.size = Pt(11)

    doc.add_paragraph()
    _add_heading(doc, "PIESE DESENATE", level=2)
    for item in [
        "Plan de incadrare in zona;",
        "Plan de situatie si amplasament al lucrarilor (scara 1:500);",
        "Schema izometrica;",
        "Profil longitudinal;",
        "Detalii piese (sectiune transversala santuri, traversari, refacere);",
    ]:
        p = doc.add_paragraph(item, style="List Bullet")
        p.runs[0].font.size = Pt(11)

    doc.add_page_break()

    # ====================================================
    # SECȚIUNEA 4 — MEMORIU TEHNIC
    # ====================================================
    _add_centered(doc, "MEMORIU TEHNIC", bold=True, size=18)
    doc.add_paragraph()

    _add_heading(doc, "1. DATE GENERALE", level=2)
    _add_para(doc, f"Denumirea lucrarii: {_get(data, 'tipul_lucrarii', 'Bransament nou proiectat (racord) de gaze naturale')}.")
    _add_para(doc, "Amplasamentul imobilului pentru care s-a solicitat racordarea:")
    _add_para(doc, f"   {_get(data, 'adresa_imobil') or _get(data, 'amplasament_strada')}, "
                   f"Nr. {_get(data, 'loc_consum_numar', '____')}, "
                   f"{_get(data, 'amplasament_localitate')}, {_get(data, 'amplasament_judet')}")
    _add_para(doc, "Amplasamentul lucrarii:")
    _add_para(doc, f"   {_get(data, 'amplasament_strada')}, {_get(data, 'amplasament_localitate')}, {_get(data, 'amplasament_judet')}")
    _add_para(doc, f"Beneficiarul investitiei: {_get(data, 'beneficiar_investitiei', 'ENGIE Romania S.A.')}")
    _add_para(doc, f"Client: {_get(data, 'nume_client')}")
    _add_para(doc, f"Proiectant: {_get(data, 'proiectant_general_firma', COMPANY['name'])}")
    _add_para(doc, f"Executant: {_get(data, 'exec_firma', _get(data, 'fact_executanta_societate', COMPANY['name']))}")

    _add_heading(doc, "2. NECESITATEA SI OPORTUNITATEA LUCRARII", level=2)
    necesitate_text = (
        f"In {_get(data, 'amplasament_strada')}, Nr. {_get(data, 'loc_consum_numar', '____')}, "
        f"{_get(data, 'amplasament_localitate')}, {_get(data, 'nume_client')} detine un imobil "
        f"pentru care a solicitat alimentarea cu gaze naturale.\n\n"
        f"Pentru alimentarea cu gaze naturale a imobilului este necesara realizarea unui bransament nou, "
        f"individual, {_get(data, 'br_presiune', 'REDUSA').upper()} PRESIUNE, din "
        f"{'PE 100 SDR 11' if (data.get('br_material') or '').lower().startswith('poliet') else 'OL'}, "
        f"Dn {_get(data, 'br_diametru_dn', '63 mm')}, cu lungimea de {_get(data, 'br_lungime_m', '4')} m "
        f"(3m + 1m raiser), pozat la adancimea {_get(data, 'br_adancime_pozare_m', '0,90')} m.\n\n"
        f"Bransamentul va fi racordat la conducta existenta de "
        f"{'PE 100 SDR 11' if (data.get('cnd_ex_material') or '').lower().startswith('poliet') else 'OL'} "
        f"Dn {_get(data, 'cnd_ex_diametru_dn', '90 mm')}, situata pe {_get(data, 'cnd_ex_amplasament', _get(data, 'amplasament_strada'))}."
    )
    _add_para(doc, necesitate_text)

    _add_heading(doc, "3. DESCRIEREA LUCRARILOR", level=2)
    _add_heading(doc, "3.1. Amplasamentul lucrarilor", level=3)
    _add_para(doc, f"Lucrarile se vor amplasa pe teritoriul administrativ al {_get(data, 'amplasament_localitate', 'localitatii')}, "
                   f"{_get(data, 'amplasament_judet')}, respectiv pe {_get(data, 'amplasament_strada')}.")
    _add_para(doc, "Bransamentul proiectat va respecta prevederile «Normelor tehnice pentru proiectarea, "
                   "executarea si exploatarea sistemelor de alimentare cu gaze naturale – ORD ANRE 89/2018», "
                   "ale legislatiei in vigoare.")
    _add_para(doc, "Bransamentul va fi amplasat in domeniul public, conform planului de amplasare "
                   "(scara 1:2000) si a planului de situatie (scara 1:500), anexate.")

    _add_heading(doc, "3.2. Categoria de importanta a lucrarilor", level=3)
    _add_para(doc, f"Conform HG 766/1997, categoria de importanta este: «{_get(data, 'cladire_cat_importanta', 'C - normala')}».")
    _add_para(doc, "Executia se va realiza cu materiale C1 (practice neinflamabile) si va avea gradul I de "
                   "rezistenta la foc.")

    doc.add_page_break()

    # ====================================================
    # SECȚIUNEA 5 — ANEXA 14 LISTA MATERIALE (AUTO-GENERATED)
    # ====================================================
    _add_centered(doc, "ANEXA 14", bold=True, size=14)
    _add_centered(doc, "LISTA MATERIALE", bold=True, size=16)
    _add_para(doc, f"Denumirea lucrarii: {_get(data, 'tipul_lucrarii')} — "
                   f"{_get(data, 'amplasament_strada')}, Nr. {_get(data, 'loc_consum_numar', '___')}")
    doc.add_paragraph()

    rows = materials_db.build_materials_table(data)
    if rows:
        # V10.6.3 — Show SAP column only if at least one row has a SAP code (Distrigaz)
        has_sap = any(r.get("sap_code") for r in rows)
        n_cols = 7 if has_sap else 6
        tbl = doc.add_table(rows=len(rows) + 1, cols=n_cols)
        tbl.style = "Light Grid Accent 1"
        hdr = tbl.rows[0]
        headers = (["Nr.", "Cod SAP", "Denumire material", "BR/CND", "Cant.", "UM", "Obs."]
                   if has_sap else
                   ["Nr.", "Denumire material (specificație tehnică)", "BR/CND", "Cant.", "UM", "Obs."])
        for i, label in enumerate(headers):
            hdr.cells[i].text = label
            for run in hdr.cells[i].paragraphs[0].runs:
                run.bold = True
        # Data rows
        for ri, row in enumerate(rows, start=1):
            cells = tbl.rows[ri].cells
            if has_sap:
                cells[0].text = str(row["nr"])
                cells[1].text = row.get("sap_code", "")
                cells[2].text = row["desc"]
                cells[3].text = row["dest"]
                cells[4].text = str(row["qty"])
                cells[5].text = row["um"]
                cells[6].text = ""
            else:
                cells[0].text = str(row["nr"])
                cells[1].text = row["desc"]
                cells[2].text = row["dest"]
                cells[3].text = str(row["qty"])
                cells[4].text = row["um"]
                cells[5].text = ""
        doc.add_paragraph()
        _add_para(doc, f"Total materiale: {len(rows)} poziții (calcul automat după specificațiile tehnice ale proiectului).", italic=True)
        if not has_sap:
            _add_para(doc, f"Operatorul Sistemului de Distribuție ({_get(data, 'osd_operator', '—')}) va completa codurile specifice de catalog la momentul predării.", italic=True)
    else:
        _add_para(doc, "Niciun material selectat. Completează în Studio: br_material, br_diametru_dn, br_lungime_m.", italic=True)

    # Footer
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(f"Generat la {_today_ro()} de {COMPANY['name']} (V10.5 EPD)").italic = True

    # Serialize
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
