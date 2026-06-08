"""Apă-Canal — Template-uri DOCX concrete pentru branșament apă potabilă + racord canalizare.

Replică pattern-ul `gas_doc_templates.py` adaptat pentru:
- Branșament apă potabilă (DN 25/32 PEHD/PVC)
- Racord canalizare menajeră (DN 110/160/200 PVC/PP)
- Norme: Legea 241/2006 + STAS 8591 + SR EN 805 + SR EN 752
- Operatori: Apa Nova București, Apa Brașov, RAJA Constanța, Compania de Apă Cluj, etc.

5 template-uri concrete.
"""
from __future__ import annotations
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

import gas_doc_templates as base


def _hdr(doc, title, subtitle=""):
    base._company_header(doc)
    base._add_heading(doc, title, level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    if subtitle:
        base._add_para(doc, subtitle, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
    doc.add_paragraph()


def cerere_bransament_apa(proj: Dict[str, Any]) -> bytes:
    data = proj.get("data") or {}
    doc = Document()
    _hdr(doc, "CERERE BRANȘAMENT APĂ POTABILĂ",
         "(Conform Legea 241/2006 + STAS 8591 + Regulament local)")
    base._add_para(doc, f"Către {base._get(data,'apa_canal_operator','OPERATOR APĂ-CANAL')}", bold=True)
    doc.add_paragraph()
    base._add_kv_table(doc, [
        ("Beneficiar", base._get(data, "beneficiar_nume")),
        ("CNP/CUI", base._get(data, "beneficiar_cnp_cui")),
        ("Adresă loc consum", base._get(data, "loc_consum_adresa")),
        ("Nr. cadastral", base._get(data, "loc_consum_cadastru")),
        ("Tip locuință / obiectiv", base._get(data, "tip_consumator")),
        ("Număr locatari / persoane", base._get(data, "ac_nr_persoane", "—")),
        ("Debit estimat (l/s)", base._get(data, "ac_debit_estimat_ls")),
        ("Diametru branșament solicitat (DN)", base._get(data, "ac_dn_apa", "DN 25 PEHD")),
        ("Tip apometru", base._get(data, "ac_tip_apometru", "Apometru clasa C, DN 20")),
    ])
    doc.add_paragraph()
    base._add_para(doc, "Anexe:", bold=True)
    for ln in [
        "• Copie act proprietate / contract",
        "• Plan încadrare + plan situație 1:500",
        "• Schema instalație interioară apă caldă/rece",
        "• Memoriu tehnic justificativ",
        "• Copie CI / CUI",
    ]:
        doc.add_paragraph(ln, style="List Bullet")
    base._footer_signature(doc, proj, data)
    return base._save(doc)


def cerere_racord_canalizare(proj: Dict[str, Any]) -> bytes:
    data = proj.get("data") or {}
    doc = Document()
    _hdr(doc, "CERERE RACORD CANALIZARE MENAJERĂ",
         "(Conform Legea 241/2006 + SR EN 752)")
    base._add_para(doc, f"Către {base._get(data,'apa_canal_operator','OPERATOR APĂ-CANAL')}", bold=True)
    doc.add_paragraph()
    base._add_kv_table(doc, [
        ("Beneficiar", base._get(data, "beneficiar_nume")),
        ("Adresă", base._get(data, "loc_consum_adresa")),
        ("Tip canalizare", base._get(data, "ac_tip_canalizare", "Menajeră")),
        ("Diametru racord (DN)", base._get(data, "ac_dn_canal", "DN 160 PVC SN8")),
        ("Pantă racord (%)", base._get(data, "ac_panta_canal", "2-3%")),
        ("Tip cămin", base._get(data, "ac_tip_camin", "Cămin tip C1, DN 1000")),
        ("Adâncime racord (cm)", base._get(data, "ac_adancime_canal", "80-120")),
        ("Debit estimat ape uzate (l/s)", base._get(data, "ac_debit_apa_uzata_ls")),
    ])
    base._footer_signature(doc, proj, data)
    return base._save(doc)


def memoriu_tehnic_ac(proj: Dict[str, Any]) -> bytes:
    data = proj.get("data") or {}
    doc = Document()
    _hdr(doc, "MEMORIU TEHNIC JUSTIFICATIV (APĂ-CANAL)",
         "(Conform STAS 1342 + STAS 1846 + SR EN 805/752)")
    base._add_kv_table(doc, [
        ("Beneficiar", base._get(data, "beneficiar_nume")),
        ("Adresă", base._get(data, "loc_consum_adresa")),
        ("Debit calculat apă (l/s)", base._get(data, "ac_debit_estimat_ls")),
        ("Coeficient simultaneitate Ks", base._get(data, "ac_ks_apa", "0.6")),
        ("Presiune disponibilă rețea (bar)", base._get(data, "ac_presiune_disponibila_bar", "3.5")),
        ("Pierdere presiune calculată (bar)", base._get(data, "ac_pierdere_calc_bar")),
    ])
    doc.add_paragraph()
    base._add_para(doc, "Calcule hidraulice:", bold=True)
    base._add_para(doc, "Q_calc = Σ Q_obiect × Ks · Δp = λ × (L/D) × (ρv²/2) · v < 2 m/s pentru apă potabilă",
                   italic=True)
    base._footer_signature(doc, proj, data)
    return base._save(doc)


def pv_receptie_ac(proj: Dict[str, Any]) -> bytes:
    data = proj.get("data") or {}
    doc = Document()
    _hdr(doc, "PROCES-VERBAL DE RECEPȚIE BRANȘAMENT APĂ-CANAL",
         "(Conform HG 273/1994 + STAS 8591)")
    base._add_kv_table(doc, [
        ("Nr. PV", base._get(data, "receptie_pv_numar")),
        ("Data", base._get(data, "receptie_pv_data")),
        ("Proba presiune (bar)", base._get(data, "ac_proba_presiune_bar", "10")),
        ("Durată probă (h)", base._get(data, "ac_proba_durata_h", "2")),
        ("Etanșeitate canalizare", base._get(data, "ac_etanseitate", "Verificat — admis")),
        ("Rezultat", base._get(data, "proba_rezultat", "Admis")),
    ])
    base._footer_signature(doc, proj, data)
    return base._save(doc)


def carte_tehnica_ac(proj: Dict[str, Any]) -> bytes:
    data = proj.get("data") or {}
    doc = Document()
    _hdr(doc, "CARTEA TEHNICĂ INSTALAȚIE APĂ-CANAL",
         "(Conform HG 273/1994)")
    for sec, key in [
        ("A. Proiectare", "ct_sectiune_A_continut"),
        ("B. Execuție", "ct_sectiune_B_continut"),
        ("C. Recepție", "ct_sectiune_C_continut"),
        ("D. Exploatare", "ct_sectiune_D_continut"),
    ]:
        base._add_heading(doc, sec, level=2)
        base._add_para(doc, base._get(data, key, "(de completat)"), italic=True)
        doc.add_paragraph()
    base._footer_signature(doc, proj, data)
    return base._save(doc)


TEMPLATES = {
    "ac_cerere_bransament": {"label": "Cerere Branșament Apă", "phase": "atr", "fn": cerere_bransament_apa, "norm": "Legea 241/2006 + STAS 8591"},
    "ac_cerere_racord":     {"label": "Cerere Racord Canalizare", "phase": "atr", "fn": cerere_racord_canalizare, "norm": "Legea 241/2006 + SR EN 752"},
    "ac_memoriu_tehnic":    {"label": "Memoriu Tehnic Apă-Canal", "phase": "dtac", "fn": memoriu_tehnic_ac, "norm": "STAS 1342 + STAS 1846"},
    "ac_pv_receptie":       {"label": "PV Recepție Apă-Canal", "phase": "receptie", "fn": pv_receptie_ac, "norm": "HG 273/1994"},
    "ac_carte_tehnica":     {"label": "Carte Tehnică Apă-Canal", "phase": "receptie", "fn": carte_tehnica_ac, "norm": "HG 273/1994"},
}


def list_templates() -> List[Dict[str, Any]]:
    return [{"id": k, **{kk: vv for kk, vv in v.items() if kk != "fn"}} for k, v in TEMPLATES.items()]


def generate(template_id: str, proj: Dict[str, Any]) -> Optional[Tuple[bytes, str]]:
    tpl = TEMPLATES.get(template_id)
    if not tpl:
        return None
    fname = f"{template_id}_{(proj.get('title','proiect')[:40]).replace(' ','_')}.docx"
    return tpl["fn"](proj), fname
