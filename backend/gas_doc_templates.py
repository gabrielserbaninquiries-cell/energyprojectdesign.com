"""Gas Natural — Template-uri DOCX cu placeholder + condiționale IF pentru
documentația tehnică COMPLETĂ (Energy Project Design SRL).

Acest modul transformă 10 ore de muncă manuală în 30 de minute. Folosește
`python-docx` programatic pentru a genera 8 template-uri legale conforme
NTPEE 2018 + HG 907/2016 + Legea 50/1991 + Ord. ANRE 89/2018 + 162/2021.

Toate placeholder-ele suportă logică condițională simplă:
- {{var}} — înlocuit cu valoarea câmpului
- {{var | default("—")}} — fallback dacă vid
- {{#if var}}...{{/if}} — afișează blocul doar dacă `var` truthy
- {{#if var == "valoare"}}...{{/if}} — comparație directă
- {{calc.recommended_dn}} — output din calc_engine

Template-uri implementate:
1. cerere_cu       — Cerere Certificat de Urbanism (către Primărie)
2. cerere_atr      — Cerere Aviz Tehnic Racordare (către OSD)
3. memoriu_tehnic  — Memoriu Tehnic Justificativ (DTAC/PT)
4. caiet_sarcini   — Caiet de sarcini execuție
5. borderou        — Borderou piese scrise + desenate
6. cerere_pif      — Cerere Punere în Funcțiune (către OSD)
7. pv_receptie     — Proces Verbal de Recepție la Terminarea Lucrărilor
8. carte_tehnica   — Cartea Tehnică a Construcției (4 secțiuni)
"""
from __future__ import annotations
import io
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

import gas_calc_engine as engine

# ============================================================================
# COMPANY HEADER — Energy Project Design SRL
# ============================================================================
COMPANY = {
    "name": "ENERGY PROJECT DESIGN S.R.L.",
    "cui": "43151074",
    "reg_com": "J40/12982/2020",
    "address": "Str. Lt. Alexandru Popescu nr. 9B, Sectorul 3, București",
    "phone": "+40 723 000 000",
    "email": "office@energyprojectdesign.com",
    "web": "energyprojectdesign.com",
    "anre_proiectant": "PDD/2022/0001",  # Atestat ANRE Proiectare Distribuție Distribuție gaze
    "anre_executant": "EDD/2022/0001",   # Atestat ANRE Execuție Distribuție gaze
}


# ----------------------------------------------------------------------------
# Helpers
# ----------------------------------------------------------------------------
def _today_ro() -> str:
    months = ["ianuarie", "februarie", "martie", "aprilie", "mai", "iunie",
              "iulie", "august", "septembrie", "octombrie", "noiembrie", "decembrie"]
    now = datetime.now()
    return f"{now.day:02d} {months[now.month-1]} {now.year}"


def _get(data: Dict[str, Any], key: str, default: str = "—") -> str:
    v = data.get(key)
    if v in (None, "", []):
        return default
    return str(v)


def _truthy(data: Dict[str, Any], key: str) -> bool:
    v = data.get(key)
    if v in (None, "", [], 0, "0"):
        return False
    return True


def _add_heading(doc: Document, text: str, level: int = 1, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    h = doc.add_heading(text, level=level)
    h.alignment = align


def _add_para(doc: Document, text: str, bold: bool = False, italic: bool = False,
              size: int = 11, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)


def _add_kv_table(doc: Document, rows: List[Tuple[str, str]], col1_cm: float = 5.5) -> None:
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (k, v) in enumerate(rows):
        cells = t.rows[i].cells
        cells[0].text = k
        cells[1].text = v
        for run in cells[0].paragraphs[0].runs:
            run.bold = True


def _company_header(doc: Document) -> None:
    """Block antet companie."""
    _add_para(doc, COMPANY["name"], bold=True, size=12)
    _add_para(doc, f"CUI {COMPANY['cui']} · {COMPANY['reg_com']}", size=9)
    _add_para(doc, COMPANY["address"], size=9)
    _add_para(doc, f"{COMPANY['phone']} · {COMPANY['email']} · {COMPANY['web']}", size=9)
    _add_para(doc, f"Atestat ANRE Proiectare gaze: {COMPANY['anre_proiectant']}", size=9, italic=True)
    doc.add_paragraph()


def _footer_signature(doc: Document, proj: Dict[str, Any], data: Dict[str, Any]) -> None:
    doc.add_paragraph()
    doc.add_paragraph()
    t = doc.add_table(rows=2, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.rows[0].cells[0].text = "Întocmit (Proiectant)"
    t.rows[0].cells[1].text = "Executant"
    t.rows[0].cells[2].text = "Verificat (VGD)"
    proiectant = _get(data, "dtac_proiectant_specialitate", COMPANY["name"])
    proiectant_firma = _get(data, "proiectant_general_firma", proiectant)
    proiectant_nr = _get(data, "proiectant_aut_nr", _get(data, "dtac_atestat_proiectant", COMPANY["anre_proiectant"]))
    proiectant_grad = _get(data, "proiectant_aut_grad", "Iea (instalații exterioare apă-gaze)")
    executant_firma = _get(data, "exec_firma")
    executant_nr = _get(data, "executant_aut_nr", "—")
    executant_grad = _get(data, "executant_aut_grad", "EDD (Execuție Distribuție Gaze)")
    vgd = _get(data, "dtac_verificator_vgd", "—")
    verificator_leg = _get(data, "verificator_legitimatie_nr", "—")

    t.rows[1].cells[0].text = (
        f"{proiectant_firma}\n"
        f"({proiectant})\n"
        f"Autorizație ANRE: {proiectant_nr} ({proiectant_grad})\n\n"
        f"Semnătură: ______________"
    )
    t.rows[1].cells[1].text = (
        f"{executant_firma}\n"
        f"Autorizație ANRE: {executant_nr} ({executant_grad})\n"
        f"RTE: {_get(data, 'exec_responsabil_tehnic', '—')}\n\n"
        f"Semnătură: ______________"
    )
    t.rows[1].cells[2].text = (
        f"{vgd}\n"
        f"Legitimație nr.: {verificator_leg}\n\n"
        f"Semnătură: ______________"
    )
    doc.add_paragraph()
    if proj.get("signature_hash"):
        _add_para(doc, f"Hash SHA-256: {proj['signature_hash']}", italic=True, size=8)
        _add_para(doc, f"Semnat digital: {proj.get('signed_at','—')}", italic=True, size=8)
    _add_para(doc, f"Document generat: {_today_ro()} · Energy Project Design", italic=True, size=8)


def _project_cartouche(doc: Document, proj: Dict[str, Any], data: Dict[str, Any]) -> None:
    """Cartouche tehnic plasat după antet — consumă identificatorii proiectului.

    Apelat în template-urile principale (memoriu, caiet, borderou, carte tehnică)
    pentru a afișa identificatori legali completi.
    """
    rows = [
        ("Denumire lucrare", _get(data, "denumire_lucrare_extinsa", proj.get("title", "—"))),
        ("Tip lucrare", _get(data, "tipul_lucrarii")),
        ("Faza de proiectare", _get(data, "faza_proiectare", "DTAC + PTH")),
        ("Nr. proiect / an", _get(data, "proiect_nr_an")),
        ("Amplasament lucrare", _get(data, "amplasament_lucrare", _get(data, "loc_consum_adresa"))),
        ("Loc consum (imobil beneficiar)", _get(data, "amplasament_imobil_consum", _get(data, "loc_consum_adresa"))),
        ("Categorie presiune", _get(data, "presiune_categorie")),
        ("DN canonic conductă", _get(data, "sf_diametru_nominal_DN")),
        ("Ordin lucru OSD", _get(data, "ordin_lucru_nr_data")),
        ("Notificare ISC", _get(data, "isc_nr_inreg")),
        ("Dispoziție de șantier necesară", _get(data, "dispozitie_necesara", "Nu")),
        ("Coduri materiale (catalog OSD)", _get(data, "materiale_catalog_codes")),
    ]
    visible = [(k, v) for k, v in rows if v not in ("—", "", None)]
    if not visible:
        return
    _add_para(doc, "CARTOUCHE PROIECT", bold=True, size=9)
    _add_kv_table(doc, visible, col1_cm=6.5)
    doc.add_paragraph()


def _save(doc: Document) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ============================================================================
# TEMPLATE 1: Cerere CU (Certificat de Urbanism)
# ============================================================================
def cerere_cu(proj: Dict[str, Any]) -> bytes:
    data = proj.get("data") or {}
    doc = Document()
    _company_header(doc)
    _add_para(doc, f"Către PRIMĂRIA {_get(data,'loc_consum_localitate','—').upper()}", bold=True, size=12)
    _add_para(doc, f"Județul {_get(data,'loc_consum_judet','—')}", size=10)
    doc.add_paragraph()
    _add_heading(doc, "CERERE PENTRU EMITEREA CERTIFICATULUI DE URBANISM", level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_para(doc, "(Conform Legii 50/1991 republicată, art. 6)", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
    doc.add_paragraph()

    _add_para(doc, "Subsemnatul/Subscrisa:", bold=True)
    _add_kv_table(doc, [
        ("Nume / Denumire beneficiar", _get(data, "beneficiar_nume")),
        ("CNP / CUI", _get(data, "beneficiar_cnp_cui")),
        ("Adresa fiscală", _get(data, "beneficiar_adresa")),
        ("Telefon", _get(data, "beneficiar_telefon")),
        ("Email", _get(data, "beneficiar_email")),
    ])

    doc.add_paragraph()
    _add_para(doc, "Solicit emiterea Certificatului de Urbanism în vederea:", bold=True)
    scop = _get(data, "scop_lucrare", "Branșament nou la rețeaua de gaze naturale")
    _add_para(doc, scop, italic=True)

    doc.add_paragraph()
    _add_para(doc, "Pentru imobilul situat la:", bold=True)
    _add_kv_table(doc, [
        ("Adresă loc de consum", _get(data, "loc_consum_adresa")),
        ("Strada și număr", _get(data, "loc_consum_strada")),
        ("Localitate", _get(data, "loc_consum_localitate")),
        ("Județ", _get(data, "loc_consum_judet")),
        ("Nr. cadastral / CF", _get(data, "loc_consum_cadastru")),
    ])

    doc.add_paragraph()
    _add_para(doc, "Date tehnice estimate:", bold=True)
    _add_kv_table(doc, [
        ("Tip consumator", _get(data, "tip_consumator")),
        ("Regim funcționare", _get(data, "regim_functionare")),
        ("Debit instalat estimat (m³/h)", _get(data, "debit_instalat_mc_h")),
        ("Consum anual estimat (m³/an)", _get(data, "consum_anual_mc")),
    ])

    doc.add_paragraph()
    _add_para(doc, "Anexe:", bold=True)
    doc.add_paragraph("• Extras CF / Act proprietate (copie)", style="List Bullet")
    doc.add_paragraph("• Plan de încadrare în zonă (scara 1:5000)", style="List Bullet")
    doc.add_paragraph("• Plan de situație (scara 1:500)", style="List Bullet")
    doc.add_paragraph("• Memoriu de necesitate", style="List Bullet")

    _footer_signature(doc, proj, data)
    return _save(doc)


# ============================================================================
# TEMPLATE 2: Cerere ATR (Aviz Tehnic Racordare) către OSD
# ============================================================================
def cerere_atr(proj: Dict[str, Any]) -> bytes:
    data = proj.get("data") or {}
    osd = _get(data, "atr_osd", "Distrigaz Sud Rețele")
    doc = Document()
    _company_header(doc)
    _add_para(doc, f"Către {osd.upper()}", bold=True, size=12)
    _add_para(doc, "OSD — Operator Sistem Distribuție gaze naturale", size=10, italic=True)
    doc.add_paragraph()

    _add_heading(doc, "CERERE PENTRU EMITEREA AVIZULUI TEHNIC DE RACORDARE (ATR)", level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_para(doc, "(Conform Ord. ANRE 89/2018 — Procedura de acces la sistemul de distribuție gaze naturale)",
              italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
    doc.add_paragraph()

    _add_para(doc, "Solicitant (beneficiar/titular contract):", bold=True)
    _add_kv_table(doc, [
        ("Nume/Denumire", _get(data, "beneficiar_nume")),
        ("CNP/CUI", _get(data, "beneficiar_cnp_cui")),
        ("Adresă fiscală", _get(data, "beneficiar_adresa")),
        ("Telefon", _get(data, "beneficiar_telefon")),
        ("Email", _get(data, "beneficiar_email")),
    ])

    doc.add_paragraph()
    _add_para(doc, "Loc consum (punct racordare):", bold=True)
    _add_kv_table(doc, [
        ("Adresă", _get(data, "loc_consum_adresa")),
        ("Localitate", _get(data, "loc_consum_localitate")),
        ("Județ", _get(data, "loc_consum_judet")),
        ("Cadastru/CF", _get(data, "loc_consum_cadastru")),
    ])

    doc.add_paragraph()
    _add_para(doc, "Date tehnice solicitate:", bold=True)
    tip = _get(data, "tip_consumator")
    debit_calc = ""
    nr_cons = data.get("nr_consumatori_simultani")
    debit_ind = data.get("debit_individual_mc_h")
    # Auto-calc debit total cu Ks (logica IF: dacă sunt consumatori multipli)
    if nr_cons and debit_ind:
        try:
            r = engine.debit_calculat(int(nr_cons), float(debit_ind), tip.lower() if tip else "casnic")
            debit_calc = f"{r['debit_calculat_mc_h']} m³/h (Ks={r['ks']})"
        except (ValueError, TypeError):
            pass

    _add_kv_table(doc, [
        ("Tip consumator", tip),
        ("Regim funcționare", _get(data, "regim_functionare")),
        ("Debit instalat (m³/h)", _get(data, "debit_instalat_mc_h")),
        ("Debit calculat cu Ks", debit_calc or "—"),
        ("Consum anual (m³/an)", _get(data, "consum_anual_mc")),
        ("Presiune solicitată", _get(data, "sf_presiune_max_op_bar") + " bar" if _truthy(data, "sf_presiune_max_op_bar") else "Joasă presiune (≤100 mbar)"),
    ])

    # Bloc condițional: dacă există SF cu soluție tehnică, includ
    if _truthy(data, "sf_solutie_tehnica"):
        doc.add_paragraph()
        _add_para(doc, "Soluție tehnică propusă (din Studiul de Fezabilitate):", bold=True)
        _add_para(doc, _get(data, "sf_solutie_tehnica"))

    doc.add_paragraph()
    _add_para(doc, "Anexe la cerere:", bold=True)
    doc.add_paragraph("• Copie CI/CUI beneficiar", style="List Bullet")
    doc.add_paragraph("• Extras CF / Act proprietate", style="List Bullet")
    doc.add_paragraph("• Plan situație 1:500 (cu poziția propusă pentru branșament)", style="List Bullet")
    doc.add_paragraph("• Plan încadrare 1:5000", style="List Bullet")
    if _truthy(data, "cu_numar"):
        doc.add_paragraph(f"• Copie Certificat Urbanism nr. {_get(data,'cu_numar')} / {_get(data,'cu_data_emitere','—')}", style="List Bullet")

    doc.add_paragraph()
    _add_para(doc, f"Data: {_today_ro()}", italic=True)

    _footer_signature(doc, proj, data)
    return _save(doc)


# ============================================================================
# TEMPLATE 3: Memoriu Tehnic Justificativ (DTAC + PT)
# ============================================================================
def memoriu_tehnic(proj: Dict[str, Any]) -> bytes:
    data = proj.get("data") or {}
    doc = Document()
    _company_header(doc)
    _project_cartouche(doc, proj, data)

    _add_heading(doc, "MEMORIU TEHNIC JUSTIFICATIV", level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_para(doc, "Proiect distribuție gaze naturale",
              italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
    _add_para(doc, f"Subdomeniu: {proj.get('subdomain', 'bransament-casnic').replace('-', ' ').title()}",
              italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
    doc.add_paragraph()

    _add_heading(doc, "1. Date generale", level=2)
    _add_kv_table(doc, [
        ("Beneficiar", _get(data, "beneficiar_nume")),
        ("Loc consum", f"{_get(data,'loc_consum_adresa')}, {_get(data,'loc_consum_localitate')}, jud. {_get(data,'loc_consum_judet')}"),
        ("Cadastru/CF", _get(data, "loc_consum_cadastru")),
        ("Proiectant", _get(data, "dtac_proiectant_specialitate", COMPANY["name"])),
        ("Atestat ANRE proiectant", _get(data, "dtac_atestat_proiectant", COMPANY["anre_proiectant"])),
        ("Tip consumator", _get(data, "tip_consumator")),
        ("Scop lucrare", _get(data, "scop_lucrare")),
    ])

    _add_heading(doc, "2. Cadru legal aplicabil", level=2)
    doc.add_paragraph("• Legea 123/2012 — Legea energiei electrice și a gazelor naturale", style="List Bullet")
    doc.add_paragraph("• NTPEE 2018 — Normele tehnice pentru proiectarea, executarea și exploatarea sistemelor de alimentare cu gaze naturale", style="List Bullet")
    doc.add_paragraph("• HG 907/2016 — Etapele de elaborare a documentațiilor tehnico-economice", style="List Bullet")
    doc.add_paragraph("• Legea 50/1991 — Autorizarea executării lucrărilor de construcții", style="List Bullet")
    doc.add_paragraph("• Ord. ANRE 89/2018 — Procedura de acces la sistemul de distribuție", style="List Bullet")
    doc.add_paragraph("• Ord. ANRE 162/2021 — Punere în funcțiune (PIF)", style="List Bullet")

    _add_heading(doc, "3. Soluție tehnică propusă", level=2)
    if _truthy(data, "sf_solutie_tehnica"):
        _add_para(doc, _get(data, "sf_solutie_tehnica"))
    else:
        _add_para(doc, "Se propune realizarea unui branșament/extindere cu conductă PE 100 SDR 11, "
                       "pozată îngropat conform NTPEE 2018 art. 56, cu trecere de la presiune medie la "
                       "presiune joasă printr-un post de reglare amplasat la limita de proprietate.", italic=True)

    _add_heading(doc, "4. Date tehnice și dimensionări", level=2)
    # Auto-calc IF length & debit available
    L = data.get("sf_lungime_conducta_m") or data.get("pt_lungime_m")
    Q = data.get("debit_instalat_mc_h")
    calc_done = False
    if L and Q:
        try:
            dim = engine.dimensionare_conducta(regime="joasa", length_m=float(L), debit_mc_h=float(Q),
                                               material=_get(data, "sf_material_conducta", "PE 100 SDR 11"))
            rec = dim.get("recommended", {})
            calc_rows = [
                ("Lungime conductă (m)", str(L)),
                ("Material", _get(data, "sf_material_conducta", "PE 100 SDR 11")),
                ("Debit calculat (m³/h)", str(Q)),
                ("DN recomandat (calc Renouard)", f"DN {rec.get('DN','—')} (Ø ext. {rec.get('od_mm','—')} mm)"),
                ("Pierdere presiune calculată", f"{rec.get('delta_p','—')} mbar"),
                ("Viteza gazului", f"{rec.get('viteza_m_s','—')} m/s"),
                ("Verdict dimensionare", "OK conform NTPEE 2018" if rec.get("ok") else "ATENȚIE - reverificare necesară"),
            ]
            _add_kv_table(doc, calc_rows)
            calc_done = True
        except (ValueError, TypeError):
            pass
    if not calc_done:
        _add_kv_table(doc, [
            ("Lungime conductă (m)", _get(data, "sf_lungime_conducta_m") or _get(data, "pt_lungime_m")),
            ("Material", _get(data, "sf_material_conducta", "PE 100 SDR 11")),
            ("Diametru nominal", _get(data, "sf_diametru_nominal_DN")),
            ("Presiune maximă operare (bar)", _get(data, "sf_presiune_max_op_bar")),
            ("Pierderi presiune calculate (bar)", _get(data, "pt_calcul_pierderi_presiune_bar")),
        ])

    # 4.1 Condiții naturale ale zonei (NTPEE 2018 art. 14)
    if any(data.get(k) for k in ["adancime_inghet_cm", "altitudine_m", "temp_medie_anuala_C",
                                  "temp_minima_iarna_C", "relief_zona"]):
        _add_heading(doc, "4.1 Condiții naturale ale amplasamentului", level=3)
        _add_kv_table(doc, [
            ("Adâncime de îngheț (cm)", _get(data, "adancime_inghet_cm", "80")),
            ("Altitudine teren (m)", _get(data, "altitudine_m")),
            ("Temperatura medie anuală (°C)", _get(data, "temp_medie_anuala_C")),
            ("Temperatura minimă de iarnă (°C)", _get(data, "temp_minima_iarna_C", "-15")),
            ("Relief / tip zonă", _get(data, "relief_zona", "Câmpie urbană")),
        ])

    # 4.2 Date seismice (P100-1/2013)
    if any(data.get(k) for k in ["seismic_acceleratie_ag", "seismic_grad_SR11100", "seismic_perioada_colt_Tc"]):
        _add_heading(doc, "4.2 Date seismice ale zonei (P100-1/2013)", level=3)
        _add_kv_table(doc, [
            ("Accelerația terenului ag", _get(data, "seismic_acceleratie_ag", "0.20g")),
            ("Grad seismic SR 11100", _get(data, "seismic_grad_SR11100", "VII")),
            ("Perioada de colț Tc (s)", _get(data, "seismic_perioada_colt_Tc", "1.0")),
        ])

    # 4.3 Categoria de importanță (HG 766/1997) + destinația clădirii
    if any(data.get(k) for k in ["categorie_importanta_HG766", "cladire_destinatie"]):
        _add_heading(doc, "4.3 Categoria de importanță și destinația", level=3)
        _add_kv_table(doc, [
            ("Categoria de importanță (HG 766/1997)", _get(data, "categorie_importanta_HG766", "Categoria C — Normală")),
            ("Destinația clădirii", _get(data, "cladire_destinatie", "Locuință familială")),
        ])

    # 4.4 Centrală termică (dacă există)
    if _truthy(data, "are_centrala_termica") or data.get("centrala_producator") or data.get("centrala_putere_kw"):
        _add_heading(doc, "4.4 Aparat consumator — centrală termică", level=3)
        _add_kv_table(doc, [
            ("Producător centrală", _get(data, "centrala_producator")),
            ("Putere nominală (kW)", _get(data, "centrala_putere_kw")),
            ("Tip", _get(data, "tip_consumator", "Centrală termică murală cu tiraj forțat")),
        ])

    _add_heading(doc, "5. Probe și verificări (NTPEE 2018 cap. 5)", level=2)
    _add_para(doc, "Înainte de PIF, conducta va fi supusă următoarelor probe:")
    p_lucru = data.get("sf_presiune_max_op_bar")
    p_rez_min, p_et_min = "6", "0.11"
    if p_lucru:
        try:
            v = engine.validare_probe(0, 0, float(p_lucru))
            p_rez_min = str(v["rezistenta_min_bar"])
            p_et_min = str(v["etanseitate_min_bar"])
        except (ValueError, TypeError):
            pass
    _add_kv_table(doc, [
        ("Proba de rezistență minimă (bar)", p_rez_min),
        ("Proba de etanșeitate minimă (bar)", p_et_min),
        ("Durată proba rezistență", "min. 1 oră (PE)"),
        ("Durată proba etanșeitate", "min. 24 ore (rețele)"),
    ])

    # Bloc condițional pe rezultate probe efective
    if _truthy(data, "proba_rezultat"):
        _add_para(doc, f"Rezultat probe efectuate: {_get(data,'proba_rezultat')}", bold=True)
        if _truthy(data, "proba_observatii"):
            _add_para(doc, f"Observații: {_get(data,'proba_observatii')}", italic=True)

    # 6. Exigențele esențiale (Legea 10/1995 art. 5)
    _add_heading(doc, "6. Exigențele esențiale ale construcției (Legea 10/1995)", level=2)
    EXIG_DEFAULTS = {
        "exig_A_rezistenta": "Conducta PE 100 SDR 11 + sudură electrofuziune asigură rezistența mecanică conform NTPEE 2018 art. 60. Adâncimea de pozare ≥ 0.8m sub nivelul terenului asigură protecție la sarcini.",
        "exig_B_siguranta_expl": "Probe rezistență și etanșeitate conform NTPEE 2018. Robinet de branșament cu acces public, marcaje vizuale și banda de avertizare la 30cm deasupra conductei.",
        "exig_C_siguranta_foc": "Distanțe minime de siguranță respectate: 3m față de fundații, 0.5m față de cabluri electrice, 1m față de canalizare. Materialul PE 100 conform EN 1555.",
        "exig_D_mediu": "Lucrările respectă normele de protecție a mediului. Refacerea peisajului după execuție. Materialele excavate sunt evacuate la rampă autorizată.",
    }
    _add_kv_table(doc, [
        ("A — Rezistență mecanică și stabilitate", _get(data, "exig_A_rezistenta", EXIG_DEFAULTS["exig_A_rezistenta"])),
        ("B — Siguranță în exploatare", _get(data, "exig_B_siguranta_expl", EXIG_DEFAULTS["exig_B_siguranta_expl"])),
        ("C — Siguranță la foc", _get(data, "exig_C_siguranta_foc", EXIG_DEFAULTS["exig_C_siguranta_foc"])),
        ("D — Igienă, sănătate, mediu", _get(data, "exig_D_mediu", EXIG_DEFAULTS["exig_D_mediu"])),
    ])

    _add_heading(doc, "7. Verificare independentă (VGD)", level=2)
    if _truthy(data, "dtac_verificator_vgd"):
        _add_para(doc, f"Verificator de proiect (VGD) atestat: {_get(data,'dtac_verificator_vgd')}", bold=True)
    else:
        _add_para(doc, "Verificator de proiect (VGD) atestat ANRE — se va completa la finalizare DTAC.",
                  italic=True)

    _footer_signature(doc, proj, data)
    return _save(doc)


# ============================================================================
# TEMPLATE 4: Caiet de Sarcini Execuție
# ============================================================================
def caiet_sarcini(proj: Dict[str, Any]) -> bytes:
    data = proj.get("data") or {}
    doc = Document()
    _company_header(doc)
    _project_cartouche(doc, proj, data)
    _add_heading(doc, "CAIET DE SARCINI — EXECUȚIE INSTALAȚIE GAZE NATURALE", level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_para(doc, "(Conform NTPEE 2018 cap. 4 + Legea 10/1995 — Calitate construcții)",
              italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
    doc.add_paragraph()

    _add_heading(doc, "1. Obiectul lucrării", level=2)
    _add_para(doc, _get(data, "scop_lucrare", "Realizare branșament gaze naturale + instalație utilizare"))

    _add_heading(doc, "2. Cerințe față de executant", level=2)
    doc.add_paragraph("• Autorizație ANRE EDD (Execuție Distribuție gaze) valabilă", style="List Bullet")
    doc.add_paragraph("• Responsabil Tehnic cu Execuția (RTE) atestat MDLPA", style="List Bullet")
    doc.add_paragraph("• Sudori autorizați (PE sau OL după caz)", style="List Bullet")
    doc.add_paragraph("• Diriginte de șantier atestat", style="List Bullet")

    _add_heading(doc, "3. Materiale aprobate", level=2)
    _add_para(doc, _get(data, "pt_lista_materiale", "Conductă PE 100 SDR 11 conform EN 1555, "
              "armături sub presiune, regulator presiune, contor cu corecție automată conform Ord. ANRE 75/2020."))

    _add_heading(doc, "4. Tehnologii de montaj", level=2)
    doc.add_paragraph("• Sudură PE prin electrofuziune sau cap-la-cap conform standardelor", style="List Bullet")
    doc.add_paragraph("• Adâncime pozare conform NTPEE 2018 art. 56:", style="List Bullet")
    doc.add_paragraph("    - 0,9 m minim sub trotuar", style="List Bullet 2")
    doc.add_paragraph("    - 1,0 m minim sub trafic auto", style="List Bullet 2")
    doc.add_paragraph("    - 0,6 m minim în spații verzi", style="List Bullet 2")
    doc.add_paragraph("• Pat de nisip 10 cm sub și deasupra conductei", style="List Bullet")
    doc.add_paragraph("• Bandă de avertizare galbenă la 30 cm deasupra conductei", style="List Bullet")

    # 4.1 Detalii tehnologice specifice (consumă 12+ placeholders tehnice)
    _add_heading(doc, "4.1 Detalii tehnologice specifice proiectului", level=3)
    tech_rows = [
        ("Tip sudură electrofuziune", _get(data, "tip_sudura", "Electrofuziune cu manșoane PE 100 (preferat) sau cap-la-cap")),
        ("Unghi minim cuplare (grade)", _get(data, "unghi_cuplare_min_grade", "90 (perpendicular pe traseu)")),
        ("Fir trasor — material", _get(data, "fir_trasor_material", "Cupru izolat PVC")),
        ("Fir trasor — secțiune (mm²)", _get(data, "fir_trasor_sectiune_mm2", "2.5")),
        ("Tub de protecție", _get(data, "tub_protectie", "PVC rigid Ø110 (la traversări drumuri, sub fundații, în pereți)")),
        ("Pat cărămizi de protecție", _get(data, "pat_caramizi", "Da — la traversare cu cabluri electrice și telefonice (Ord. ANRE)")),
        ("Distanță minimă față de limita de proprietate (m)", _get(data, "pozare_distanta_limita", "0.5")),
    ]
    _add_kv_table(doc, tech_rows)

    # 4.2 Traseu și lungimi (consumă lungime_*, traseu_pe_drum)
    if any(data.get(k) for k in ["lungime_pe_drum_m", "lungime_raiser_m", "traseu_pe_drum"]):
        _add_heading(doc, "4.2 Repartiție traseu conductă", level=3)
        traseu_rows = [
            ("Lungime conductă în carosabil/drum (m)", _get(data, "lungime_pe_drum_m")),
            ("Lungime coloană ascendentă raiser (m)", _get(data, "lungime_raiser_m")),
            ("Traseu pozare predominant", _get(data, "traseu_pe_drum", "Trotuar / spațiu verde / carosabil — vezi planul de situație D2")),
        ]
        _add_kv_table(doc, traseu_rows)

    # 4.3 Conductă existentă (la branșament în rețea existentă)
    if data.get("conducta_existenta_strada") or data.get("conducta_existenta_caracteristici"):
        _add_heading(doc, "4.3 Conductă existentă (punct de racord)", level=3)
        _add_kv_table(doc, [
            ("Stradă conductă existentă", _get(data, "conducta_existenta_strada")),
            ("Caracteristici (DN / material / presiune)", _get(data, "conducta_existenta_caracteristici")),
        ])

    _add_heading(doc, "5. Controale și verificări", level=2)
    doc.add_paragraph("• Verificare materiale cu certificate de calitate", style="List Bullet")
    doc.add_paragraph("• Verificare sudori și buletine de autorizație", style="List Bullet")
    doc.add_paragraph("• Probă de rezistență (apă sau aer) — min. 1,5 × P_op, min. 4 bar pentru PE", style="List Bullet")
    doc.add_paragraph("• Probă de etanșeitate — 1,1 × P_op, min. 24 ore", style="List Bullet")
    doc.add_paragraph("• Probă PIF (împreună cu OSD)", style="List Bullet")

    if _truthy(data, "exec_firma"):
        _add_heading(doc, "6. Executant desemnat", level=2)
        _add_kv_table(doc, [
            ("Firmă executantă", _get(data, "exec_firma")),
            ("RTE", _get(data, "exec_responsabil_tehnic")),
            ("Diriginte șantier", _get(data, "exec_diriginte_santier")),
            ("Data start lucrări", _get(data, "exec_data_start")),
            ("Data terminare", _get(data, "exec_data_terminare")),
        ])

    _footer_signature(doc, proj, data)
    return _save(doc)


# ============================================================================
# TEMPLATE 5: Borderou Piese
# ============================================================================
def borderou(proj: Dict[str, Any]) -> bytes:
    data = proj.get("data") or {}
    doc = Document()
    _company_header(doc)
    _project_cartouche(doc, proj, data)
    _add_heading(doc, "BORDEROU PIESE SCRISE ȘI DESENATE", level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_para(doc, f"Proiect: {proj.get('title','—')}", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    _add_heading(doc, "A. Piese scrise", level=2)
    pieces = [
        ("1", "Foaie de capăt"),
        ("2", "Borderou"),
        ("3", "Tema de proiectare"),
        ("4", "Memoriu tehnic justificativ"),
        ("5", "Calcule de dimensionare"),
        ("6", "Caiet de sarcini"),
        ("7", "Liste de cantități de lucrări (antemăsurătoare)"),
        ("8", "Programe de verificare (PCC)"),
    ]
    if _truthy(data, "dtac_verificator_vgd"):
        pieces.append(("9", f"Raport verificare VGD: {_get(data,'dtac_verificator_vgd')}"))
    t = doc.add_table(rows=1, cols=3)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    hdr[0].text = "Nr."
    hdr[1].text = "Denumire piesă"
    hdr[2].text = "Format"
    for n, lbl in pieces:
        r = t.add_row().cells
        r[0].text = n
        r[1].text = lbl
        r[2].text = "A4"

    doc.add_paragraph()
    _add_heading(doc, "B. Piese desenate", level=2)
    drawings = [
        ("D1", "Plan încadrare în zonă", "1:5000", "A3"),
        ("D2", "Plan de situație", "1:500", "A3"),
        ("D3", "Profil longitudinal conductă", "1:500/1:100", "A2"),
        ("D4", "Detalii branșament", "1:20", "A3"),
        ("D5", "Schemă izometrică instalație utilizare", "—", "A3"),
        ("D6", "Detalii post reglare-măsurare", "1:10", "A3"),
    ]
    t = doc.add_table(rows=1, cols=4)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    hdr[0].text = "Cod"
    hdr[1].text = "Denumire"
    hdr[2].text = "Scară"
    hdr[3].text = "Format"
    for c, lbl, sc, fmt in drawings:
        r = t.add_row().cells
        r[0].text = c
        r[1].text = lbl
        r[2].text = sc
        r[3].text = fmt

    if _truthy(data, "pt_numar_planse"):
        doc.add_paragraph()
        _add_para(doc, f"Număr total planșe PT: {_get(data,'pt_numar_planse')}", bold=True)

    # C. Materiale principale și furnizori
    doc.add_paragraph()
    _add_heading(doc, "C. Materiale principale și furnizori", level=2)
    mat_rows = [
        ("Conductă PE 100 SDR 11", _get(data, "sf_diametru_nominal_DN", "DN 32"),
         _get(data, "mat_furnizor_teava", "—"),
         "EN 1555"),
        ("Robinet branșament sub presiune", _get(data, "mat_serie_robinet_br", "—"),
         _get(data, "mat_furnizor_robinet", "—"),
         "EN 1555-4"),
        ("Regulator presiune", "DN 25/JP",
         _get(data, "mat_marca_regulator", "—"),
         "EN 88-1"),
        ("Contor gaze", _get(data, "mat_marca_contor", "—"),
         "—", "EN 1359 / Ord. ANRE 75/2020"),
    ]
    tm = doc.add_table(rows=1, cols=4)
    tm.style = "Light Grid Accent 1"
    hdr = tm.rows[0].cells
    hdr[0].text = "Element"; hdr[1].text = "Tip / DN / Serie"; hdr[2].text = "Furnizor"; hdr[3].text = "Standard"
    for el, tp, furn, std in mat_rows:
        r = tm.add_row().cells
        r[0].text = el; r[1].text = tp; r[2].text = furn; r[3].text = std
    if _truthy(data, "mat_certif_conformitate"):
        _add_para(doc, f"Certificate de conformitate: {_get(data,'mat_certif_conformitate')}", italic=True, size=9)

    # D. Avize utilități obținute
    AVIZE = [
        ("Aviz E-Distribuție", "aviz_edistr_nr_data"),
        ("Aviz Telekom", "aviz_telekom_nr_data"),
        ("Aviz Apa Nova", "aviz_apa_nr_data"),
        ("Aviz STB transport", "aviz_stb_nr_data"),
        ("Aviz NetCity / fibră", "aviz_netcity_nr_data"),
        ("Aviz Luxten iluminat", "aviz_luxten_nr_data"),
        ("Aviz străzi PMB", "aviz_strazi_nr_data"),
        ("Aviz circulație PMB", "aviz_circulatie_pmb_nr_data"),
        ("Aviz mediu PMB", "aviz_mediu_pmb_nr_data"),
        ("Aviz APM (Mediu)", "aviz_apm_nr_data"),
        ("Acord acces proprietate", "acord_acces_nr_data"),
    ]
    visible_avize = [(lbl, _get(data, k)) for lbl, k in AVIZE if data.get(k)]
    if visible_avize:
        doc.add_paragraph()
        _add_heading(doc, "D. Avize utilități obținute", level=2)
        ta = doc.add_table(rows=1, cols=2)
        ta.style = "Light Grid Accent 1"
        hdr = ta.rows[0].cells
        hdr[0].text = "Aviz"; hdr[1].text = "Nr. / Data"
        for lbl, v in visible_avize:
            r = ta.add_row().cells
            r[0].text = lbl; r[1].text = v

    _footer_signature(doc, proj, data)
    return _save(doc)
# ============================================================================
def cerere_pif(proj: Dict[str, Any]) -> bytes:
    data = proj.get("data") or {}
    osd = _get(data, "atr_osd", "Distrigaz Sud Rețele")
    doc = Document()
    _company_header(doc)
    _add_para(doc, f"Către {osd.upper()}", bold=True, size=12)
    _add_para(doc, "Serviciu Punere în Funcțiune", size=10, italic=True)
    doc.add_paragraph()

    _add_heading(doc, "CERERE PUNERE ÎN FUNCȚIUNE (PIF)", level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_para(doc, "(Conform Ord. ANRE 162/2021)", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
    doc.add_paragraph()

    _add_para(doc, "Subscrisa Energy Project Design S.R.L., în calitate de proiectant atestat, "
                   "vă solicită prin prezenta efectuarea Punerii în Funcțiune pentru lucrarea:", italic=True)

    doc.add_paragraph()
    _add_kv_table(doc, [
        ("Beneficiar", _get(data, "beneficiar_nume")),
        ("CNP/CUI", _get(data, "beneficiar_cnp_cui")),
        ("Adresă loc consum", _get(data, "loc_consum_adresa")),
        ("Cod loc consum (din ATR)", _get(data, "atr_numar")),
        ("Nr. AC", _get(data, "ac_numar")),
        ("Data emiterii AC", _get(data, "ac_data_emitere")),
        ("Firmă executantă", _get(data, "exec_firma")),
        ("Data terminării lucrărilor", _get(data, "exec_data_terminare")),
    ])

    doc.add_paragraph()
    _add_para(doc, "Anexăm prezenta cerere cu următoarele documente:", bold=True)
    docs_anex = [
        f"• Proces Verbal Recepție la Terminarea Lucrărilor nr. {_get(data,'receptie_pv_numar','—')} / {_get(data,'receptie_pv_data','—')}",
        f"• Buletine probe rezistență ({_get(data,'proba_rezistenta_bar','—')} bar / {_get(data,'proba_rezistenta_durata_min','—')} min)",
        f"• Buletine probe etanșeitate ({_get(data,'proba_etanseitate_bar','—')} bar / {_get(data,'proba_etanseitate_durata_h','—')} ore)",
        "• Certificate calitate materiale",
        "• Buletine autorizație sudori",
    ]
    if _truthy(data, "as_built_anexat") and data["as_built_anexat"] == "Da":
        docs_anex.append("• Documentație as-built")
    for d in docs_anex:
        doc.add_paragraph(d, style="List Bullet")

    if _truthy(data, "proba_rezultat"):
        doc.add_paragraph()
        verdict = _get(data, "proba_rezultat")
        _add_para(doc, f"Rezultat global probe: {verdict}", bold=True,
                  italic=(verdict == "Respins"))

    doc.add_paragraph()
    _add_para(doc, f"Data cererii: {_today_ro()}", italic=True)

    _footer_signature(doc, proj, data)
    return _save(doc)


# ============================================================================
# TEMPLATE 7: PV Recepție Terminarea Lucrărilor (PVRTL)
# ============================================================================
def pv_receptie(proj: Dict[str, Any]) -> bytes:
    data = proj.get("data") or {}
    doc = Document()
    _company_header(doc)
    _add_heading(doc, "PROCES VERBAL DE RECEPȚIE LA TERMINAREA LUCRĂRILOR", level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    nr = _get(data, "receptie_pv_numar", "—")
    dt = _get(data, "receptie_pv_data", _today_ro())
    _add_para(doc, f"Nr. {nr} / {dt}", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
    _add_para(doc, "(Conform HG 273/1994 + Ord. MLPAT 770/1997)",
              italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
    doc.add_paragraph()

    _add_heading(doc, "1. Identificare obiect recepționat", level=2)
    _add_kv_table(doc, [
        ("Denumire lucrare", proj.get("title", "—")),
        ("Beneficiar", _get(data, "beneficiar_nume")),
        ("Adresă", _get(data, "loc_consum_adresa")),
        ("Autorizație Construire", f"AC nr. {_get(data,'ac_numar','—')} / {_get(data,'ac_data_emitere','—')}"),
        ("Proiectant", _get(data, "dtac_proiectant_specialitate", COMPANY["name"])),
        ("Executant", _get(data, "exec_firma")),
        ("RTE", _get(data, "exec_responsabil_tehnic")),
        ("Diriginte șantier", _get(data, "exec_diriginte_santier")),
        ("Data începere", _get(data, "exec_data_start")),
        ("Data terminare", _get(data, "exec_data_terminare")),
    ])

    _add_heading(doc, "2. Comisia de recepție", level=2)
    _add_para(doc, _get(data, "receptie_comisia",
              "Reprezentant beneficiar, Diriginte șantier, Reprezentant executant, "
              "Reprezentant OSD (la cerere), Verificator VGD"))

    _add_heading(doc, "3. Verificări efectuate", level=2)
    doc.add_paragraph("• Conformitate execuție cu proiectul tehnic", style="List Bullet")
    doc.add_paragraph("• Probe de rezistență și etanșeitate", style="List Bullet")
    doc.add_paragraph("• Certificate calitate materiale", style="List Bullet")
    doc.add_paragraph("• Buletine sudori", style="List Bullet")

    _add_heading(doc, "4. Concluzia comisiei", level=2)
    rez = _get(data, "proba_rezultat", "Admis")
    if rez == "Admis":
        _add_para(doc, "✓ Lucrarea se RECEPȚIONEAZĂ și se predă beneficiarului în condiții de funcționare conform proiectului.",
                  bold=True)
    elif rez == "Admis cu observații":
        _add_para(doc, "⚠ Lucrarea se RECEPȚIONEAZĂ CU OBSERVAȚII. Beneficiarul va remedia neconformitățile în 30 de zile.",
                  bold=True)
    else:
        _add_para(doc, "✗ Lucrarea NU se RECEPȚIONEAZĂ. Se va întocmi listă de remedieri obligatorii.",
                  bold=True)

    if _truthy(data, "proba_observatii"):
        doc.add_paragraph()
        _add_para(doc, "Observații/Remedieri necesare:", bold=True)
        _add_para(doc, _get(data, "proba_observatii"))

    _footer_signature(doc, proj, data)
    return _save(doc)


# ============================================================================
# TEMPLATE 8: Cartea Tehnică a Construcției
# ============================================================================
def carte_tehnica(proj: Dict[str, Any]) -> bytes:
    data = proj.get("data") or {}
    doc = Document()
    _company_header(doc)
    _project_cartouche(doc, proj, data)
    _add_heading(doc, "CARTEA TEHNICĂ A CONSTRUCȚIEI", level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_para(doc, f"Obiect: {proj.get('title','—')}", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_para(doc, "(Conform HG 273/1994 + Ord. MLPAT 770/1997 — 4 secțiuni obligatorii)",
              italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
    doc.add_paragraph()

    # Date identificatoare obiectiv
    _add_heading(doc, "Date identificatoare", level=2)
    _add_kv_table(doc, [
        ("Beneficiar", _get(data, "beneficiar_nume")),
        ("Loc consum", f"{_get(data,'loc_consum_adresa')}, {_get(data,'loc_consum_localitate')}, jud. {_get(data,'loc_consum_judet')}"),
        ("Cadastru / CF", _get(data, "loc_consum_cadastru")),
        ("Tip lucrare", _get(data, "tipul_lucrarii")),
        ("Categorie importanță (HG 766/1997)", _get(data, "categorie_importanta_HG766", "C")),
        ("Operator distribuție", _get(data, "atr_osd")),
        ("Nr. proiect / an", _get(data, "proiect_nr_an")),
        ("Exemplare", _get(data, "exemplare_nr", "4 (2 beneficiar + 1 OSD + 1 arhivă)")),
    ])

    _add_heading(doc, "Secțiunea A — Documentația de proiectare", level=2)
    _add_kv_table(doc, [
        ("Proiectant", _get(data, "dtac_proiectant_specialitate", COMPANY["name"])),
        ("Atestat ANRE", _get(data, "dtac_atestat_proiectant", COMPANY["anre_proiectant"])),
        ("Data întocmire DTAC", _get(data, "dtac_data_intocmire")),
        ("Verificator VGD", _get(data, "dtac_verificator_vgd")),
        ("CU nr.", f"{_get(data,'cu_numar')} / {_get(data,'cu_data_emitere')}"),
        ("ATR nr.", f"{_get(data,'atr_numar')} / {_get(data,'atr_data')}"),
        ("AC nr.", f"{_get(data,'ac_numar')} / {_get(data,'ac_data_emitere')}"),
    ])
    # Conținut narativ pentru Secțiunea A
    ct_a_default = ("Documentația conține: memoriu tehnic justificativ, planuri de situație 1:500 și 1:200, "
                    "calcule hidraulice (debit, viteză, pierdere presiune), schema izometrică, lista materialelor, "
                    "caietul de sarcini de execuție, referat de verificare independentă (VGD).")
    _add_para(doc, _get(data, "ct_sectiune_A_continut", ct_a_default), italic=True)

    _add_heading(doc, "Secțiunea B — Documentația de execuție", level=2)
    _add_kv_table(doc, [
        ("Executant", _get(data, "exec_firma")),
        ("RTE", _get(data, "exec_responsabil_tehnic")),
        ("Diriginte șantier", _get(data, "exec_diriginte_santier")),
        ("Data începere", _get(data, "exec_data_start")),
        ("Data terminare", _get(data, "exec_data_terminare")),
        ("Materiale principale", _get(data, "pt_lista_materiale")),
        ("Certificate conformitate materiale", _get(data, "mat_certif_conformitate")),
    ])
    ct_b_default = ("Documentația cuprinde: ordin de începere lucrări, PV predare-primire amplasament, "
                    "Program de control al calității, certificate calitate materiale, fișe sudori autorizați, "
                    "PV-uri faze determinante, PV lucrări ascunse, jurnale de șantier și buletine probe.")
    _add_para(doc, _get(data, "ct_sectiune_B_continut", ct_b_default), italic=True)

    _add_heading(doc, "Secțiunea C — Documentația de recepție", level=2)
    _add_kv_table(doc, [
        ("PVRTL nr.", f"{_get(data,'receptie_pv_numar')} / {_get(data,'receptie_pv_data')}"),
        ("Comisia recepție", _get(data, "receptie_comisia")),
        ("  - Președinte", _get(data, "com_recep_presedinte")),
        ("  - Reprezentant OSD", _get(data, "com_recep_reprez_osd")),
        ("  - Reprezentant ISC", _get(data, "com_recep_reprez_isc")),
        ("  - Reprezentant beneficiar", _get(data, "com_recep_reprez_beneficiar")),
        ("Probe rezistență (bar)", _get(data, "proba_rezistenta_bar")),
        ("Probe etanșeitate (bar)", _get(data, "proba_etanseitate_bar")),
        ("Rezultat probe", _get(data, "proba_rezultat")),
        ("As-built anexat", _get(data, "as_built_anexat")),
    ])
    ct_c_default = ("Documentația cuprinde: PV recepție la terminarea lucrărilor (PVRTL), PV recepție finală, "
                    "buletine probe etanșeitate și rezistență, declarații de conformitate, planuri as-built, "
                    "PV PIF emis de OSD, contract de furnizare gaze naturale.")
    _add_para(doc, _get(data, "ct_sectiune_C_continut", ct_c_default), italic=True)

    _add_heading(doc, "Secțiunea D — Urmărirea în timp a comportării construcției", level=2)
    _add_para(doc, "Verificări periodice (NTPEE 2018 art. 78):")
    doc.add_paragraph("• Verificare etanșeitate la 2 ani — în sarcina beneficiarului prin firmă autorizată", style="List Bullet")
    doc.add_paragraph("• Revizie tehnică instalație utilizare la 10 ani — Ord. ANRE 16/2015", style="List Bullet")
    doc.add_paragraph("• Verificare regulator / contor — la fiecare reverificare metrologică", style="List Bullet")
    doc.add_paragraph("• Defecțiuni majore — notificare OSD în max 24 ore", style="List Bullet")
    ct_d_default = ("Beneficiarul / administratorul construcției are obligația de a păstra Cartea Tehnică pe toată "
                    "durata existenței obiectivului. Orice intervenție ulterioară (extindere, mutare contor, "
                    "modificare traseu, scoatere din funcțiune) se consemnează în Secțiunea D prin proces-verbal.")
    _add_para(doc, _get(data, "ct_sectiune_D_continut", ct_d_default), italic=True)

    if _truthy(data, "pif_data"):
        _add_heading(doc, "PIF (Punere în Funcțiune)", level=2)
        _add_kv_table(doc, [
            ("Data PIF", _get(data, "pif_data")),
            ("OSD prezent", _get(data, "pif_osd")),
            ("Responsabil OSD", _get(data, "pif_responsabil_osd")),
            ("Serie contor", _get(data, "pif_contor_serie")),
            ("Index inițial", _get(data, "pif_contor_index_initial")),
            ("Contract furnizare", _get(data, "pif_contract_furnizare")),
        ])

    _footer_signature(doc, proj, data)
    return _save(doc)


# ============================================================================
# TEMPLATE 9-14: Cereri Avize Utilități + Anunțuri (5 noi)
# ============================================================================
def _cerere_aviz_generic(proj: Dict[str, Any], destinatar: str, denumire_aviz: str,
                          cadru_legal: str, anexe: List[str],
                          extra_block: Optional[str] = None) -> bytes:
    """Builder generic pentru cereri de aviz utilități."""
    data = proj.get("data") or {}
    doc = Document()
    _company_header(doc)
    _add_para(doc, f"Către {destinatar.upper()}", bold=True, size=12)
    doc.add_paragraph()

    _add_heading(doc, f"CERERE PENTRU EMITEREA {denumire_aviz.upper()}", level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_para(doc, f"({cadru_legal})", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
    doc.add_paragraph()

    _add_para(doc, "Solicitant (proiectant gaze naturale):", bold=True)
    _add_kv_table(doc, [
        ("Denumire", COMPANY["name"]),
        ("CUI / Reg. Comerțului", f"{COMPANY['cui']} / {COMPANY['reg_com']}"),
        ("Atestat ANRE", COMPANY["anre_proiectant"]),
        ("Email", COMPANY["email"]),
        ("Telefon", COMPANY["phone"]),
    ])

    doc.add_paragraph()
    _add_para(doc, "În numele beneficiarului:", bold=True)
    _add_kv_table(doc, [
        ("Beneficiar", _get(data, "beneficiar_nume")),
        ("CNP/CUI", _get(data, "beneficiar_cnp_cui")),
        ("Telefon", _get(data, "beneficiar_telefon")),
        ("Email", _get(data, "beneficiar_email")),
    ])

    doc.add_paragraph()
    _add_para(doc, "Pentru obiectivul:", bold=True)
    _add_kv_table(doc, [
        ("Tip lucrare", _get(data, "scop_lucrare", "Branșament gaze naturale")),
        ("Adresă", _get(data, "loc_consum_adresa")),
        ("Localitate", _get(data, "loc_consum_localitate")),
        ("Județ", _get(data, "loc_consum_judet")),
        ("Nr. cadastral / CF", _get(data, "loc_consum_cadastru")),
    ])

    if extra_block:
        doc.add_paragraph()
        _add_para(doc, extra_block, italic=True)

    if _truthy(data, "cu_numar"):
        doc.add_paragraph()
        _add_para(doc, f"În baza Certificatului de Urbanism nr. {_get(data,'cu_numar')} din {_get(data,'cu_data_emitere','—')}", bold=True)

    doc.add_paragraph()
    _add_para(doc, "Anexăm prezenta cerere cu următoarele documente:", bold=True)
    for a in anexe:
        doc.add_paragraph(a, style="List Bullet")

    doc.add_paragraph()
    _add_para(doc, f"Data: {_today_ro()}", italic=True)
    _footer_signature(doc, proj, data)
    return _save(doc)


def cerere_aviz_apa(proj: Dict[str, Any]) -> bytes:
    data = proj.get("data") or {}
    destinatar = _get(data, "apa_canal_operator", "RAJA / Apa Nova / SC Compania Județeană Apă")
    return _cerere_aviz_generic(proj, destinatar, "Avizului tehnic de amplasament — APĂ-CANALIZARE",
        "Conform Legii 241/2006 — Serviciul de alimentare cu apă și de canalizare",
        ["Plan situație 1:500 cu poziția conductei propuse",
         "Plan încadrare 1:5000",
         "Memoriu tehnic gaze naturale (sinteză)",
         "Copie Certificat Urbanism"],
        extra_block="Solicităm avizarea traseului conductei de gaze pentru evitarea intersecțiilor cu rețelele existente de apă-canal. "
                    "Distanța minimă conform STAS 8591/1997: 0.5 m de la canalizare, 0.3 m de la apă potabilă.")


def cerere_aviz_electrica(proj: Dict[str, Any]) -> bytes:
    data = proj.get("data") or {}
    destinatar = _get(data, "electrica_operator", "E-Distribuție / Electrica / Delgaz Grid")
    return _cerere_aviz_generic(proj, destinatar, "Avizului tehnic de amplasament — REȚELE ELECTRICE",
        "Conform Legii 123/2012 + Ord. ANRE 11/2014",
        ["Plan situație 1:500 cu traseul conductei",
         "Plan încadrare 1:5000",
         "Memoriu tehnic (sinteză)",
         "Copie CU"],
        extra_block="Pentru evitarea intersecțiilor cu LEA / LES și menținerea distanțelor de siguranță conform PE 101/85: "
                    "min. 1 m LES JT, 3 m LES MT, 5 m LEA MT, 10 m LEA IT.")


def cerere_aviz_drumuri(proj: Dict[str, Any]) -> bytes:
    data = proj.get("data") or {}
    destinatar = _get(data, "drumuri_administrator", "Direcția de Drumuri și Poduri / Compania Națională de Administrare Infrastructură Rutieră")
    return _cerere_aviz_generic(proj, destinatar, "Avizului de spargere / pozare în carosabil",
        "Conform OG 43/1997 — Regimul drumurilor + HG 540/2000",
        ["Plan situație 1:500 cu traseul de spargere",
         "Memoriu tehnic — descriere lucrare + tehnologie refacere",
         "Schiță detaliu refacere îmbrăcăminte",
         "Garanție bună execuție (dacă e cazul)",
         "Copie CU + AC (la trimitere)"],
        extra_block="Solicităm avizul pentru spargerea carosabilului în vederea pozării unei conducte PE 100 SDR 11. "
                    "Refacerea îmbrăcămintei se va executa conform AND 605:2016 și SR EN 13108-1.")


def cerere_aviz_politie(proj: Dict[str, Any]) -> bytes:
    data = proj.get("data") or {}
    destinatar = _get(data, "politie_unitatea", "Inspectoratul de Poliție al Județului — Serviciul Rutier")
    return _cerere_aviz_generic(proj, destinatar, "Avizului de la Poliția Rutieră — semnalizare lucrare",
        "Conform OUG 195/2002 — Codul Rutier + HG 1391/2006",
        ["Plan situație 1:500 cu zona lucrării",
         "Schemă de semnalizare temporară (conform SR 1848-7)",
         "Program de execuție + ore de lucru propuse",
         "Copie AC"],
        extra_block="Solicităm avizarea schemei de semnalizare temporară pentru perioada de execuție a lucrărilor de gaze "
                    "în zona carosabilului public, conform STAS SR 1848-7 — Semnalizare rutieră.")


def cerere_aviz_mediu(proj: Dict[str, Any]) -> bytes:
    data = proj.get("data") or {}
    destinatar = _get(data, "apm_unitate", "Agenția pentru Protecția Mediului — Județul " + _get(data, "loc_consum_judet", "—"))
    return _cerere_aviz_generic(proj, destinatar, "Punctului de vedere / Acordului de mediu",
        "Conform Legii 292/2018 + OUG 195/2005 — Protecția mediului",
        ["Memoriu de prezentare (conform anexa Ord. MMP 19/2010)",
         "Plan situație 1:500",
         "Plan încadrare 1:5000",
         "Copie CU",
         "Dovada achitării taxei"],
        extra_block="Solicităm punctul de vedere al APM pentru lucrările de instalație gaze naturale propuse, "
                    "lucrările nefiind supuse procedurii EIM conform anexa nr. 2 din Legea 292/2018 (categoria 10/c — "
                    "rețele utilități publice cu lungime sub 1 km).")


def cerere_aviz_iscir(proj: Dict[str, Any]) -> bytes:
    data = proj.get("data") or {}
    destinatar = _get(data, "iscir_unitate", "ISCIR — Inspecția de Stat pentru Controlul Cazanelor, Recipientelor sub Presiune și Instalațiilor de Ridicat")
    return _cerere_aviz_generic(proj, destinatar, "Avizului ISCIR pentru centrala termică",
        "Conform Legii 64/2008 + PT C9-2010 (centrale termice)",
        ["Schemă termo-hidraulică",
         "Specificație tehnică centrală + arzător",
         "Schiță amplasare în cameră",
         "Calcul ventilație și evacuare gaze arse",
         "Copie CU + Memoriu tehnic"],
        extra_block="Solicităm avizul ISCIR pentru montarea/punerea în funcțiune a centralei termice "
                    "racordată la rețeaua de gaze naturale, conform PT C9 — Cerințe tehnice pentru centrale termice.")


def anunt_incepere(proj: Dict[str, Any]) -> bytes:
    """Anunț începere lucrări — către ISC + Primărie."""
    data = proj.get("data") or {}
    doc = Document()
    _company_header(doc)
    _add_para(doc, "Către:", bold=True, size=11)
    _add_para(doc, "1. INSPECTORATUL DE STAT ÎN CONSTRUCȚII — Inspectoratul Județean", bold=True, size=10)
    _add_para(doc, f"2. PRIMĂRIA {_get(data,'loc_consum_localitate','—').upper()}", bold=True, size=10)
    doc.add_paragraph()
    _add_heading(doc, "ANUNȚ ÎNCEPERE LUCRĂRI", level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_para(doc, "(Conform Legii 50/1991 art. 7 alin. (8))", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
    doc.add_paragraph()

    _add_para(doc, "Vă comunicăm începerea lucrărilor de construire pentru obiectivul:", bold=True)
    _add_kv_table(doc, [
        ("Denumire lucrare", proj.get("title", "—")),
        ("Adresă obiectiv", _get(data, "loc_consum_adresa")),
        ("Localitate", _get(data, "loc_consum_localitate")),
        ("Judet", _get(data, "loc_consum_judet")),
        ("Beneficiar", _get(data, "beneficiar_nume")),
        ("Autorizație Construire", f"AC nr. {_get(data,'ac_numar')} / {_get(data,'ac_data_emitere')}"),
        ("Emitent AC", _get(data, "ac_emitent")),
        ("Data programată începere", _get(data, "exec_data_start", _today_ro())),
        ("Durată estimată execuție", _get(data, "ac_termen_executie") + " luni" if _truthy(data, "ac_termen_executie") else "—"),
        ("Proiectant", _get(data, "dtac_proiectant_specialitate", COMPANY["name"])),
        ("Executant", _get(data, "exec_firma", COMPANY["name"])),
        ("RTE atestat", _get(data, "exec_responsabil_tehnic")),
        ("Diriginte șantier", _get(data, "exec_diriginte_santier")),
    ])

    doc.add_paragraph()
    _add_para(doc, "Anexăm:", bold=True)
    doc.add_paragraph("• Copie Autorizație de Construire", style="List Bullet")
    doc.add_paragraph("• Copie contract execuție", style="List Bullet")
    doc.add_paragraph("• Dovada atestare RTE și Diriginte", style="List Bullet")
    doc.add_paragraph("• Copie ATR (Aviz Tehnic Racordare)", style="List Bullet")

    _footer_signature(doc, proj, data)
    return _save(doc)


def predare_amplasament(proj: Dict[str, Any]) -> bytes:
    """Proces verbal de predare-primire amplasament."""
    data = proj.get("data") or {}
    doc = Document()
    _company_header(doc)
    _add_heading(doc, "PROCES VERBAL DE PREDARE-PRIMIRE AMPLASAMENT", level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_para(doc, f"Nr. ___ / {_today_ro()}", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
    _add_para(doc, "(Conform Legii 10/1995 art. 13 + Reg. recepție HG 273/1994)",
              italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
    doc.add_paragraph()

    _add_para(doc, "Subsemnatii:", bold=True)
    doc.add_paragraph(f"• Reprezentant beneficiar: {_get(data, 'beneficiar_nume')}, CNP/CUI {_get(data,'beneficiar_cnp_cui')}", style="List Bullet")
    doc.add_paragraph(f"• Reprezentant executant: {_get(data, 'exec_firma', COMPANY['name'])}", style="List Bullet")
    doc.add_paragraph(f"• Diriginte de șantier: {_get(data, 'exec_diriginte_santier')}", style="List Bullet")
    doc.add_paragraph(f"• Reprezentant proiectant: {_get(data, 'dtac_proiectant_specialitate', COMPANY['name'])}", style="List Bullet")
    doc.add_paragraph("• Reprezentant ISC (la cerere): _____________", style="List Bullet")

    doc.add_paragraph()
    _add_para(doc, "Am procedat la predarea-primirea amplasamentului pentru:", bold=True)
    _add_kv_table(doc, [
        ("Adresă", _get(data, "loc_consum_adresa")),
        ("Localitate / Județ", f"{_get(data,'loc_consum_localitate')}, {_get(data,'loc_consum_judet')}"),
        ("Nr. cadastral", _get(data, "loc_consum_cadastru")),
        ("AC nr.", f"{_get(data,'ac_numar','—')} / {_get(data,'ac_data_emitere','—')}"),
    ])

    doc.add_paragraph()
    _add_para(doc, "Cu această ocazie s-au constatat:", bold=True)
    doc.add_paragraph("1. Reperele topografice corespund cu cele din planul de amplasament", style="List Number")
    doc.add_paragraph("2. Nu există rețele subterane nesemnalizate care să împiedice execuția", style="List Number")
    doc.add_paragraph("3. Beneficiarul a pus la dispoziție terenul liber de servituți și sarcini", style="List Number")
    doc.add_paragraph("4. S-au prezentat: AC original, planurile DTAC, programul de execuție", style="List Number")
    doc.add_paragraph("5. Executantul a verificat conformitatea cu DTAC și are toate avizele necesare", style="List Number")

    doc.add_paragraph()
    _add_para(doc, "Începând cu data prezentă, executantul preia răspunderea integrală pentru lucrările "
                   "de execuție conform Legii 10/1995 și a normelor tehnice în vigoare.", italic=True)

    _footer_signature(doc, proj, data)
    return _save(doc)


def dispozitie_santier(proj: Dict[str, Any]) -> bytes:
    """Dispoziție de șantier (modificare/clarificare detaliu pe parcursul execuției)."""
    data = proj.get("data") or {}
    doc = Document()
    _company_header(doc)
    _add_heading(doc, "DISPOZIȚIE DE ȘANTIER", level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_para(doc, f"Nr. ___ / {_today_ro()}", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
    _add_para(doc, "(Conform Legii 10/1995 + Ord. MLPAT 24/N/1997)",
              italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
    doc.add_paragraph()

    _add_kv_table(doc, [
        ("Obiectiv", proj.get("title", "—")),
        ("Beneficiar", _get(data, "beneficiar_nume")),
        ("AC", f"{_get(data,'ac_numar','—')} / {_get(data,'ac_data_emitere','—')}"),
        ("Proiectant", _get(data, "dtac_proiectant_specialitate", COMPANY["name"])),
        ("Executant", _get(data, "exec_firma")),
    ])

    doc.add_paragraph()
    _add_para(doc, "Obiect dispoziție:", bold=True)
    _add_para(doc, _get(data, "dispozitie_obiect",
              "Se dispune ___________________________________________________________________________ "
              "(detaliază: modificarea traseului, schimbarea diametrului, soluție alternativă fitting, "
              "consolidare locală etc.)"), italic=True)

    doc.add_paragraph()
    _add_para(doc, "Justificare tehnică:", bold=True)
    _add_para(doc, _get(data, "dispozitie_justificare",
              "Motivul tehnic care impune dispoziția: ___________________________________________"),
              italic=True)

    doc.add_paragraph()
    _add_para(doc, "Implicații:", bold=True)
    doc.add_paragraph("• Calitate: respectă normele NTPEE 2018 + cerințele esențiale Legea 10/1995", style="List Bullet")
    doc.add_paragraph(f"• Cost: {_get(data,'dispozitie_cost_impact','fără impact / impact minor')}", style="List Bullet")
    doc.add_paragraph(f"• Termen: {_get(data,'dispozitie_termen_impact','fără impact')}", style="List Bullet")

    doc.add_paragraph()
    _add_para(doc, "Aprobat de proiectant atestat ANRE, conform competențelor:", bold=True)

    _footer_signature(doc, proj, data)
    return _save(doc)


def _footer_signature_legacy_stub(*_a, **_kw):  # noqa - kept for backwards-compat if referenced
    pass


# ============================================================================
# REGISTRY
# ============================================================================
TEMPLATES: Dict[str, Dict[str, Any]] = {
    "cerere_cu":              {"label": "Cerere Certificat de Urbanism",                  "phase": "cu",       "fn": cerere_cu,            "norm": "Legea 50/1991"},
    "cerere_atr":             {"label": "Cerere Aviz Tehnic Racordare (OSD)",             "phase": "cu",       "fn": cerere_atr,           "norm": "Ord. ANRE 89/2018"},
    "cerere_aviz_apa":        {"label": "Cerere aviz amplasament — APĂ-CANAL",            "phase": "cu",       "fn": cerere_aviz_apa,      "norm": "Legea 241/2006 + STAS 8591"},
    "cerere_aviz_electrica":  {"label": "Cerere aviz amplasament — REȚELE ELECTRICE",     "phase": "cu",       "fn": cerere_aviz_electrica,"norm": "Ord. ANRE 11/2014 + PE 101/85"},
    "cerere_aviz_drumuri":    {"label": "Cerere aviz Drumuri (spargere/pozare carosabil)","phase": "cu",       "fn": cerere_aviz_drumuri,  "norm": "OG 43/1997 + AND 605:2016"},
    "cerere_aviz_politie":    {"label": "Cerere aviz Poliția Rutieră (semnalizare)",      "phase": "cu",       "fn": cerere_aviz_politie,  "norm": "OUG 195/2002 + SR 1848-7"},
    "cerere_aviz_mediu":      {"label": "Cerere punct de vedere / acord MEDIU (APM)",     "phase": "cu",       "fn": cerere_aviz_mediu,    "norm": "Legea 292/2018"},
    "cerere_aviz_iscir":      {"label": "Cerere aviz ISCIR (centrale termice)",           "phase": "cu",       "fn": cerere_aviz_iscir,    "norm": "Legea 64/2008 + PT C9"},
    "memoriu_tehnic":         {"label": "Memoriu Tehnic Justificativ (DTAC + PT)",        "phase": "dtac",     "fn": memoriu_tehnic,       "norm": "HG 907/2016 + NTPEE 2018"},
    "caiet_sarcini":          {"label": "Caiet de Sarcini Execuție",                       "phase": "pt",       "fn": caiet_sarcini,        "norm": "NTPEE 2018 cap. 4"},
    "borderou":               {"label": "Borderou Piese (scrise + desenate)",             "phase": "pt",       "fn": borderou,             "norm": "HG 907/2016"},
    "anunt_incepere":         {"label": "Anunț începere lucrări (ISC + Primărie)",        "phase": "executie", "fn": anunt_incepere,       "norm": "Legea 50/1991 art. 7"},
    "predare_amplasament":    {"label": "PV Predare-Primire Amplasament",                 "phase": "executie", "fn": predare_amplasament,  "norm": "Legea 10/1995 art. 13"},
    "dispozitie_santier":     {"label": "Dispoziție de șantier (opțional)",               "phase": "executie", "fn": dispozitie_santier,   "norm": "Ord. MLPAT 24/N/1997"},
    "cerere_pif":             {"label": "Cerere Punere în Funcțiune (OSD)",               "phase": "pif",      "fn": cerere_pif,           "norm": "Ord. ANRE 162/2021"},
    "pv_receptie":            {"label": "PV Recepție la Terminarea Lucrărilor",           "phase": "receptie", "fn": pv_receptie,          "norm": "HG 273/1994"},
    "carte_tehnica":          {"label": "Cartea Tehnică a Construcției (4 secțiuni)",     "phase": "receptie", "fn": carte_tehnica,        "norm": "HG 273/1994 + Ord. MLPAT 770/1997"},
}

# V10.5 — Register the MASTER document: Proiect Bransament Complet (replică XLS user)
try:
    from gas_doc_proiect_complet import proiect_bransament_complet as _proiect_complet_fn

    def _proiect_complet_wrap(proj):
        return _proiect_complet_fn(proj)

    TEMPLATES["proiect_bransament_complet"] = {
        "label": "Proiect Bransament COMPLET (Referat + Foaie capat + Borderou + Memoriu + ANEXA 14 materiale auto)",
        "phase": "dtac",
        "fn": _proiect_complet_wrap,
        "norm": "NTPEE 2018 + HG 907/2016 + ANEXA 13 (554 SAP materials)",
    }
except Exception as _e:  # pragma: no cover
    pass


def list_templates() -> List[Dict[str, Any]]:
    """Return list of {id, label, phase, norm} — UI consumes this list."""
    return [{"id": k, **{kk: vv for kk, vv in v.items() if kk != "fn"}} for k, v in TEMPLATES.items()]


def generate(template_id: str, proj: Dict[str, Any]) -> Optional[Tuple[bytes, str]]:
    """Generate a DOCX. Returns (bytes, filename) or None if template_id unknown."""
    tpl = TEMPLATES.get(template_id)
    if not tpl:
        return None
    data = tpl["fn"](proj)
    safe_title = (proj.get("title", "proiect")[:40]).replace("/", "-").replace(" ", "_")
    fname = f"{template_id}_{safe_title}.docx"
    return data, fname


# Auto-register extra V6.4 templates (PV LA, PV FD, PCC, RVT, ISC, as-built)
try:
    import gas_doc_templates_extra as _extra  # noqa: E402
    _extra.register_into(__import__(__name__))
except Exception:  # pragma: no cover - defensive
    pass
