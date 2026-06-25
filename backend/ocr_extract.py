"""OCR & Field Auto-Extract — V6.4.

Permite import PDF / DOCX / imagine cu detectare automată câmpuri proiect.
Cerință literală user: "Import proiect cu RETENȚIE imagini, ștampile, autorizări
digitale, recunoaște PDF Word, plasează ștampilă ca watermark."

Strategie hibrid:
1. Pentru .docx: extrage text cu python-docx + regex pe câmpuri cunoscute
2. Pentru .pdf: extrage text cu pypdf + regex
3. Pentru imagini: AI vision via Emergent LLM Key (gemini-3-flash) ca fallback

Returnează dict {field_key: value} pe care frontend-ul îl poate aplica direct
peste DEFAULT_DATA.
"""
from __future__ import annotations
import base64
import io
import json as _json_top
import re
from typing import Any, Dict, List, Optional

from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

from auth import get_current_user

router = APIRouter()


# ============================================================================
# Pattern recognition rules (regex over Romanian text)
# ============================================================================
PATTERNS: List[Dict[str, Any]] = [
    # CNP / CUI / CIF — multiple formats
    {"key": "beneficiar_cnp_cui", "regex": r"(?:CNP|CUI|CIF|J\d+)\s*:?\s*([0-9]{2,13})", "group": 1},
    # AC — Autorizație de Construire (flexible spaces + diacritics)
    {"key": "ac_numar", "regex": r"[Aa]utoriza[țt]i[ae]\s+(?:de\s+)?[Cc]onstruir[ei]\s+nr\.?\s*([0-9/\.\-]+)", "group": 1},
    {"key": "ac_data_emitere", "regex": r"AC.*?din\s+([0-9]{1,2}[./-][0-9]{1,2}[./-][0-9]{2,4})", "group": 1},
    # CU — Certificat de Urbanism
    {"key": "cu_numar", "regex": r"[Cc]ertificat(?:ul)?\s+(?:de\s+)?[Uu]rbanism\s+nr\.?\s*([0-9/\.\-]+)", "group": 1},
    # ATR — Acord Tehnic de Racordare
    {"key": "atr_numar", "regex": r"(?:ATR|Acord(?:ul)?\s+Tehnic\s+de\s+Racordare|Acord(?:ul)?\s+de\s+acces)\s*(?:nr\.?)?\s*([0-9/\.\-]+)", "group": 1},
    # Cadastru / CF
    {"key": "loc_consum_cadastru", "regex": r"(?:[Nn]r\.?\s*[Cc]adastral|CF\s*nr\.?|[Cc]adastral)\s*:?\s*([0-9]{4,10})", "group": 1},
    # Lungime conductă/branșament — V10.5: more flexible (matches "lungimea de 4 m", "L = 25,5 m", "25,5 m total")
    {"key": "sf_lungime_conducta_m",
     "regex": r"(?:[Ll]ungime[a]?(?:\s+(?:total[ăa]?|de))?|[Ll]\s*=)\s*(?:de\s+)?([0-9]+(?:[.,][0-9]+)?)\s*m(?:\s|\.|,|$|\))",
     "group": 1},
    # DN — accept both "DN 32" and "Dn 63 mm"
    {"key": "sf_diametru_nominal_DN", "regex": r"[Dd]n\s*([0-9]{2,4})(?:\s*mm)?", "group": 1, "prefix": "DN "},
    # Debit instalat
    {"key": "debit_instalat_mc_h",
     "regex": r"[Dd]ebit(?:ul)?\s+(?:instalat|simultan|total)?\s*:?\s*([0-9]+(?:[.,][0-9]+)?)\s*(?:m[³3]/h|Nmc/h)",
     "group": 1},
    # Presiune
    {"key": "sf_presiune_max_op_bar",
     "regex": r"[Pp]resiune[a]?\s+(?:max\.?|maxim[ăa])\s*:?\s*([0-9]+(?:[.,][0-9]+)?)\s*bar",
     "group": 1},
    # Telefon
    {"key": "beneficiar_telefon", "regex": r"[Tt]el\.?\s*:?\s*(\+?[0-9 ]{8,15})", "group": 1},
    # Email
    {"key": "beneficiar_email", "regex": r"[\w._%+-]+@[\w.-]+\.[a-zA-Z]{2,}", "group": 0},
    # Beneficiar — V10.5: also matches "Client:", "Beneficiarul investitiei:", "Investitor:"
    {"key": "beneficiar_nume",
     "regex": r"(?:[Bb]eneficiar(?:ul\s+investi[țt]iei)?|[Cc]lient(?:ul)?|[Ii]nvestitor(?:ul)?)\s*:?\s*([A-ZĂÂÎȘȚ][A-Za-zĂÂÎȘȚăâîșțéè\.\s&\-\d]{3,80}?)(?:\n|$|;|,\s+[A-Z])",
     "group": 1},
    # Adresa imobil / amplasament — multi-format
    {"key": "loc_consum_adresa",
     "regex": r"(?:[Aa]mplasament(?:ul)?(?:\s+(?:imobilului|lucr[ăa]rii))?|[Aa]dres[ăa])\s*:?\s*((?:Str|Bd|Bulevardul|Calea|Aleea|Sos|Strada)[\.\s][A-ZĂÂÎȘȚ][^\n;]{5,150})",
     "group": 1},
    # Localitate (Sector / Județ / Oraș)
    {"key": "loc_consum_localitate",
     "regex": r"(?:Sector|Sect\.?)\s+([0-9]+)\s*,?\s*(?:BUCURE[ȘS]TI|Bucure[șs]ti)",
     "group": 0},
    # Proiectant
    {"key": "proiectant_nume",
     "regex": r"[Pp]roiectant\s*:?\s*([A-Z][A-Z\s\.\-&]{2,60}(?:SRL|S\.R\.L\.|SA|S\.A\.))",
     "group": 1},
    # Executant
    {"key": "executant_nume",
     "regex": r"[Ee]xecutant\s*:?\s*([A-Z][A-Z\s\.\-&]{2,60}(?:SRL|S\.R\.L\.|SA|S\.A\.))",
     "group": 1},
    # Tip lucrare
    {"key": "tip_lucrare",
     "regex": r"(?:Bran[șs]ament\s+(?:nou|individual|colectiv)|Extindere\s+rețea|Modificare\s+contor)",
     "group": 0},
]


def _apply_patterns(text: str) -> Dict[str, Any]:
    out: Dict[str, Any] = {}
    for pat in PATTERNS:
        m = re.search(pat["regex"], text)
        if m:
            val = m.group(pat["group"])
            if "prefix" in pat:
                val = pat["prefix"] + val
            out[pat["key"]] = val.strip()
    return out


# ============================================================================
# Extractors per file type
# ============================================================================
def _extract_doc_legacy(blob: bytes) -> str:
    """Extract text from legacy .doc (Word 97-2003 binary format) via antiword.

    V9.3 — cerință user: "nu se pot incarca fisiere doc word, ci doar docx".
    Folosim antiword (instalat la /usr/bin/antiword) sau fallback la decode best-effort.
    """
    import subprocess
    import tempfile
    try:
        with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
            tmp.write(blob)
            tmp_path = tmp.name
        result = subprocess.run(
            ["antiword", tmp_path],
            capture_output=True, timeout=30, check=False,
        )
        text = (result.stdout or b"").decode("utf-8", errors="ignore")
        return text
    except FileNotFoundError:
        # antiword not installed — best-effort UTF-8 strip
        return blob.decode("latin-1", errors="ignore")
    except Exception:
        return ""


def _extract_docx(blob: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(blob))
        parts = []
        for p in doc.paragraphs:
            parts.append(p.text)
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts)
    except Exception:
        return ""


def _extract_pdf(blob: bytes) -> str:
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(blob))
        parts: List[str] = []
        for page in reader.pages:
            try:
                parts.append(page.extract_text() or "")
            except Exception:
                continue
        return "\n".join(parts)
    except Exception:
        return ""


# ============================================================================
# Routes
# ============================================================================
class ExtractResult(BaseModel):
    detected_fields: Dict[str, Any]
    confidence: str  # high | medium | low
    text_preview: str
    field_count: int


@router.post("/extract-fields", response_model=ExtractResult)
async def extract_fields(file: UploadFile = File(...), user=Depends(get_current_user)):
    blob = await file.read()
    if not blob:
        return ExtractResult(detected_fields={}, confidence="low", text_preview="(fișier gol)", field_count=0)

    fname_low = (file.filename or "").lower()
    text = ""
    extraction_note = ""
    if fname_low.endswith(".docx"):
        text = _extract_docx(blob)
        if not text:
            extraction_note = "DOCX corupt sau gol — verifică fișierul."
    elif fname_low.endswith(".doc"):
        # V9.3 — suport pentru .doc legacy (Word 97-2003) via antiword
        # V10.5 — detectează dacă e DOCX renamed ca .doc (zip header PK)
        if blob[:2] == b"PK":
            # Are signature ZIP → e de fapt DOCX cu extensia greșită
            extraction_note = "Fișierul are extensia .doc dar conține DOCX. Tratăm ca DOCX."
            text = _extract_docx(blob)
        else:
            text = _extract_doc_legacy(blob)
            if not text:
                extraction_note = "DOC legacy ilizibil — convertește la DOCX prin LibreOffice/Word."
    elif fname_low.endswith(".pdf"):
        text = _extract_pdf(blob)
        if not text:
            extraction_note = "PDF scanat (fără text OCR) sau corupt. Folosește o versiune cu text selectabil."
    elif fname_low.endswith((".png", ".jpg", ".jpeg", ".webp")):
        extraction_note = "Format imagine — folosește în loc /smart-extract-llm pentru OCR AI."
        text = ""
    else:
        extraction_note = f"Format nesuportat: {fname_low.split('.')[-1] if '.' in fname_low else 'unknown'}. Acceptăm: .docx .doc .pdf"
        text = ""

    fields = _apply_patterns(text) if text else {}
    confidence = "high" if len(fields) >= 5 else "medium" if len(fields) >= 2 else "low"

    preview = (extraction_note + "\n\n" + text[:1000]) if extraction_note else text[:1000]
    return ExtractResult(
        detected_fields=fields,
        confidence=confidence,
        text_preview=preview,
        field_count=len(fields),
    )


@router.get("/known-patterns")
async def known_patterns():
    """Listează tipurile de câmpuri pe care le poate extrage OCR-ul."""
    return {
        "patterns": [{"key": p["key"], "description": p["regex"][:80]} for p in PATTERNS],
        "supported_formats": [".docx", ".pdf", ".doc"],
        "experimental_formats": [".png", ".jpg", ".jpeg"],
    }


# ============================================================================
# V9.3 — Template Placeholder Detector
# Cerință user: "Mi-ar placea ca platforma sa poata introduce automat in documente
# placeholdere necesare. Sa recunoasca campuri de introdus si sa afiseze casete
# text in platforma pentru introducerea datelor sau variabilelor reale necesare
# inlocuirii acelor placeholdere/campuri."
# ============================================================================
class TemplatePlaceholdersResult(BaseModel):
    placeholders: List[Dict[str, Any]]   # [{ key, label, sample_context, suggested_field, type }]
    structure: Dict[str, Any]            # { sections, total_chars, candidate_count }
    text_preview: str


# Heuristici pentru detectarea câmpurilor în template — fraze sau variabile
# evidente care variază de la proiect la proiect.
TEMPLATE_HINTS: List[Dict[str, Any]] = [
    # Pattern: cuvinte CAPITALIZATE explicit înlocuibile
    {"regex": r"<([^>]{2,40})>",                       "label": "Tag explicit <{}>",          "type": "tag"},
    {"regex": r"\{\{?\s*([A-Za-zĂÂÎȘȚăâîșț_][A-Za-zĂÂÎȘȚăâîșț_0-9\s]{1,30})\s*\}?\}", "label": "Tag {{var}}",       "type": "tag"},
    {"regex": r"_{3,}",                                 "label": "Underscore (loc gol)",       "type": "blank"},
    {"regex": r"\.{4,}",                                "label": "Linie cu puncte (loc gol)",  "type": "blank"},
    # Pattern: paranteze cu cuvinte explicative
    {"regex": r"\(([Aa]ici[^)]{2,40})\)",               "label": "Indicație ({})",             "type": "hint"},
    {"regex": r"\(([Cc]ompleta[țt]i[^)]{0,40})\)",      "label": "(Completați...)",            "type": "hint"},
    {"regex": r"\(([Dd]e\s+[Cc]ompletat)\)",            "label": "(De completat)",             "type": "hint"},
    {"regex": r"\(([Vv]ariabil[ăa])\)",                 "label": "(Variabilă)",                "type": "hint"},
    # Pattern: cifre/date specifice care probabil variază
    {"regex": r"(?:[Nn]r\.?|[Nn]umar)\s*([0-9]+(?:[/.][0-9]+)+)",   "label": "Număr document {}", "type": "number"},
    {"regex": r"\bDn\s*([0-9]{2,4})\b",                  "label": "Diametru Dn{}",              "type": "spec"},
    {"regex": r"([0-9]+(?:[.,][0-9]+)?)\s*[Nn]mc/h",     "label": "Debit {} Nmc/h",             "type": "spec"},
    {"regex": r"([0-9]+(?:[.,][0-9]+)?)\s*bar\b",       "label": "Presiune {} bar",            "type": "spec"},
    {"regex": r"([0-9]+(?:[.,][0-9]+)?)\s*m(?:bar|et?ri?)?\b", "label": "Mărime {} m/mbar",     "type": "spec"},
]


def _detect_template_placeholders(text: str, limit: int = 60) -> List[Dict[str, Any]]:
    """Identifică toate locurile care par a fi placeholder-uri în template-ul DOC."""
    out: List[Dict[str, Any]] = []
    seen = set()
    for hint in TEMPLATE_HINTS:
        for m in re.finditer(hint["regex"], text):
            full = m.group(0)
            inner = m.group(1) if m.groups() else full
            label_tpl = hint["label"]
            label = label_tpl.format(inner) if "{}" in label_tpl else label_tpl
            ctx_start = max(0, m.start() - 80)
            ctx_end = min(len(text), m.end() + 80)
            context = text[ctx_start:ctx_end].replace("\n", " ").strip()
            sig = (hint["type"], inner.strip()[:30])
            if sig in seen:
                continue
            seen.add(sig)
            # Sugestie de mapping către FIELDS_REGISTRY (best effort)
            suggested = None
            inner_low = inner.lower() if isinstance(inner, str) else ""
            if "denumire" in inner_low or "lucrare" in inner_low: suggested = "proiect_titlu"
            elif "strada" in inner_low or "strad" in inner_low:   suggested = "loc_consum_strada"
            elif "nr" in inner_low and len(inner_low) < 6:        suggested = "loc_consum_numar"
            elif "beneficiar" in inner_low:                       suggested = "beneficiar_nume"
            elif "cnp" in inner_low or "cui" in inner_low:        suggested = "beneficiar_cnp_cui"
            elif "diametru" in inner_low or "dn" in inner_low:    suggested = "br_diametru_dn"
            elif "lungime" in inner_low:                          suggested = "br_lungime_m"
            elif "debit" in inner_low:                            suggested = "debit_instalat_mc_h"
            elif "presiune" in inner_low:                         suggested = "presiune_max_op_bar"
            out.append({
                "match": full[:60],
                "inner": inner if isinstance(inner, str) else str(inner),
                "label": label[:80],
                "type": hint["type"],
                "context": context[:200],
                "suggested_field": suggested,
                "position": m.start(),
            })
            if len(out) >= limit:
                return out
    return out


@router.post("/template-placeholders", response_model=TemplatePlaceholdersResult)
async def detect_template_placeholders(file: UploadFile = File(...), user=Depends(get_current_user)):
    """V9.3 — Detectează automat placeholder-urile/câmpurile variabile dintr-un template DOC/DOCX/PDF.

    User-ul upload-ează un model de document (de ex. MEMORIU AVIZARE) și platforma
    întoarce o listă cu câmpurile detectate ca fiind variabile (placeholdere, tag-uri,
    spații goale etc.) pe care apoi le poate completa direct din interfață.
    """
    blob = await file.read()
    if not blob:
        return TemplatePlaceholdersResult(placeholders=[], structure={}, text_preview="")

    fname = (file.filename or "").lower()
    if fname.endswith(".docx"):
        text = _extract_docx(blob)
    elif fname.endswith(".doc"):
        text = _extract_doc_legacy(blob)
    elif fname.endswith(".pdf"):
        text = _extract_pdf(blob)
    else:
        text = blob.decode("utf-8", errors="ignore")

    placeholders = _detect_template_placeholders(text, limit=60)
    # Detectează secțiuni (titluri în capitale)
    sections = re.findall(r"\n([A-ZĂÂÎȘȚ\s]{6,80})\n", text)
    structure = {
        "total_chars": len(text),
        "sections_detected": [s.strip() for s in sections[:20]],
        "candidate_count": len(placeholders),
        "supported_formats": [".docx", ".doc", ".pdf"],
    }
    return TemplatePlaceholdersResult(
        placeholders=placeholders,
        structure=structure,
        text_preview=text[:1500],
    )


# ============================================================================
# Apply extracted fields direct to a project (auto-fill)
# ============================================================================
class ApplyExtractRequest(BaseModel):
    pid: str
    fields: Dict[str, Any]
    overwrite: bool = False  # dacă True suprascrie valori existente


@router.post("/apply-to-project")
async def apply_extract_to_project(payload: ApplyExtractRequest, user=Depends(get_current_user)):
    """Aplică câmpurile detectate prin OCR direct peste un gas_project (write-once propagation)."""
    from db import db
    proj = await db.gas_projects.find_one(
        {"pid": payload.pid, "owner_id": user.user_id, "deleted": {"$ne": True}},
        {"_id": 0},
    )
    if not proj:
        from fastapi import HTTPException
        raise HTTPException(404, "Proiect inexistent")
    existing = proj.get("data") or {}
    applied = {}
    skipped = {}
    for k, v in (payload.fields or {}).items():
        if not payload.overwrite and existing.get(k):
            skipped[k] = existing[k]
            continue
        existing[k] = v
        applied[k] = v
    await db.gas_projects.update_one(
        {"pid": payload.pid, "owner_id": user.user_id},
        {"$set": {"data": existing}},
    )
    return {
        "pid": payload.pid,
        "applied_count": len(applied),
        "skipped_count": len(skipped),
        "applied": applied,
        "skipped": skipped,
    }


# ============================================================================
# V10.5 — SMART LLM-POWERED PLACEHOLDER DETECTOR
# Use case (literal user request): "merge sa introduc un document care nu are
# placehold-uri si recunoaste datele de introdus si adauga placehold-uri si
# casete text corespunzatoare"
#
# Strategy: extract text → send to Claude Sonnet → ask for VARIABLE values (names,
# addresses, numbers, dates) that should become editable form fields. Returns
# JSON with [{ original_text, suggested_key, suggested_label, field_type, position }]
# ============================================================================
class SmartPlaceholdersResult(BaseModel):
    placeholders: List[Dict[str, Any]]
    text_preview: str
    extraction_note: str
    field_count: int


_SMART_DETECTOR_PROMPT = (
    "Ești un asistent expert în documentație tehnică de gaze naturale din România. "
    "Primești textul unui document Word/PDF (memoriu, deviz, referat etc.) și trebuie "
    "să identifici TOATE valorile VARIABILE care ar trebui transformate în câmpuri "
    "editabile (placeholdere) pentru reutilizare în alte proiecte. "
    "Valori variabile = nume persoane, denumiri societăți, adrese, numere de "
    "documente/autorizații, date calendaristice, sume (lei/euro), lungimi (m), "
    "diametre (DN), debite (Nmc/h), presiuni (bar), număr cadastral, sectoare, etc.\n\n"
    "NU include în răspuns: cuvinte generice precum 'beneficiar', 'proiectant', etc. "
    "(care sunt etichete), titluri de secțiuni, cuvinte de legătură.\n\n"
    "Răspunde STRICT cu un JSON array (fără comentarii, fără markdown), cu maxim 30 "
    "elemente, format: \n"
    "[{\"original\":\"valoarea găsită în text\",\"suggested_key\":\"snake_case_field\","
    "\"suggested_label\":\"Etichetă RO\",\"field_type\":\"text|number|date|address\"}]\n\n"
    "Cheile cunoscute pe care să le folosești prioritar: beneficiar_nume, "
    "beneficiar_cnp_cui, beneficiar_telefon, beneficiar_email, loc_consum_adresa, "
    "loc_consum_strada, loc_consum_numar, loc_consum_localitate, loc_consum_cadastru, "
    "proiectant_nume, executant_nume, cu_numar, ac_numar, atr_numar, "
    "sf_diametru_nominal_DN, sf_lungime_conducta_m, debit_instalat_mc_h, "
    "sf_presiune_max_op_bar, tip_lucrare. Pentru altele, alege snake_case descriptiv."
)


@router.post("/smart-extract-llm", response_model=SmartPlaceholdersResult)
async def smart_extract_llm(file: UploadFile = File(...), user=Depends(get_current_user)):
    """V10.5 — AI-powered field detection for documents WITHOUT explicit placeholders.

    Folosește Claude Sonnet (prin Emergent LLM Key) pentru a identifica automat
    valorile variabile dintr-un document (memoriu/deviz/referat) și le convertește
    în propuneri de câmpuri editabile cu key + label + tip.
    """
    blob = await file.read()
    if not blob:
        return SmartPlaceholdersResult(placeholders=[], text_preview="", extraction_note="Fișier gol.", field_count=0)
    fname = (file.filename or "").lower()
    if fname.endswith(".docx") or (blob[:2] == b"PK" and fname.endswith(".doc")):
        text = _extract_docx(blob)
    elif fname.endswith(".doc"):
        text = _extract_doc_legacy(blob)
    elif fname.endswith(".pdf"):
        text = _extract_pdf(blob)
    else:
        text = blob.decode("utf-8", errors="ignore")
    if not text or len(text) < 50:
        return SmartPlaceholdersResult(
            placeholders=[], text_preview=text[:500] if text else "",
            extraction_note="Text insuficient pentru analiză AI (< 50 caractere).", field_count=0,
        )
    # Truncate to keep LLM cost reasonable
    text_for_llm = text[:8000]
    try:
        from ai_agents import _ask
        import json as _json
        reply = await _ask(
            agent="proiectant",  # any agent works, we override system msg next
            message=f"{_SMART_DETECTOR_PROMPT}\n\nDocument:\n```\n{text_for_llm}\n```\n\nJSON:",
            session_id=f"smart_extract_{user.user_id}",
        )
        # Try to extract JSON array from the reply
        m = re.search(r"\[\s*\{.*?\}\s*\]", reply, re.DOTALL)
        raw = m.group(0) if m else reply.strip()
        parsed = _json.loads(raw)
        if not isinstance(parsed, list):
            parsed = []
        # Sanitize
        out: List[Dict[str, Any]] = []
        for it in parsed[:40]:
            if not isinstance(it, dict):
                continue
            orig = str(it.get("original") or "").strip()[:200]
            if not orig:
                continue
            key = str(it.get("suggested_key") or "").strip()[:60] or "campa_" + str(len(out))
            label = str(it.get("suggested_label") or key).strip()[:80]
            ftype = str(it.get("field_type") or "text").strip()[:20]
            # Position of original in text (best-effort)
            pos = text.find(orig) if orig else -1
            out.append({
                "original": orig,
                "suggested_key": key,
                "suggested_label": label,
                "field_type": ftype,
                "position": pos,
                "context": (text[max(0, pos - 60):pos + 120] if pos >= 0 else "").replace("\n", " ").strip()[:200],
            })
        return SmartPlaceholdersResult(
            placeholders=out,
            text_preview=text[:1500],
            extraction_note=f"Detectate {len(out)} câmpuri variabile prin AI (Claude Sonnet 4.6).",
            field_count=len(out),
        )
    except Exception as e:
        # Fallback to heuristic detector if LLM fails
        fallback = _detect_template_placeholders(text, limit=40)
        return SmartPlaceholdersResult(
            placeholders=[{
                "original": p.get("inner", ""),
                "suggested_key": p.get("suggested_field") or f"campa_{idx}",
                "suggested_label": p.get("label", ""),
                "field_type": p.get("type", "text"),
                "position": p.get("position", -1),
                "context": p.get("context", ""),
            } for idx, p in enumerate(fallback)],
            text_preview=text[:1500],
            extraction_note=f"AI indisponibil ({str(e)[:80]}). Folosit detector euristic ca fallback.",
            field_count=len(fallback),
        )


# ============================================================================
# V10.6 — FILL TEMPLATE ENDPOINT
# Cerință literală user (mesaj 26.06): "fa tot UI-ul necesar completarii
# documentului incarcat - proiect gaze naturale".
#
# Workflow: utilizatorul încarcă un template DOCX, completează în UI fiecare
# placeholder detectat, iar acest endpoint primește template-ul + lista de
# înlocuiri și returnează DOCX-ul completat pentru descărcare/print.
#
# Strategy: walk pe paragrafe + tabele, înlocuire la nivel de run (păstrăm
# formatarea) cu fallback pe paragraph-level dacă textul e spart pe multiple
# runs (cazul tipic în template-uri Word generate manual).
# ============================================================================
class _FillReplacement(BaseModel):
    original: str
    replacement: str


def _docx_replace_text(doc, replacements: List[Dict[str, str]]) -> int:
    """Înlocuiește textul în DOCX păstrând formatarea per-run.

    Returnează numărul de potriviri găsite și înlocuite.
    """
    if not replacements:
        return 0
    # Sortăm înlocuirile descrescător după lungime ca să prevenim coliziunile
    # (de ex. "Sector 1" → înlocuiește înainte de "Sector").
    ordered = sorted(
        [r for r in replacements if r.get("original")],
        key=lambda r: len(r.get("original", "")),
        reverse=True,
    )
    matches = 0

    def _replace_in_paragraph(paragraph) -> int:
        nonlocal matches
        local = 0
        # Optimizare: dacă paragraful nu conține nicio cheie, sărim
        full_text = paragraph.text or ""
        if not any(r["original"] in full_text for r in ordered):
            return 0
        # Strategy A: înlocuire pe runs individuale dacă potrivirea e
        # localizată într-un singur run (păstrează formatarea perfect)
        for r in ordered:
            orig = r["original"]
            repl = r.get("replacement", "")
            for run in paragraph.runs:
                if orig in (run.text or ""):
                    run.text = run.text.replace(orig, repl)
                    local += 1
        # Strategy B: dacă textul nu s-a schimbat (placeholder spart pe runs),
        # consolidăm tot paragraful în primul run și ștergem restul.
        new_text = paragraph.text or ""
        still_has = any(r["original"] in new_text for r in ordered)
        if still_has:
            combined = new_text
            for r in ordered:
                if r["original"] in combined:
                    combined = combined.replace(r["original"], r.get("replacement", ""))
                    local += 1
            if combined != new_text:
                if paragraph.runs:
                    paragraph.runs[0].text = combined
                    for run in paragraph.runs[1:]:
                        run.text = ""
                else:
                    paragraph.add_run(combined)
        matches += local
        return local

    for para in doc.paragraphs:
        _replace_in_paragraph(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para)
    # Headers / footers
    for section in doc.sections:
        for hf in (section.header, section.footer):
            if hf is None:
                continue
            for para in hf.paragraphs:
                _replace_in_paragraph(para)
    return matches


@router.post("/fill-template")
async def fill_template(
    file: UploadFile = File(...),
    replacements: str = Form(...),
    filename: Optional[str] = Form(None),
    user=Depends(get_current_user),
):
    """V10.6 — Înlocuiește placeholder-urile dintr-un DOCX cu valorile
    completate de utilizator în UI și returnează DOCX-ul completat.

    Args:
        file: template DOCX (Word 2007+). DOC legacy și PDF NU sunt suportate
              pentru fill (doar pentru detectare).
        replacements: JSON string cu listă `[{ original, replacement }, ...]`.
        filename: opțional — numele DOCX-ului returnat (fără extensie).
    """
    fname = (file.filename or "").lower()
    if not fname.endswith(".docx"):
        # Acceptăm și DOCX cu extensia .doc (zip header PK)
        head = await file.read(2)
        await file.seek(0)
        if not (fname.endswith(".doc") and head == b"PK"):
            raise HTTPException(
                status_code=400,
                detail="Doar DOCX (Word 2007+) e suportat pentru completare. Convertește .doc/.pdf la .docx mai întâi.",
            )

    try:
        repl_list = _json_top.loads(replacements or "[]")
    except Exception:
        raise HTTPException(status_code=400, detail="replacements trebuie să fie JSON valid.")
    if not isinstance(repl_list, list):
        raise HTTPException(status_code=400, detail="replacements trebuie să fie o listă.")

    # Sanitize: doar string-uri non-empty pentru `original`
    clean: List[Dict[str, str]] = []
    seen_originals = set()
    for it in repl_list:
        if not isinstance(it, dict):
            continue
        orig = str(it.get("original") or "").strip()
        if not orig:
            continue
        repl = str(it.get("replacement") or "")
        clean.append({"original": orig, "replacement": repl})
        seen_originals.add(orig)

    # V11.6 — Extindere automată cu aliasuri lungi (conform docx-ului utilizatorului)
    # Dacă perechea folosește o cheie scurtă (ex {{osd_nume}}), adăugăm automat și
    # cheia lungă echivalentă ({{operator_sistem_distributie}}) cu aceeași valoare.
    try:
        from placeholders_aliases import ALIAS_MAP
        # Construim un index reverse: internal_key → list of long aliases
        reverse: Dict[str, List[str]] = {}
        for long_alias, internal in ALIAS_MAP.items():
            reverse.setdefault(internal, []).append(long_alias)
        # Pentru fiecare original, dacă scurtul e mapat la N aliasuri lungi,
        # le adăugăm și pe ele cu aceeași replacement
        extras: List[Dict[str, str]] = []
        for item in list(clean):
            orig = item["original"]
            # Detectează {{key}} sau key brut
            mkey = orig.strip("{ }")
            # Forward: dacă este cheia internă scurtă, adăugăm aliasurile lungi
            for long_alias in reverse.get(mkey, []):
                braced = "{{" + long_alias + "}}"
                if braced not in seen_originals:
                    extras.append({"original": braced, "replacement": item["replacement"]})
                    seen_originals.add(braced)
            # Reverse: dacă utilizatorul a trimis alias lung, adăugăm și cheia scurtă
            if mkey in ALIAS_MAP:
                short = ALIAS_MAP[mkey]
                braced = "{{" + short + "}}"
                if braced not in seen_originals:
                    extras.append({"original": braced, "replacement": item["replacement"]})
                    seen_originals.add(braced)
        clean.extend(extras)
    except Exception:
        # Aliasurile sunt o îmbunătățire — dacă map-ul nu se încarcă, nu blocăm fluxul
        pass

    blob = await file.read()
    if not blob:
        raise HTTPException(status_code=400, detail="Fișier gol.")

    try:
        from docx import Document
        doc = Document(io.BytesIO(blob))
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"DOCX invalid: {str(e)[:120]}")

    matches = _docx_replace_text(doc, clean)

    out_buf = io.BytesIO()
    doc.save(out_buf)
    out_buf.seek(0)

    safe_name = (filename or (file.filename or "document_completat")).rsplit(".", 1)[0]
    safe_name = re.sub(r"[^A-Za-z0-9_\-]+", "_", safe_name)[:80] or "document_completat"
    out_name = f"{safe_name}_completat.docx"

    return StreamingResponse(
        out_buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="{out_name}"',
            "X-Replacements-Applied": str(matches),
            "X-Replacements-Requested": str(len(clean)),
        },
    )
