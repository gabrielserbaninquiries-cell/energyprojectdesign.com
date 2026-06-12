"""
gas_doc_templates_legal — 7 documente legale ROMÂNEȘTI obligatorii lipsă din V8.0:

1. declaratie_conformitate    — Declarație de conformitate executant (Lege 10/1995 art. 25 + EN 1555)
2. buletin_proba_rezistenta   — Buletin probă rezistență (NTPEE 2018 cap. 5)
3. buletin_proba_etanseitate  — Buletin probă etanșeitate 24h (NTPEE 2018 cap. 5)
4. pv_receptie_finala         — PV Recepție Finală PVRF (HG 273/1994 art. 36)
5. pv_pif_semnat              — PV Punere în Funcțiune semnat OSD (Ord. ANRE 162/2021)
6. fisa_sudor                 — Fișa sudor autorizat per sudor (ANRE Ord. 79/2014 + EN 13067)
7. plan_ssm                   — Plan SSM Securitate Sănătate Muncă (Legea 319/2006)
"""
from typing import Any, Dict
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

import gas_doc_templates as base
COMPANY = base.COMPANY


def _declaratie_conformitate(proj: Dict[str, Any]) -> bytes:
    data = proj.get("data") or {}
    doc = Document()
    base._company_header(doc)
    base._add_heading(doc, "DECLARAȚIE DE CONFORMITATE A EXECUTANTULUI", level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    base._add_para(doc, "(Conform Legea 10/1995 art. 25 + EN 1555 + NTPEE 2018)",
                   italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
    doc.add_paragraph()
    base._add_kv_table(doc, [
        ("DC nr. / dată",          base._get(data, "dc_numar_data", base._today_ro())),
        ("Obiectiv",               proj.get("title", "—")),
        ("Beneficiar",             base._get(data, "beneficiar_nume")),
        ("Adresă",                 base._get(data, "loc_consum_adresa")),
        ("Operator distribuție",   base._get(data, "atr_osd")),
        ("AC nr.",                 f"{base._get(data,'ac_numar')} / {base._get(data,'ac_data_emitere')}"),
        ("Executant",              base._get(data, "exec_firma")),
        ("Autorizație ANRE",       f"{base._get(data,'executant_aut_nr')} ({base._get(data,'executant_aut_grad','EDD')})"),
        ("Responsabil tehnic (RTE)", base._get(data, "exec_responsabil_tehnic")),
    ])
    doc.add_paragraph()
    base._add_para(doc, "Subsemnatul, reprezentant legal al executantului, declar pe propria răspundere că:", bold=True)
    declaratii = [
        "1. Lucrările au fost executate conform proiectului tehnic aprobat și a NTPEE 2018.",
        "2. Materialele puse în operă (țeavă PE 100 SDR 11 EN 1555, robinet branșament EN 1555-4, regulator EN 88-1, contor EN 1359) au certificate de conformitate atașate.",
        "3. Sudurile au fost efectuate de sudori autorizați ANRE conform EN 13067 (fișe sudori atașate).",
        "4. Probele de rezistență și etanșeitate au fost efectuate cu rezultate ADMISIBILE (buletine atașate).",
        "5. Au fost respectate distanțele minime de siguranță față de alte rețele (electric 0.5m, apă-canal 1.0m, fundații 3.0m).",
        "6. Adâncimea de pozare respectă NTPEE 2018 art. 56 (min. 0.9m sub trotuar, 1.0m sub trafic).",
        "7. Banda de avertizare galbenă este poziționată la 30cm deasupra conductei pe tot traseul.",
        "8. Documentele de execuție (jurnal șantier, PV-uri faze determinante, PV lucrări ascunse, declarații sudori) sunt complete și atașate.",
    ]
    for d in declaratii:
        doc.add_paragraph(d, style="List Number")
    doc.add_paragraph()
    base._add_para(doc, "Garanție acordată:", bold=True)
    base._add_kv_table(doc, [
        ("Conducte și sudurile aferente", base._get(data, "garantie_conducta", "10 ani de la PVRTL")),
        ("Robinet branșament + regulator", base._get(data, "garantie_robinet", "5 ani de la PVRTL")),
        ("Manoperă execuție",              base._get(data, "garantie_manopera", "2 ani de la PVRTL")),
    ])
    base._footer_signature(doc, proj, data)
    return base._save(doc)


def _buletin_proba_rezistenta(proj: Dict[str, Any]) -> bytes:
    data = proj.get("data") or {}
    doc = Document()
    base._company_header(doc)
    base._add_heading(doc, "BULETIN DE PROBĂ — REZISTENȚĂ MECANICĂ", level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    base._add_para(doc, "(Conform NTPEE 2018 cap. 5 art. 80)", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
    doc.add_paragraph()
    base._add_kv_table(doc, [
        ("Buletin nr. / dată",     base._get(data, "buletin_rez_nr_data", base._today_ro())),
        ("Obiectiv",               proj.get("title", "—")),
        ("Loc consum",             base._get(data, "loc_consum_adresa")),
        ("Tronson testat",         base._get(data, "buletin_rez_tronson", "Conductă PE 100 DN " + base._get(data, "sf_diametru_nominal_DN", "32"))),
        ("Lungime tronson (m)",    base._get(data, "sf_lungime_conducta_m") or base._get(data, "pt_lungime_m")),
        ("Material",               base._get(data, "sf_material_conducta", "PE 100 SDR 11")),
        ("Categorie presiune",     base._get(data, "presiune_categorie", "REDUSA PRESIUNE")),
    ])
    doc.add_paragraph()
    base._add_heading(doc, "Parametri proba de rezistență", level=2)
    base._add_kv_table(doc, [
        ("Presiune de probă (bar)",       base._get(data, "proba_rezistenta_bar", "6")),
        ("Coeficient de siguranță (KSP)", base._get(data, "ksp_rezistenta", "1.5 × Pmax_op")),
        ("Mediu de încercare",            base._get(data, "mediu_proba_rez", "Aer comprimat tehnic")),
        ("Durată proba (min)",            base._get(data, "durata_proba_rez_min", "60")),
        ("Manometru utilizat",            base._get(data, "manometru_serie", "Manometru calibrat clasa 1.0, seria CAL-2026/47")),
        ("Temperatură ambiantă (°C)",     base._get(data, "temp_proba_C", "18")),
    ])
    doc.add_paragraph()
    base._add_heading(doc, "Rezultate măsurători", level=2)
    base._add_kv_table(doc, [
        ("Presiune inițială (bar)",         base._get(data, "p_initiala_rez", "6.05")),
        ("Presiune după 60 min (bar)",      base._get(data, "p_finala_rez", "6.04")),
        ("Cădere presiune măsurată (bar)",  base._get(data, "delta_p_rez", "0.01")),
        ("Cădere admisă maxim (bar)",       base._get(data, "delta_p_max_rez", "0.10")),
        ("VERDICT",                         base._get(data, "verdict_rez", "ADMIS")),
    ])
    doc.add_paragraph()
    base._add_para(doc, "Participanți la probă:", bold=True)
    base._add_kv_table(doc, [
        ("Reprezentant proiectant", base._get(data, "dtac_proiectant_specialitate", COMPANY["name"])),
        ("Reprezentant executant",  base._get(data, "exec_firma")),
        ("Responsabil tehnic (RTE)", base._get(data, "exec_responsabil_tehnic")),
        ("Reprezentant OSD",        base._get(data, "atr_osd")),
    ])
    base._footer_signature(doc, proj, data)
    return base._save(doc)


def _buletin_proba_etanseitate(proj: Dict[str, Any]) -> bytes:
    data = proj.get("data") or {}
    doc = Document()
    base._company_header(doc)
    base._add_heading(doc, "BULETIN DE PROBĂ — ETANȘEITATE", level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    base._add_para(doc, "(Conform NTPEE 2018 cap. 5 art. 82 — proba 24 ore)", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
    doc.add_paragraph()
    base._add_kv_table(doc, [
        ("Buletin nr. / dată",     base._get(data, "buletin_et_nr_data", base._today_ro())),
        ("Obiectiv",               proj.get("title", "—")),
        ("Loc consum",             base._get(data, "loc_consum_adresa")),
        ("Tronson testat",         base._get(data, "buletin_et_tronson", "Tronson complet branșament + instalație utilizare")),
        ("Lungime tronson (m)",    base._get(data, "sf_lungime_conducta_m")),
        ("Categorie presiune",     base._get(data, "presiune_categorie")),
    ])
    doc.add_paragraph()
    base._add_heading(doc, "Parametri proba de etanșeitate", level=2)
    base._add_kv_table(doc, [
        ("Presiune de probă (bar)",     base._get(data, "proba_etanseitate_bar", "0.11")),
        ("Durată proba (ore)",          base._get(data, "durata_proba_et_h", "24")),
        ("Mediu de încercare",          base._get(data, "mediu_proba_et", "Aer comprimat tehnic")),
        ("Temperatură inițială (°C)",   base._get(data, "temp_proba_et_inceput", "18.5")),
        ("Temperatură finală (°C)",     base._get(data, "temp_proba_et_sfarsit", "17.2")),
        ("Compensare termică aplicată", base._get(data, "compensare_termica_et", "Da — corecție 0.1% / °C")),
    ])
    doc.add_paragraph()
    base._add_heading(doc, "Rezultate măsurători (24 ore)", level=2)
    base._add_kv_table(doc, [
        ("Presiune inițială (bar)",         base._get(data, "p_initiala_et", "0.110")),
        ("Presiune după 24h (bar)",         base._get(data, "p_finala_et",   "0.108")),
        ("Cădere presiune (bar)",           base._get(data, "delta_p_et",    "0.002")),
        ("Cădere admisă maxim (bar)",       base._get(data, "delta_p_max_et", "0.005")),
        ("VERDICT",                         base._get(data, "verdict_et", "ADMIS")),
    ])
    doc.add_paragraph()
    base._add_para(doc, "Mențiuni:", bold=True)
    base._add_para(doc, base._get(data, "mentiuni_et",
                                  "Probă efectuată conform NTPEE 2018 art. 82. Conducta a fost izolată "
                                  "de orice obiect care ar fi putut influența rezultatul (consumatori, atmosferă)."),
                   italic=True)
    base._footer_signature(doc, proj, data)
    return base._save(doc)


def _pv_receptie_finala(proj: Dict[str, Any]) -> bytes:
    data = proj.get("data") or {}
    doc = Document()
    base._company_header(doc)
    base._add_heading(doc, "PROCES-VERBAL DE RECEPȚIE FINALĂ (PVRF)", level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    base._add_para(doc, "(Conform HG 273/1994 art. 36 — recepție la 1-3 ani de la PVRTL)",
                   italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
    doc.add_paragraph()
    base._add_kv_table(doc, [
        ("PVRF nr.",          base._get(data, "pvrf_numar")),
        ("Data recepție finală", base._get(data, "pvrf_data", base._today_ro())),
        ("Obiectiv",          proj.get("title", "—")),
        ("Beneficiar",        base._get(data, "beneficiar_nume")),
        ("Loc consum",        base._get(data, "loc_consum_adresa")),
        ("PVRTL anterior",    f"{base._get(data,'receptie_pv_numar')} / {base._get(data,'receptie_pv_data')}"),
        ("AC nr.",            f"{base._get(data,'ac_numar')} / {base._get(data,'ac_data_emitere')}"),
        ("Perioadă garanție", base._get(data, "perioada_garantie", "10 ani conductă / 5 ani robinet / 2 ani manoperă")),
    ])
    doc.add_paragraph()
    base._add_heading(doc, "Comisia de recepție finală:", level=2)
    base._add_kv_table(doc, [
        ("Președinte",                base._get(data, "com_recep_presedinte")),
        ("Reprezentant OSD",          base._get(data, "com_recep_reprez_osd")),
        ("Reprezentant ISC",          base._get(data, "com_recep_reprez_isc")),
        ("Reprezentant beneficiar",   base._get(data, "com_recep_reprez_beneficiar")),
        ("Reprezentant executant",    base._get(data, "exec_firma")),
        ("Reprezentant proiectant",   base._get(data, "dtac_proiectant_specialitate", COMPANY["name"])),
    ])
    doc.add_paragraph()
    base._add_heading(doc, "Verificări efectuate", level=2)
    doc.add_paragraph("• Funcționare instalație de la PVRTL la PVRF — fără incidente raportate", style="List Bullet")
    doc.add_paragraph("• Verificare etanșeitate periodică (la 2 ani conform NTPEE 2018 art. 78)", style="List Bullet")
    doc.add_paragraph("• Verificare stare conductă, robinet, regulator, contor", style="List Bullet")
    doc.add_paragraph("• Verificare marcaje (banda avertizare, plăcuțe semnalizare)", style="List Bullet")
    doc.add_paragraph("• Verificare existență și conformitate Carte Tehnică", style="List Bullet")
    doc.add_paragraph()
    base._add_para(doc, "Constatări:", bold=True)
    base._add_para(doc, base._get(data, "pvrf_constatari",
                                  "Instalația a funcționat fără defecțiuni pe perioada de garanție. "
                                  "Nu s-au constatat scurgeri sau abateri de la parametrii de proiect. "
                                  "Beneficiarul a respectat obligațiile contractuale privind verificările periodice."),
                   italic=True)
    base._add_para(doc, "Concluzie:", bold=True)
    base._add_para(doc, base._get(data, "pvrf_concluzie",
                                  "Se aprobă RECEPȚIA FINALĂ a lucrării. Garanția acordată de executant "
                                  "se consideră ÎNCHEIATĂ. Beneficiarul preia integral responsabilitatea "
                                  "exploatării, mentenanței și siguranței instalației conform NTPEE 2018."),
                   bold=True)
    base._footer_signature(doc, proj, data)
    return base._save(doc)


def _pv_pif_semnat(proj: Dict[str, Any]) -> bytes:
    data = proj.get("data") or {}
    doc = Document()
    base._company_header(doc)
    base._add_heading(doc, "PROCES-VERBAL DE PUNERE ÎN FUNCȚIUNE (PIF)", level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    base._add_para(doc, "(Conform Ord. ANRE 162/2021 — semnat de OSD)",
                   italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
    doc.add_paragraph()
    base._add_kv_table(doc, [
        ("PV PIF nr.",            base._get(data, "pif_numar")),
        ("Data PIF",              base._get(data, "pif_data", base._today_ro())),
        ("Obiectiv",              proj.get("title", "—")),
        ("Beneficiar",            base._get(data, "beneficiar_nume")),
        ("Loc consum",            base._get(data, "loc_consum_adresa")),
        ("CLC (cod loc consum)",  base._get(data, "clc_cod")),
        ("Operator distribuție",  base._get(data, "atr_osd")),
        ("ATR nr.",               f"{base._get(data,'atr_numar')} / {base._get(data,'atr_data')}"),
        ("PVRTL nr.",             f"{base._get(data,'receptie_pv_numar')} / {base._get(data,'receptie_pv_data')}"),
        ("Contor seria",          base._get(data, "contor_serie")),
        ("Contor index PIF",      base._get(data, "contor_index_pif", "0.000 m³")),
        ("Presiune intrare (bar)",base._get(data, "pif_presiune_intrare", base._get(data, "sf_presiune_max_op_bar"))),
        ("Categorie consumator",  base._get(data, "categorie_consumator", "Casnic")),
    ])
    doc.add_paragraph()
    base._add_para(doc, "OSD certifică prin prezentul PV că:", bold=True)
    doc.add_paragraph("1. Instalația a fost verificată și corespunde proiectului DTAC + PTH aprobat.", style="List Number")
    doc.add_paragraph("2. Probele de etanșeitate înainte de PIF au fost efectuate și sunt ADMISE.", style="List Number")
    doc.add_paragraph("3. Contorul a fost montat și sigilat metrologic; indexul inițial este consemnat.", style="List Number")
    doc.add_paragraph("4. S-a făcut purjarea de aer și umplerea cu gaze naturale a instalației.", style="List Number")
    doc.add_paragraph("5. Beneficiarul a fost instruit privind utilizarea în siguranță (ANRE Ord. 16/2015).", style="List Number")
    doc.add_paragraph("6. Contractul de furnizare gaze a fost semnat cu furnizorul ales de beneficiar.", style="List Number")
    doc.add_paragraph()
    base._add_para(doc, "Mențiuni speciale OSD:", bold=True)
    base._add_para(doc, base._get(data, "pif_mentiuni_osd",
                                  "Instalația este pusă în funcțiune și activată în Sistemul Național de Transport Gaze. "
                                  "Beneficiarul are obligația revizuirii instalației la 2 ani (etanșeitate) și 10 ani (revizie integrală)."),
                   italic=True)
    doc.add_paragraph()
    sig_t = doc.add_table(rows=2, cols=3)
    sig_t.rows[0].cells[0].text = "OSD (semnatar)"
    sig_t.rows[0].cells[1].text = "Beneficiar"
    sig_t.rows[0].cells[2].text = "Executant prezent"
    sig_t.rows[1].cells[0].text = f"{base._get(data,'atr_osd','—')}\nReprezentant: {base._get(data,'pif_osd_reprezentant','—')}\n\nSemnătură + ștampilă:"
    sig_t.rows[1].cells[1].text = f"{base._get(data,'beneficiar_nume','—')}\n\nSemnătură:"
    sig_t.rows[1].cells[2].text = f"{base._get(data,'exec_firma','—')}\nRTE: {base._get(data,'exec_responsabil_tehnic','—')}\n\nSemnătură:"
    base._footer_signature(doc, proj, data)
    return base._save(doc)


def _fisa_sudor(proj: Dict[str, Any]) -> bytes:
    data = proj.get("data") or {}
    doc = Document()
    base._company_header(doc)
    base._add_heading(doc, "FIȘĂ SUDOR AUTORIZAT", level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    base._add_para(doc, "(Conform ANRE Ord. 79/2014 + EN 13067 + ISO 9606-1)",
                   italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
    doc.add_paragraph()
    base._add_kv_table(doc, [
        ("Fișa nr.",                   base._get(data, "sudor_fisa_numar")),
        ("Obiectiv",                   proj.get("title", "—")),
        ("Loc consum",                 base._get(data, "loc_consum_adresa")),
        ("Executant",                  base._get(data, "exec_firma")),
        ("Nume sudor",                 base._get(data, "sudor_nume")),
        ("CNP",                        base._get(data, "sudor_cnp")),
        ("Calificare (EN 13067)",      base._get(data, "sudor_calificare", "Sudor PE 100 — proceduri electrofuziune (EF) + cap-la-cap (BW)")),
        ("Autorizație ANRE nr.",       base._get(data, "sudor_autorizatie_nr")),
        ("Valabilitate autorizație",   base._get(data, "sudor_valabilitate", "2 ani de la emitere")),
        ("Organism atestare",          base._get(data, "sudor_organism", "ANRE / RAR / ISCIR")),
    ])
    doc.add_paragraph()
    base._add_heading(doc, "Date probe atestare", level=2)
    base._add_kv_table(doc, [
        ("Procedeu sudură",        base._get(data, "sudor_procedeu", "Electrofuziune cu manșoane (EF) + Sudură cap-la-cap (BW)")),
        ("Material încercat",      base._get(data, "sudor_material", "PE 100 SDR 11")),
        ("Diametre acoperite",     base._get(data, "sudor_diametre", "DN 20 ... DN 110")),
        ("Echipamente utilizate",  base._get(data, "sudor_echipamente", "Mașină EF Plasson tip 4400, Termofuzionator Wavin 110")),
        ("Test tracțiune (MPa)",   base._get(data, "sudor_tractiune", "23.5 (admis ≥20)")),
        ("Test impact",            base._get(data, "sudor_impact", "ADMIS — fără cedare la 4kJ/m²")),
    ])
    doc.add_paragraph()
    base._add_heading(doc, "Suduri efectuate pe acest obiectiv", level=2)
    base._add_kv_table(doc, [
        ("Total suduri",                base._get(data, "sudor_total_suduri", base._get(data, "asb_numar_sudari", "—"))),
        ("Verificate vizual (OK)",      base._get(data, "sudor_verificate_vizual", "100% — toate")),
        ("Verificate cu ultrasunete",   base._get(data, "sudor_us_verificate", "—")),
        ("Defecte constatate",          base._get(data, "sudor_defecte", "Niciun defect")),
    ])
    base._footer_signature(doc, proj, data)
    return base._save(doc)


def _plan_ssm(proj: Dict[str, Any]) -> bytes:
    data = proj.get("data") or {}
    doc = Document()
    base._company_header(doc)
    base._add_heading(doc, "PLAN DE SECURITATE ȘI SĂNĂTATE ÎN MUNCĂ (SSM)", level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    base._add_para(doc, "(Conform Legea 319/2006 + HG 1425/2006 + HG 300/2006)",
                   italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
    doc.add_paragraph()
    base._add_kv_table(doc, [
        ("Plan SSM nr.",           base._get(data, "ssm_numar_data", base._today_ro())),
        ("Obiectiv",               proj.get("title", "—")),
        ("Beneficiar",             base._get(data, "beneficiar_nume")),
        ("Loc consum",             base._get(data, "loc_consum_adresa")),
        ("Executant lucrare",      base._get(data, "exec_firma")),
        ("Responsabil SSM executant", base._get(data, "ssm_responsabil_exec", "—")),
        ("Coordonator SSM (HG 300/2006)", base._get(data, "ssm_coordonator")),
        ("Durată estimată lucrări (zile)", base._get(data, "ssm_durata_zile", "5")),
        ("Număr lucrători angajați",       base._get(data, "ssm_nr_lucratori", "3")),
    ])
    doc.add_paragraph()
    base._add_heading(doc, "1. Identificarea riscurilor", level=2)
    riscuri = [
        ("Săpătură șanț (cădere de la înălțime, surpare maluri)", "Mediu", "Sprijinire maluri >1.5m, balustradă perimetru, semnalizare"),
        ("Lucrul cu gaz (explozie, intoxicație CO/CH4)",          "Mare",  "Detector gaze portabil, ventilație forțată, interzicere foc deschis"),
        ("Manipulare materiale grele (PE coils, sudor)",          "Mediu", "Echipament EPP, instructaj, ridicare cu 2 persoane >20kg"),
        ("Trafic auto pe șantier (semnalizare carosabil)",         "Mediu", "Aviz Poliția Rutieră, panouri SR 1848-7, semnalizator"),
        ("Electrocutare (rețele subterane existente)",             "Mare",  "Detector cabluri, săpătură manuală în zone de risc"),
        ("Sudură PE (fum, temperatură 220°C aparat)",              "Mic",   "Mănuși termice, mască vapori, distanță de siguranță"),
    ]
    t = doc.add_table(rows=1, cols=3)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    hdr[0].text = "Risc identificat"; hdr[1].text = "Nivel"; hdr[2].text = "Măsuri de prevenire"
    for r, n, m in riscuri:
        rw = t.add_row().cells
        rw[0].text = r; rw[1].text = n; rw[2].text = m
    doc.add_paragraph()
    base._add_heading(doc, "2. Echipament Individual de Protecție (EIP)", level=2)
    doc.add_paragraph("• Cască de protecție EN 397 (clasă G450V)", style="List Bullet")
    doc.add_paragraph("• Vestă reflectorizantă EN ISO 20471 clasa 2", style="List Bullet")
    doc.add_paragraph("• Mănuși de protecție EN 388 (mecanic) + EN 12477 (termic la sudură)", style="List Bullet")
    doc.add_paragraph("• Bocanci de protecție EN ISO 20345 S3 (anti-perforare)", style="List Bullet")
    doc.add_paragraph("• Detector portabil gaze CH4 + CO (calibrat conform EN 60079)", style="List Bullet")
    doc.add_paragraph("• Mască respiratorie FFP2 (la sudură PE)", style="List Bullet")
    doc.add_paragraph()
    base._add_heading(doc, "3. Instructaj SSM obligatoriu", level=2)
    doc.add_paragraph("• Instructaj la angajare (Legea 319/2006 art. 20)", style="List Bullet")
    doc.add_paragraph("• Instructaj la locul de muncă (înainte de începere fiecare zi de lucru)", style="List Bullet")
    doc.add_paragraph("• Instructaj periodic (lunar / NSPM specifice gaze)", style="List Bullet")
    doc.add_paragraph("• Verificare cunoștințe (test scris bianual)", style="List Bullet")
    doc.add_paragraph()
    base._add_heading(doc, "4. Plan de intervenție urgență", level=2)
    base._add_kv_table(doc, [
        ("ISU (Pompieri)",          "112 / 021-411.96.96"),
        ("Ambulanță",                "112 / 961"),
        ("Distrigaz Sud avarii",     "021-9281 / dispecerat 24/7"),
        ("E-Distribuție avarii",     "021-9291"),
        ("Coordonator SSM proiect",  base._get(data, "ssm_coordonator", "—")),
        ("Responsabil executant",    base._get(data, "exec_responsabil_tehnic")),
    ])
    doc.add_paragraph()
    base._add_heading(doc, "5. Semnalizare șantier", level=2)
    doc.add_paragraph("• Panou de identificare șantier la intrare (vizibil de la 50m)", style="List Bullet")
    doc.add_paragraph("• Marcaje pentru aliniament traseu conductă în execuție", style="List Bullet")
    doc.add_paragraph("• Semnalizare circulație pietonală + auto (cf. SR 1848-7)", style="List Bullet")
    doc.add_paragraph("• Iluminare temporară zona de lucru (>50 lux)", style="List Bullet")
    base._footer_signature(doc, proj, data)
    return base._save(doc)


# === REGISTRY ADDONS ===
LEGAL_TEMPLATES = {
    "declaratie_conformitate":   {"label": "Declarație de Conformitate Executant",         "phase": "receptie", "fn": _declaratie_conformitate,   "norm": "Legea 10/1995 art. 25 + EN 1555"},
    "buletin_proba_rezistenta":  {"label": "Buletin Probă Rezistență Mecanică",            "phase": "executie", "fn": _buletin_proba_rezistenta,  "norm": "NTPEE 2018 cap. 5 art. 80"},
    "buletin_proba_etanseitate": {"label": "Buletin Probă Etanșeitate (24h)",              "phase": "executie", "fn": _buletin_proba_etanseitate, "norm": "NTPEE 2018 cap. 5 art. 82"},
    "pv_receptie_finala":        {"label": "PV Recepție Finală (PVRF)",                    "phase": "receptie", "fn": _pv_receptie_finala,        "norm": "HG 273/1994 art. 36"},
    "pv_pif_semnat":             {"label": "PV Punere în Funcțiune semnat OSD",            "phase": "pif",      "fn": _pv_pif_semnat,             "norm": "Ord. ANRE 162/2021"},
    "fisa_sudor":                {"label": "Fișă Sudor Autorizat",                          "phase": "executie", "fn": _fisa_sudor,                "norm": "ANRE Ord. 79/2014 + EN 13067"},
    "plan_ssm":                  {"label": "Plan SSM (Securitate Sănătate Muncă)",         "phase": "executie", "fn": _plan_ssm,                  "norm": "Legea 319/2006 + HG 1425/2006 + HG 300/2006"},
}


def register_into(base_module) -> None:
    if not hasattr(base_module, "TEMPLATES"):
        return
    for k, v in LEGAL_TEMPLATES.items():
        base_module.TEMPLATES.setdefault(k, v)
