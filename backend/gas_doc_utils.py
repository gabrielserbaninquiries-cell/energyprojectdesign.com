"""Shared utilities for gas DOCX builders.

Extracted from gas_doc_templates.py to break the circular import:
    gas_doc_templates ↔ gas_doc_proiect_complet

Both files now import only from this module (no cross-imports).
Backward-compatible: gas_doc_templates re-exports the symbols for legacy callers.
"""
from __future__ import annotations

import io
from datetime import datetime
from typing import Any, Dict, List, Tuple

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


# ============================================================================
# COMPANY HEADER — Energy Project Design SRL
# ============================================================================
COMPANY: Dict[str, str] = {
    "name": "ENERGY PROJECT DESIGN S.R.L.",
    "cui": "43151074",
    "reg_com": "J40/12982/2020",
    "address": "Str. Lt. Alexandru Popescu nr. 9B, Sectorul 3, București",
    "phone": "+40 723 000 000",
    "email": "office@energyprojectdesign.com",
    "web": "energyprojectdesign.com",
    "anre_proiectant": "PDD/2022/0001",
    "anre_executant": "EDD/2022/0001",
}


# ----------------------------------------------------------------------------
# Helpers (pure — no external module deps beyond python-docx)
# ----------------------------------------------------------------------------
def _today_ro() -> str:
    months = ["ianuarie", "februarie", "martie", "aprilie", "mai", "iunie",
              "iulie", "august", "septembrie", "octombrie", "noiembrie", "decembrie"]
    now = datetime.now()
    return f"{now.day:02d} {months[now.month-1]} {now.year}"


def _get(data: Dict[str, Any], key: str, default: str = "—") -> str:
    v = data.get(key)
    if v in (None, "", []):
        return default
    return str(v)


def _truthy(data: Dict[str, Any], key: str) -> bool:
    v = data.get(key)
    if v in (None, "", [], 0, "0"):
        return False
    return True


def _add_heading(doc: Document, text: str, level: int = 1,
                 align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    h = doc.add_heading(text, level=level)
    h.alignment = align


def _add_para(doc: Document, text: str, bold: bool = False, italic: bool = False,
              size: int = 11, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)


def _add_kv_table(doc: Document, rows: List[Tuple[str, str]],
                  col1_cm: float = 5.5) -> None:
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (k, v) in enumerate(rows):
        cells = t.rows[i].cells
        cells[0].text = k
        cells[1].text = v
        for run in cells[0].paragraphs[0].runs:
            run.bold = True


def _save(doc: Document) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


__all__ = [
    "COMPANY",
    "_today_ro",
    "_get",
    "_truthy",
    "_add_heading",
    "_add_para",
    "_add_kv_table",
    "_save",
]
