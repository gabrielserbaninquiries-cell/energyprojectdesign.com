"""System-seeded DOCX templates for each industry/subdomain.

These are built once on first server start and made available read-only to
every user via /api/templates (with is_system=true flag). Users can also
upload their own templates as before.

Templates are generated in-memory using python-docx with Romanian gas
engineering boilerplate text + placeholders. Replace these later with
your real legal forms.
"""
import io
import base64
from datetime import datetime, timezone
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _h(doc, text, size=18):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)


def _p(doc, text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11)
    if bold:
        r.bold = True


def _build_cerere_racordare() -> bytes:
    d = Document()
    _h(d, "CERERE DE RACORDARE LA REȚEAUA DE DISTRIBUȚIE GAZE NATURALE")
    d.add_paragraph()
    _p(d, "Către: {{osd}}")
    d.add_paragraph()
    _p(d, "Subsemnatul/Subscrisa {{beneficiar}}, cu domiciliul/sediul în localitatea {{localitate}}, județul {{judet}}, adresa {{adresa_lucrare}}, telefon {{telefon}}, email {{email}}, vă rog să aprobați racordarea la rețeaua de distribuție gaze naturale a imobilului situat la adresa menționată.")
    d.add_paragraph()
    _p(d, "Date tehnice ale lucrării:", bold=True)
    _p(d, "• Tip lucrare: {{tip_lucrare}}")
    _p(d, "• Debit instalat: {{debit_instalat}} mc/h")
    _p(d, "• Putere instalată: {{putere_instalata_kw}} kW")
    _p(d, "• Lungime estimată branșament: {{lungime_bransament}} m")
    _p(d, "• Presiune regim solicitată: {{presiune_regim}}")
    _p(d, "• Contor propus: {{contor_orientativ}}")
    d.add_paragraph()
    _p(d, "Date contract: nr. {{numar_contract}} din {{data_contract}}.")
    _p(d, "Proiectant: {{proiectant}}")
    _p(d, "Executant: {{executant}}")
    d.add_paragraph()
    _p(d, "Observații: {{observatii}}")
    d.add_paragraph()
    _p(d, "Data: {{data_document}}")
    _p(d, "Semnătura beneficiar: ________________________")
    out = io.BytesIO()
    d.save(out)
    return out.getvalue()


def _build_memoriu_tehnic() -> bytes:
    d = Document()
    _h(d, "MEMORIU TEHNIC — INSTALAȚIE GAZE NATURALE", size=16)
    d.add_paragraph()
    _p(d, "1. DATE GENERALE", bold=True)
    _p(d, "Beneficiar: {{beneficiar}}")
    _p(d, "Adresa lucrării: {{adresa_lucrare}}, {{localitate}}, jud. {{judet}}")
    _p(d, "Telefon contact: {{telefon}} | Email: {{email}}")
    _p(d, "Operator distribuție (OSD): {{osd}}")
    _p(d, "Tip lucrare: {{tip_lucrare}}")
    _p(d, "Contract: nr. {{numar_contract}} / {{data_contract}}")
    d.add_paragraph()
    _p(d, "2. ECHIPA TEHNICĂ AUTORIZATĂ", bold=True)
    _p(d, "Proiectant ANRE: {{proiectant}}")
    _p(d, "Executant ANRE: {{executant}}")
    _p(d, "Verificator documentație (VGD): {{verificator_vgd}}")
    _p(d, "Responsabil tehnic execuție (RTE): {{responsabil_rte}}")
    d.add_paragraph()
    _p(d, "3. DATE TEHNICE", bold=True)
    _p(d, "Debit instalat: {{debit_instalat}} mc/h")
    _p(d, "Debit calculat: {{debit_calculat_mc_h}} mc/h")
    _p(d, "Debit recomandat (cu marjă 10%): {{debit_recomandat_mc_h}} mc/h")
    _p(d, "Putere instalată estimată: {{putere_instalata_kw}} kW")
    _p(d, "Presiune regim: {{presiune_regim}}")
    _p(d, "Diametru conductă: {{diametru_conducta}}")
    _p(d, "Material conductă: {{material_conducta}}")
    _p(d, "Lungime branșament: {{lungime_bransament}} m")
    _p(d, "Punct racordare: {{punct_racordare}}")
    _p(d, "Post reglare: {{post_reglare}}")
    _p(d, "Contor recomandat: {{contor_orientativ}}")
    _p(d, "Risc presiune: {{risc_presiune}}")
    _p(d, "Estimare cost branșament: {{estimare_cost}} RON")
    d.add_paragraph()
    _p(d, "4. OBSERVAȚII TEHNICE", bold=True)
    _p(d, "{{observatii}}")
    d.add_paragraph()
    _p(d, "Data emiterii: {{data_document}}")
    _p(d, "Întocmit: {{proiectant}}")
    out = io.BytesIO()
    d.save(out)
    return out.getvalue()


def _build_borderou() -> bytes:
    d = Document()
    _h(d, "BORDEROU DOCUMENTE — DOSAR TEHNIC")
    d.add_paragraph()
    _p(d, "Beneficiar: {{beneficiar}}")
    _p(d, "Adresa lucrare: {{adresa_lucrare}}, {{localitate}}, {{judet}}")
    _p(d, "Tip lucrare: {{tip_lucrare}}")
    _p(d, "Contract: {{numar_contract}} / {{data_contract}}")
    d.add_paragraph()
    _p(d, "Conținut dosar:", bold=True)
    _p(d, "1. Cerere de racordare")
    _p(d, "2. Memoriu tehnic")
    _p(d, "3. Plan de încadrare (anexă)")
    _p(d, "4. Plan de situație (anexă)")
    _p(d, "5. Schemă izometrică (anexă)")
    _p(d, "6. Acord acces / declarație proprietate (anexă)")
    _p(d, "7. Verificare VGD: {{verificator_vgd}}")
    _p(d, "8. Confirmare RTE: {{responsabil_rte}}")
    d.add_paragraph()
    _p(d, "Întocmit: {{proiectant}} | Data: {{data_document}}")
    out = io.BytesIO()
    d.save(out)
    return out.getvalue()


def _build_adresa_osd() -> bytes:
    d = Document()
    _h(d, "ADRESĂ CĂTRE OPERATORUL SISTEMULUI DE DISTRIBUȚIE", size=14)
    d.add_paragraph()
    _p(d, "Către: {{osd}}")
    _p(d, "În atenția: Departamentul Racordări")
    d.add_paragraph()
    _p(d, "Stimată/Stimate domnule director,")
    d.add_paragraph()
    _p(d, "Prin prezenta, vă transmitem documentația tehnică aferentă lucrării de {{tip_lucrare}} pentru beneficiarul {{beneficiar}}, situat la adresa {{adresa_lucrare}}, {{localitate}}, județul {{judet}}.")
    d.add_paragraph()
    _p(d, "Detalii lucrare:", bold=True)
    _p(d, "• Contract nr. {{numar_contract}} din {{data_contract}}")
    _p(d, "• Debit instalat: {{debit_instalat}} mc/h")
    _p(d, "• Lungime branșament: {{lungime_bransament}} m")
    _p(d, "• Proiectant ANRE: {{proiectant}}")
    _p(d, "• Executant ANRE: {{executant}}")
    d.add_paragraph()
    _p(d, "Vă rugăm să analizați documentația și să comunicați avizul dvs. în termenul legal.")
    d.add_paragraph()
    _p(d, "Cu deosebită stimă,")
    _p(d, "{{proiectant}}")
    _p(d, "Data: {{data_document}}")
    out = io.BytesIO()
    d.save(out)
    return out.getvalue()


def _build_certificare_vgd() -> bytes:
    d = Document()
    _h(d, "CERTIFICARE INTERNĂ — VERIFICATOR DOCUMENTAȚIE (VGD)", size=14)
    d.add_paragraph()
    _p(d, "Beneficiar lucrare: {{beneficiar}}")
    _p(d, "Adresa lucrării: {{adresa_lucrare}}, {{localitate}}, {{judet}}")
    _p(d, "Tip lucrare: {{tip_lucrare}}")
    _p(d, "Contract: nr. {{numar_contract}} / {{data_contract}}")
    d.add_paragraph()
    _p(d, "1. VERIFICATOR DOCUMENTAȚIE", bold=True)
    _p(d, "Nume: {{verificator_vgd}}")
    _p(d, "Atestat ANRE: {{atestat_vgd}}")
    _p(d, "Data verificării: {{data_verificare_vgd}}")
    _p(d, "Status verificare: {{status_vgd}}")
    d.add_paragraph()
    _p(d, "2. OBSERVAȚII VERIFICARE", bold=True)
    _p(d, "{{observatii_vgd}}")
    d.add_paragraph()
    _p(d, "Subsemnatul, în calitate de Verificator Documentație autorizat ANRE, certific verificarea completă a documentației tehnice aferente lucrării de mai sus, conform reglementărilor în vigoare.")
    d.add_paragraph()
    _p(d, "Data certificării: {{data_document}}")
    _p(d, "Semnătură VGD: ________________________")
    _p(d, "Ștampilă: ________________________")
    out = io.BytesIO()
    d.save(out)
    return out.getvalue()


def _build_certificare_rte() -> bytes:
    d = Document()
    _h(d, "CERTIFICARE INTERNĂ — RESPONSABIL TEHNIC EXECUȚIE (RTE)", size=14)
    d.add_paragraph()
    _p(d, "Beneficiar lucrare: {{beneficiar}}")
    _p(d, "Adresa lucrării: {{adresa_lucrare}}, {{localitate}}, {{judet}}")
    _p(d, "Tip lucrare: {{tip_lucrare}}")
    _p(d, "Executant: {{executant}}")
    _p(d, "Contract: nr. {{numar_contract}} / {{data_contract}}")
    d.add_paragraph()
    _p(d, "1. RESPONSABIL TEHNIC EXECUȚIE", bold=True)
    _p(d, "Nume: {{responsabil_rte}}")
    _p(d, "Autorizație ANRE: {{autorizatie_rte}}")
    _p(d, "Data verificării execuției: {{data_verificare_rte}}")
    _p(d, "Status execuție: {{status_rte}}")
    d.add_paragraph()
    _p(d, "2. OBSERVAȚII EXECUȚIE", bold=True)
    _p(d, "{{observatii_rte}}")
    d.add_paragraph()
    _p(d, "Subsemnatul, în calitate de Responsabil Tehnic cu Execuția autorizat ANRE, certific execuția conformă a lucrării de mai sus și respectarea documentației tehnice verificate.")
    d.add_paragraph()
    _p(d, "Data certificării: {{data_document}}")
    _p(d, "Semnătură RTE: ________________________")
    _p(d, "Ștampilă: ________________________")
    out = io.BytesIO()
    d.save(out)
    return out.getvalue()


def _build_cerere_racordare_fv() -> bytes:
    d = Document()
    _h(d, "CERERE DE RACORDARE LA REȚEAUA ELECTRICĂ DE DISTRIBUȚIE — INSTALAȚIE FOTOVOLTAICĂ", size=14)
    d.add_paragraph()
    _p(d, "Către: {{operator_distributie}}")
    _p(d, "În atenția: Departamentul Racordări Prosumatori / Producători")
    d.add_paragraph()
    _p(d, "Subsemnatul/Subscrisa {{beneficiar}}, cu domiciliul/sediul în {{localitate}}, jud. {{judet}}, "
          "adresa {{adresa_lucrare}}, telefon {{telefon}}, email {{email}}, "
          "solicit racordarea la rețeaua electrică de distribuție a unei instalații fotovoltaice "
          "cu putere instalată de {{fv_p_kwp}} kWp, încadrată în categoria ANRE {{fv_categorie_anre}} "
          "({{fv_categorie_label}}).")
    d.add_paragraph()
    _p(d, "1. DATE TEHNICE INSTALAȚIE FV", bold=True)
    _p(d, "• Putere instalată: {{fv_p_kwp}} kWp")
    _p(d, "• Număr panouri: {{fv_n_panouri}} × {{fv_p_panou_wp}} Wp")
    _p(d, "• Configurație: {{fv_n_string}} string-uri × {{fv_n_serie}} panouri/string")
    _p(d, "• Tensiune DC string (STC): {{fv_u_string_v}} V")
    _p(d, "• Invertor recomandat: {{fv_invertor_kw}} kW (raport DC/AC = {{fv_raport_dc_ac}})")
    _p(d, "• Producție anuală estimată: {{fv_productie_anuala_kwh}} kWh "
          "({{fv_productie_specifica}} kWh/kWp · iradiație {{fv_iradiatie_kwh_m2_an}} kWh/m²/an)")
    d.add_paragraph()
    _p(d, "2. REGIM DE FUNCȚIONARE SOLICITAT", bold=True)
    _p(d, "Regim: {{fv_regim}}")
    _p(d, "Tip racord: {{fv_tip_racord}}")
    _p(d, "Aviz necesar (conform ANRE Ord. 34/2024): {{fv_aviz_necesar}}")
    d.add_paragraph()
    _p(d, "{IF fv_p_kwp <= 10.8: Solicit înregistrarea ca prosumator casnic, cu compensare cantitativă a energiei livrate, conform Legii 220/2008 și Ord. ANRE 15/2022 actualizat. ELSE Solicit emiterea ATR conform procedurii standard pentru categoria ANRE corespunzătoare puterii instalate.}")
    d.add_paragraph()
    _p(d, "{IF fv_p_kwp > 200: ATENȚIE — instalație de tip parc fotovoltaic. Anexez studiu de soluție, acord mediu și solicitare licență producere energie electrică ANRE. ELSE Anexez documentația tehnică standard conform Ord. 34/2024.}")
    d.add_paragraph()
    _p(d, "3. PROIECTANT / EXECUTANT", bold=True)
    _p(d, "Proiectant atestat ANRE: {{proiectant}} ({{atestat_proiectant}})")
    _p(d, "Executant atestat ANRE: {{executant}} ({{atestat_executant}})")
    d.add_paragraph()
    _p(d, "Contract intern: nr. {{numar_contract}} / {{data_contract}}")
    _p(d, "Observații: {{observatii}}")
    d.add_paragraph()
    _p(d, "Data: {{data_document}}")
    _p(d, "Semnătură beneficiar: ________________________")
    _p(d, "Ștampilă proiectant: ________________________")
    out = io.BytesIO()
    d.save(out)
    return out.getvalue()


def _build_memoriu_tehnic_fv() -> bytes:
    d = Document()
    _h(d, "MEMORIU TEHNIC — INSTALAȚIE FOTOVOLTAICĂ", size=16)
    d.add_paragraph()
    _p(d, "1. DATE GENERALE", bold=True)
    _p(d, "Beneficiar: {{beneficiar}}")
    _p(d, "Adresa lucrării: {{adresa_lucrare}}, {{localitate}}, jud. {{judet}}")
    _p(d, "Telefon / Email: {{telefon}} / {{email}}")
    _p(d, "Operator distribuție: {{operator_distributie}}")
    _p(d, "Contract: nr. {{numar_contract}} / {{data_contract}}")
    d.add_paragraph()
    _p(d, "2. ÎNCADRARE ANRE", bold=True)
    _p(d, "Putere instalată: {{fv_p_kwp}} kWp")
    _p(d, "Categorie: {{fv_categorie_anre}} — {{fv_categorie_label}}")
    _p(d, "Regim funcționare: {{fv_regim}}")
    _p(d, "Tip racord: {{fv_tip_racord}}")
    _p(d, "Aviz/avize necesare: {{fv_aviz_necesar}}")
    d.add_paragraph()
    _p(d, "{IF fv_categorie_anre == C1: Instalație de tip PROSUMATOR CASNIC. Procedură simplificată conform ANRE Ord. 34/2024 — ATR + CR emise în max. 30 de zile. Beneficiarul are dreptul la compensare cantitativă a energiei livrate în rețea, conform Legii 220/2008 cu modificările ulterioare. ELSE Aplicabile prevederi extinse ANRE — vezi paragrafele de mai jos.}")
    d.add_paragraph()
    _p(d, "{IF fv_categorie_anre == C2: Instalație de tip PROSUMATOR NON-CASNIC. ATR cu studiu de soluție simplificat. Racord trifazat 3x400V obligatoriu. Compensare cantitativă disponibilă pentru excedent. ELSE  }")
    _p(d, "{IF fv_categorie_anre == C3: Instalație de tip PRODUCĂTOR MIC (27-200 kWp). Studiu de soluție complet, ATR cu cerințe extinse. Pentru P > 100 kWp este obligatorie obținerea licenței de producere energie electrică de la ANRE. Releu de protecție conform EN 50549-1 obligatoriu. ELSE  }")
    _p(d, "{IF fv_categorie_anre == C4: Instalație de tip PARC FOTOVOLTAIC (> 200 kWp). Racord pe MT 20 kV. Necesită: studiu de soluție complet, acord de mediu, licență producere ANRE, celulă MT 24 kV, telegestiune cu protocol Modbus/IEC 61850. ELSE  }")
    d.add_paragraph()
    _p(d, "3. CONFIGURAȚIE TEHNICĂ", bold=True)
    _p(d, "Panouri fotovoltaice: {{fv_n_panouri}} buc × {{fv_p_panou_wp}} Wp "
          "(module monocristaline TOPCon/HJT, eficiență ≥ 21.5%)")
    _p(d, "Configurație string-uri: {{fv_n_string}} string-uri × {{fv_n_serie}} module în serie")
    _p(d, "Tensiune string STC (25°C): {{fv_u_string_v}} V")
    _p(d, "Invertor: {{fv_invertor_kw}} kW (raport DC/AC = {{fv_raport_dc_ac}})")
    _p(d, "{IF fv_p_kwp <= 5: Conectare monofazată 230V admisă (P ≤ 5 kWp). ELSE Conectare trifazată 3x400V obligatorie.}")
    d.add_paragraph()
    _p(d, "4. CABLURI ȘI PROTECȚII", bold=True)
    _p(d, "Cablu DC (string → invertor): {{fv_cablu_dc_mm2}} mm² — tip H1Z2Z2-K solar 1500V "
          "(cădere de tensiune calculată: {{fv_cablu_dc_caderea_pct}}%, sub limita admisă 1%)")
    _p(d, "Cablu AC (invertor → tablou racordare): {{fv_cablu_ac_mm2}} mm² — tip CYY-F/N2XH 0.6/1 kV "
          "(curent calculat: {{fv_cablu_ac_curent_a}} A, cădere admisă 1.5% conf. NTI-TEL-007)")
    d.add_paragraph()
    _p(d, "Protecții electrice obligatorii (conform I7-2011 și SR EN IEC 62548):")
    _p(d, "{{fv_protectii_lista}}")
    d.add_paragraph()
    _p(d, "5. PRODUCȚIE ESTIMATĂ", bold=True)
    _p(d, "Iradiație zonă: {{fv_iradiatie_kwh_m2_an}} kWh/m²/an (sursa: PVGIS-SARAH3, JRC)")
    _p(d, "Performance Ratio (PR): 0.78 (mediu, include pierderi cabluri/invertor/mismatch/T°/murdărire)")
    _p(d, "Producție anuală estimată: {{fv_productie_anuala_kwh}} kWh")
    _p(d, "Producție specifică: {{fv_productie_specifica}} kWh/kWp")
    _p(d, "Factor de utilizare anual: {{fv_factor_utilizare_pct}}%")
    d.add_paragraph()
    _p(d, "{IF fv_p_kwp >= 100: ATENȚIE: instalație ≥ 100 kWp — sunt obligatorii smart-meter bidirecțional clasa 1, telegestiune Modbus/M-Bus și relee de protecție EN 50549-1 (anti-islanding, U/f, ROCOF). ELSE Echipamentele standard de protecție specificate la pct. 4 sunt suficiente.}")
    d.add_paragraph()
    _p(d, "6. ECHIPA TEHNICĂ AUTORIZATĂ", bold=True)
    _p(d, "Proiectant ANRE: {{proiectant}} ({{atestat_proiectant}})")
    _p(d, "Executant ANRE: {{executant}} ({{atestat_executant}})")
    _p(d, "Verificator documentație: {{verificator_vgd}}")
    _p(d, "Responsabil tehnic execuție (RTE): {{responsabil_rte}}")
    d.add_paragraph()
    _p(d, "7. OBSERVAȚII", bold=True)
    _p(d, "{{observatii}}")
    d.add_paragraph()
    _p(d, "Data emiterii: {{data_document}}")
    _p(d, "Întocmit: {{proiectant}}")
    _p(d, "Semnătură și ștampilă: ________________________")
    out = io.BytesIO()
    d.save(out)
    return out.getvalue()


def _build_adresa_od_fv() -> bytes:
    d = Document()
    _h(d, "ADRESĂ CĂTRE OPERATORUL DE DISTRIBUȚIE — INSTALAȚIE FOTOVOLTAICĂ", size=13)
    d.add_paragraph()
    _p(d, "Către: {{operator_distributie}}")
    _p(d, "În atenția: Departamentul Racordări Prosumatori / Producători")
    d.add_paragraph()
    _p(d, "Stimată/Stimate domnule director,")
    d.add_paragraph()
    _p(d, "Prin prezenta, depunem documentația tehnică completă aferentă instalației fotovoltaice "
          "ce urmează a fi racordată pentru beneficiarul {{beneficiar}}, situat la adresa "
          "{{adresa_lucrare}}, {{localitate}}, jud. {{judet}}.")
    d.add_paragraph()
    _p(d, "Date sintetice:", bold=True)
    _p(d, "• Putere instalată: {{fv_p_kwp}} kWp — categorie ANRE {{fv_categorie_anre}}")
    _p(d, "• Regim funcționare: {{fv_regim}}")
    _p(d, "• Producție anuală estimată: {{fv_productie_anuala_kwh}} kWh")
    _p(d, "• Invertor recomandat: {{fv_invertor_kw}} kW")
    _p(d, "• Contract: nr. {{numar_contract}} / {{data_contract}}")
    d.add_paragraph()
    _p(d, "{IF fv_p_kwp <= 10.8: Solicit emiterea simplificată a ATR și CR (procedură prosumator casnic ≤ 10.8 kWp), conform ANRE Ord. 34/2024 — termen legal 30 de zile. ELSE Solicit deschiderea procedurii standard de racordare cu studiu de soluție conform categoriei ANRE indicate.}")
    d.add_paragraph()
    _p(d, "{IF fv_p_kwp > 100: Anexez în mod suplimentar: studiu de soluție complet, schema unifilară aprobată VGD, fișa tehnică invertor și panouri, certificate de conformitate EN 50549-1, dovada inițiere licență producere ANRE. ELSE Anexez documentația standard: cerere racordare, memoriu tehnic, schema unifilară, fișe tehnice echipamente, declarație proprietate.}")
    d.add_paragraph()
    _p(d, "Echipa autorizată:")
    _p(d, "• Proiectant ANRE: {{proiectant}}")
    _p(d, "• Executant ANRE: {{executant}}")
    _p(d, "• Verificator documentație: {{verificator_vgd}}")
    d.add_paragraph()
    _p(d, "Vă mulțumim pentru disponibilitate și rămânem deschiși oricăror clarificări tehnice.")
    d.add_paragraph()
    _p(d, "Cu stimă,")
    _p(d, "{{proiectant}}")
    _p(d, "Data: {{data_document}}")
    out = io.BytesIO()
    d.save(out)
    return out.getvalue()


SYSTEM_TEMPLATES = [
    {
        "key": "sys_cerere_racordare_gaz",
        "name": "Cerere racordare gaze naturale (sistem)",
        "industry": "gas_engineering",
        "subdomain": "bransamente_gaz",
        "builder": _build_cerere_racordare,
    },
    {
        "key": "sys_memoriu_tehnic_gaz",
        "name": "Memoriu tehnic instalație gaze (sistem)",
        "industry": "gas_engineering",
        "subdomain": "instalatii_utilizare",
        "builder": _build_memoriu_tehnic,
    },
    {
        "key": "sys_borderou_documente_gaz",
        "name": "Borderou documente dosar gaze (sistem)",
        "industry": "gas_engineering",
        "subdomain": "bransamente_gaz",
        "builder": _build_borderou,
    },
    {
        "key": "sys_adresa_osd_gaz",
        "name": "Adresă către OSD (sistem)",
        "industry": "gas_engineering",
        "subdomain": "bransamente_gaz",
        "builder": _build_adresa_osd,
    },
    {
        "key": "sys_certificare_vgd_gaz",
        "name": "Certificare internă VGD (sistem)",
        "industry": "gas_engineering",
        "subdomain": "bransamente_gaz",
        "builder": _build_certificare_vgd,
    },
    {
        "key": "sys_certificare_rte_gaz",
        "name": "Certificare internă RTE (sistem)",
        "industry": "gas_engineering",
        "subdomain": "bransamente_gaz",
        "builder": _build_certificare_rte,
    },
    {
        "key": "sys_cerere_racordare_fv",
        "name": "Cerere racordare fotovoltaic (sistem)",
        "industry": "photovoltaics",
        "subdomain": "racordare_fv",
        "builder": _build_cerere_racordare_fv,
    },
    {
        "key": "sys_memoriu_tehnic_fv",
        "name": "Memoriu tehnic fotovoltaic — smart IF/ELSE (sistem)",
        "industry": "photovoltaics",
        "subdomain": "racordare_fv",
        "builder": _build_memoriu_tehnic_fv,
    },
    {
        "key": "sys_adresa_od_fv",
        "name": "Adresă către OD pentru racordare FV (sistem)",
        "industry": "photovoltaics",
        "subdomain": "racordare_fv",
        "builder": _build_adresa_od_fv,
    },
]


async def seed_system_templates(db):
    """Ensure all system templates exist (idempotent)."""
    now = datetime.now(timezone.utc).isoformat()
    from docx_processor import extract_placeholders
    for tpl in SYSTEM_TEMPLATES:
        existing = await db.system_templates.find_one({"key": tpl["key"]})
        data = tpl["builder"]()
        placeholders = extract_placeholders(data)
        doc = {
            "key": tpl["key"],
            "template_id": tpl["key"],
            "name": tpl["name"],
            "industry": tpl["industry"],
            "subdomain": tpl["subdomain"],
            "placeholders": placeholders,
            "size_bytes": len(data),
            "data_b64": base64.b64encode(data).decode("ascii"),
            "is_system": True,
            "created_at": existing["created_at"] if existing else now,
            "updated_at": now,
        }
        if existing:
            await db.system_templates.update_one({"key": tpl["key"]}, {"$set": doc})
        else:
            await db.system_templates.insert_one(doc)
