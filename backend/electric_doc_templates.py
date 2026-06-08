"""Electric (LES/LEA) — Template-uri DOCX concrete pentru distribuție electrică joasă/medie tensiune.

Replică pattern-ul `gas_doc_templates.py` adaptat pentru:
- Branșament electric LES (subteran) / LEA (aerian)
- Norme: Ord. ANRE 11/2014 + PE 101/85 + PE 107/95 + STAS 8779
- Operatori: E-Distribuție (Muntenia/Dobrogea/Banat), Delgaz Grid, DEER

6 template-uri concrete + 4 stub-uri pentru extindere ulterioară.
"""
from __future__ import annotations
import io
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

import gas_doc_templates as base  # reuse helpers

COMPANY = base.COMPANY  # same EPD company header


def _hdr(doc: Document, title: str, subtitle: str = "") -> None:
    base._company_header(doc)
    base._add_heading(doc, title, level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    if subtitle:
        base._add_para(doc, subtitle, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
    doc.add_paragraph()


def cerere_atr_electric(proj: Dict[str, Any]) -> bytes:
    data = proj.get("data") or {}
    doc = Document()
    _hdr(doc, "CERERE AVIZ TEHNIC RACORDARE (ELECTRIC)",
         "(Conform Ord. ANRE 11/2014 + Regulamentul ANRE de racordare)")
    base._add_para(doc, f"Către {base._get(data,'electrica_operator','OPERATOR DE DISTRIBUȚIE')}", bold=True)
    doc.add_paragraph()
    base._add_kv_table(doc, [
        ("Beneficiar", base._get(data, "beneficiar_nume")),
        ("CNP/CUI", base._get(data, "beneficiar_cnp_cui")),
        ("Adresă loc consum", base._get(data, "loc_consum_adresa")),
        ("Nr. cadastral", base._get(data, "loc_consum_cadastru")),
        ("Tensiune solicitată", base._get(data, "el_tensiune", "0.4 kV (Joasă tensiune)")),
        ("Putere instalată (kW)", base._get(data, "el_putere_instalata_kw")),
        ("Putere absorbită (kW)", base._get(data, "el_putere_absorbita_kw")),
        ("Coeficient simultaneitate", base._get(data, "el_coef_simultaneitate", "0.85")),
        ("Tip racordare", base._get(data, "el_tip_racordare", "Monofazat 230V")),
        ("Tip consumator", base._get(data, "tip_consumator")),
    ])
    doc.add_paragraph()
    base._add_para(doc, "Anexe:", bold=True)
    for ln in [
        "• Copie act proprietate / contract", "• Extras CF",
        "• Plan de încadrare 1:5000 + plan situație 1:500",
        "• Schema electrică principală + listă consumatori",
        "• Memoriu tehnic instalație electrică",
    ]:
        doc.add_paragraph(ln, style="List Bullet")
    base._footer_signature(doc, proj, data)
    return base._save(doc)


def memoriu_electric(proj: Dict[str, Any]) -> bytes:
    data = proj.get("data") or {}
    doc = Document()
    _hdr(doc, "MEMORIU TEHNIC JUSTIFICATIV (INSTALAȚIE ELECTRICĂ)",
         "(Conform PE 101/85 + PE 107/95 + Ord. ANRE 11/2014)")
    base._add_kv_table(doc, [
        ("Beneficiar", base._get(data, "beneficiar_nume")),
        ("Adresă", base._get(data, "loc_consum_adresa")),
        ("Tip lucrare", base._get(data, "tipul_lucrarii", "Branșament electric")),
        ("Tip cablu", base._get(data, "el_tip_cablu", "ACYABY 4x16 mm²")),
        ("Secțiune cablu (mm²)", base._get(data, "el_sectiune_cablu_mm2")),
        ("Lungime cablu (m)", base._get(data, "el_lungime_cablu_m")),
        ("Tip pozare", base._get(data, "el_tip_pozare", "Subteran LES / Aerian LEA")),
        ("Cădere tensiune max. (%)", base._get(data, "el_cadere_tensiune_pct", "3%")),
    ])
    doc.add_paragraph()
    base._add_para(doc, "Calcule electrice:", bold=True)
    base._add_para(doc, "I = P / (U × cos φ) = ... A · ΔU = I × R × L / 1000 < 3% × Un",
                   italic=True)
    doc.add_paragraph()
    base._add_para(doc, "Soluție propusă:", bold=True)
    base._add_para(doc, base._get(data, "sf_solutie_tehnica",
                   "Racordare la rețeaua de distribuție prin firidă tip E1 + contor electronic."),
                   italic=True)
    base._footer_signature(doc, proj, data)
    return base._save(doc)


def caiet_sarcini_electric(proj: Dict[str, Any]) -> bytes:
    data = proj.get("data") or {}
    doc = Document()
    _hdr(doc, "CAIET DE SARCINI EXECUȚIE INSTALAȚIE ELECTRICĂ",
         "(Conform PE 107/95 + I 7/2011)")
    base._add_kv_table(doc, [
        ("Executant (autorizat ANRE)", base._get(data, "exec_firma")),
        ("RTE", base._get(data, "exec_responsabil_tehnic")),
        ("Diriginte", base._get(data, "exec_diriginte_santier")),
        ("Tip cablu", base._get(data, "el_tip_cablu")),
        ("Adâncime pozare LES (cm)", base._get(data, "el_adancime_les_cm", "80")),
    ])
    doc.add_paragraph()
    base._add_para(doc, "Materiale principale:", bold=True)
    for ln in [
        "• Cablu ACYABY / NYY conform STAS 8779",
        "• Tub PVC protecție / țeavă metalică",
        "• Firida electrică tip E1 / E2 / E3 conform proiect",
        "• Contor electronic monofazat / trifazat",
        "• Disjunctoare diferențiale 30 mA + automate",
    ]:
        doc.add_paragraph(ln, style="List Bullet")
    base._footer_signature(doc, proj, data)
    return base._save(doc)


def cerere_pif_electric(proj: Dict[str, Any]) -> bytes:
    data = proj.get("data") or {}
    doc = Document()
    _hdr(doc, "CERERE PUNERE ÎN FUNCȚIUNE INSTALAȚIE ELECTRICĂ",
         "(Către operatorul de distribuție electrică)")
    base._add_para(doc, f"Către {base._get(data,'electrica_operator','OPERATOR DE DISTRIBUȚIE')}", bold=True)
    doc.add_paragraph()
    base._add_kv_table(doc, [
        ("Obiectiv", proj.get("title", "—")),
        ("Beneficiar", base._get(data, "beneficiar_nume")),
        ("Adresă", base._get(data, "loc_consum_adresa")),
        ("ATR nr.", f"{base._get(data,'atr_numar','—')} / {base._get(data,'atr_data','—')}"),
        ("AC nr.", f"{base._get(data,'ac_numar','—')} / {base._get(data,'ac_data_emitere','—')}"),
        ("PV recepție nr.", f"{base._get(data,'receptie_pv_numar','—')} / {base._get(data,'receptie_pv_data','—')}"),
        ("Serie contor instalat", base._get(data, "pif_contor_serie")),
        ("Index inițial (kWh)", base._get(data, "pif_contor_index_initial")),
    ])
    doc.add_paragraph()
    base._add_para(doc, "Solicităm efectuarea PIF și încheierea contractului de furnizare.", italic=True)
    base._footer_signature(doc, proj, data)
    return base._save(doc)


def carte_tehnica_electric(proj: Dict[str, Any]) -> bytes:
    data = proj.get("data") or {}
    doc = Document()
    _hdr(doc, "CARTEA TEHNICĂ A INSTALAȚIEI ELECTRICE",
         "(Conform HG 273/1994 + Ord. MLPAT 770/1997 + Legea 10/1995)")
    for sec, key in [
        ("A. Documentația de proiectare", "ct_sectiune_A_continut"),
        ("B. Documentația de execuție", "ct_sectiune_B_continut"),
        ("C. Documentația de recepție", "ct_sectiune_C_continut"),
        ("D. Documentația de exploatare", "ct_sectiune_D_continut"),
    ]:
        base._add_heading(doc, sec, level=2)
        base._add_para(doc, base._get(data, key, "(de completat)"), italic=True)
        doc.add_paragraph()
    base._footer_signature(doc, proj, data)
    return base._save(doc)


def pv_receptie_electric(proj: Dict[str, Any]) -> bytes:
    data = proj.get("data") or {}
    doc = Document()
    _hdr(doc, "PROCES-VERBAL DE RECEPȚIE INSTALAȚIE ELECTRICĂ",
         "(Conform HG 273/1994 + Legea 10/1995)")
    base._add_kv_table(doc, [
        ("Nr. PV", base._get(data, "receptie_pv_numar")),
        ("Data", base._get(data, "receptie_pv_data")),
        ("Beneficiar", base._get(data, "beneficiar_nume")),
        ("Executant", base._get(data, "exec_firma")),
        ("Cădere tensiune măsurată (%)", base._get(data, "el_cadere_tensiune_masurata", "<3%")),
        ("Rezistență izolație (MΩ)", base._get(data, "el_rezistenta_izolatie", ">0.5")),
        ("Rezistență priză pământ (Ω)", base._get(data, "el_priza_pamant", "<4")),
        ("Rezultat verificări", base._get(data, "proba_rezultat", "Admis")),
    ])
    base._footer_signature(doc, proj, data)
    return base._save(doc)


TEMPLATES = {
    "el_cerere_atr":       {"label": "Cerere ATR Electric", "phase": "atr",      "fn": cerere_atr_electric,  "norm": "Ord. ANRE 11/2014"},
    "el_memoriu_tehnic":   {"label": "Memoriu Tehnic Electric", "phase": "dtac",  "fn": memoriu_electric,     "norm": "PE 101/85 + PE 107/95"},
    "el_caiet_sarcini":    {"label": "Caiet Sarcini Electric", "phase": "pt",     "fn": caiet_sarcini_electric,"norm": "I 7/2011"},
    "el_cerere_pif":       {"label": "Cerere PIF Electric", "phase": "pif",       "fn": cerere_pif_electric,  "norm": "Ord. ANRE 11/2014"},
    "el_pv_receptie":      {"label": "PV Recepție Electric", "phase": "receptie", "fn": pv_receptie_electric, "norm": "HG 273/1994"},
    "el_carte_tehnica":    {"label": "Carte Tehnică Electric", "phase": "receptie","fn": carte_tehnica_electric,"norm": "HG 273/1994 + Ord. MLPAT 770/1997"},
}


def list_templates() -> List[Dict[str, Any]]:
    return [{"id": k, **{kk: vv for kk, vv in v.items() if kk != "fn"}} for k, v in TEMPLATES.items()]


def generate(template_id: str, proj: Dict[str, Any]) -> Optional[Tuple[bytes, str]]:
    tpl = TEMPLATES.get(template_id)
    if not tpl:
        return None
    fname = f"{template_id}_{(proj.get('title','proiect')[:40]).replace(' ','_')}.docx"
    return tpl["fn"](proj), fname
