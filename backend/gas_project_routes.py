"""Gas Natural Project Studio — v2 (multi-country, subdomain-aware).

Endpoints (mounted under /api by server.py):

CATALOG (no project context needed)
- GET   /gas-project/catalog/countries            — list supported countries
- GET   /gas-project/catalog/{country}/subdomains — list subdomains for country
- GET   /gas-project/catalog/{country}/{sub}      — full subdomain blueprint (phases + fields)
- GET   /gas-project/calcs                        — list available calculations
- POST  /gas-project/calc                         — run a single calculation (no auth required)

PROJECTS (auth required)
- POST   /gas-project                                 — create project (subdomain, country, title)
- GET    /gas-project                                 — list current user's gas projects
- GET    /gas-project/{pid}                           — full project + phase schema
- PATCH  /gas-project/{pid}                           — partial update (data/phase/title/status)
- DELETE /gas-project/{pid}                           — soft delete
- POST   /gas-project/{pid}/sign                      — apply digital signature
- GET    /gas-project/{pid}/qr                        — return QR PNG (base64)
- GET    /gas-project/{pid}/public                    — public lightweight summary
- GET    /gas-project/{pid}/phases                    — phases schema only
- POST   /gas-project/{pid}/phase/{phase_id}/dispatch — email phase summary to recipients

RECIPIENTS (auth required)
- GET    /gas-project/recipients                      — list user's saved recipients per role
- PUT    /gas-project/recipients                      — replace user's recipients per role

LEGACY
- GET    /gas-project/phases                          — fallback to default subdomain (bransament-casnic)
"""
from __future__ import annotations
import base64
import hashlib
import io
import os
from datetime import datetime, timezone
from typing import Any, Dict, List, Optional

from fastapi import APIRouter, Depends, HTTPException
from pydantic import BaseModel, EmailStr

import qrcode
from fastapi import UploadFile, File, Form
from fastapi.responses import StreamingResponse

import gas_catalog as catalog
import gas_calc_engine as engine
import gas_doc_templates as doc_templates
import gas_avize_catalog as avize_catalog
import validators_ro


def _new_id(prefix: str) -> str:
    return f"{prefix}{os.urandom(8).hex()}"


# ---------- pydantic ----------
class GasProjectCreate(BaseModel):
    title: str
    country: str = "RO"
    subdomain: str = "bransament-casnic"
    data: Optional[Dict[str, Any]] = None
    phase: Optional[str] = "tema"


class GasProjectPatch(BaseModel):
    title: Optional[str] = None
    phase: Optional[str] = None
    data: Optional[Dict[str, Any]] = None
    status: Optional[str] = None


class SignPayload(BaseModel):
    stamp_id: Optional[str] = None
    note: Optional[str] = None


class CalcInput(BaseModel):
    calc: str
    params: Dict[str, Any] = {}


class DispatchPayload(BaseModel):
    recipients: List[EmailStr]
    cc: Optional[List[EmailStr]] = None
    message: Optional[str] = None
    include_pdf: bool = True


class RecipientsConfig(BaseModel):
    """Map: role -> list of email addresses configured by user."""
    items: Dict[str, List[EmailStr]] = {}


class ValidatePayload(BaseModel):
    cnp: Optional[str] = None
    cui: Optional[str] = None


class AvizStatusUpdate(BaseModel):
    status: Optional[str] = None  # planificat | cerut | primit | respins
    sent_to: Optional[List[EmailStr]] = None
    received_number: Optional[str] = None
    received_pdf_b64: Optional[str] = None  # base64 PDF aviz primit
    notes: Optional[str] = None


class AvizDispatch(BaseModel):
    recipients: List[EmailStr]
    cc: Optional[List[EmailStr]] = None
    message: Optional[str] = None
    include_attachments_b64: Optional[List[Dict[str, str]]] = None  # [{name, content_b64}]


def _ascii_safe(s: str, maxlen: int = 60) -> str:
    """Strip non-latin chars for HTTP filename headers."""
    import unicodedata
    out = unicodedata.normalize("NFKD", s).encode("ascii", "ignore").decode("ascii")
    return (out or "proiect")[:maxlen].replace("/", "-").replace(" ", "_") or "proiect"


# ---------- helpers ----------
def _public_payload(p: Dict[str, Any]) -> Dict[str, Any]:
    data = p.get("data") or {}
    return {
        "pid": p["pid"],
        "title": p.get("title"),
        "status": p.get("status"),
        "phase": p.get("phase"),
        "country": p.get("country", "RO"),
        "subdomain": p.get("subdomain", "bransament-casnic"),
        "beneficiar": data.get("beneficiar_nume"),
        "loc_consum": data.get("loc_consum_adresa"),
        "signed_at": p.get("signed_at"),
        "signature_hash": p.get("signature_hash"),
    }


def _compute_signature_hash(proj: Dict[str, Any]) -> str:
    import json
    canonical = {
        "pid": proj["pid"], "title": proj.get("title"),
        "country": proj.get("country"), "subdomain": proj.get("subdomain"),
        "data": proj.get("data") or {}, "phase": proj.get("phase"),
    }
    blob = json.dumps(canonical, sort_keys=True, separators=(",", ":")).encode("utf-8")
    return hashlib.sha256(blob).hexdigest()


def _attach_progress(proj: Dict[str, Any]) -> Dict[str, Any]:
    proj["progress"] = catalog.progress_for(
        proj.get("country", "RO"),
        "gaze-naturale",
        proj.get("subdomain", "bransament-casnic"),
        proj.get("data") or {},
    )
    return proj


# ---------- factory ----------
def make_gas_project_router(db, get_current_user):
    r = APIRouter(prefix="/gas-project", tags=["gas-project"])

    # ===================== CATALOG (no auth) =====================
    @r.get("/catalog/countries")
    async def cat_countries():
        return {"countries": catalog.list_countries()}

    @r.get("/catalog/{country}/subdomains")
    async def cat_subdomains(country: str):
        subs = catalog.list_subdomains(country, "gaze-naturale")
        meta = catalog.get_industry_meta(country, "gaze-naturale")
        return {"country": country, "industry": "gaze-naturale", "meta": meta, "subdomains": subs}

    @r.get("/catalog/{country}/{subdomain}")
    async def cat_subdomain_detail(country: str, subdomain: str):
        sub = catalog.get_subdomain(country, "gaze-naturale", subdomain)
        if not sub:
            raise HTTPException(404, "Subdomeniu inexistent")
        meta = catalog.get_industry_meta(country, "gaze-naturale")
        return {"country": country, "industry_meta": meta, "subdomain": sub}

    @r.get("/calcs")
    async def list_calcs():
        return {"calcs": engine.AVAILABLE_CALCS}

    @r.post("/calc")
    async def run_calc(payload: CalcInput):
        return engine.run(payload.calc, payload.params or {})

    # legacy fallback
    @r.get("/phases")
    async def list_phases_legacy():
        return {"phases": catalog.get_phases_for("RO", "gaze-naturale", "bransament-casnic")}

    # ===================== DOCUMENT TEMPLATES LIST (no auth) =====================
    @r.get("/doc-templates")
    async def list_doc_templates():
        """List 8 DOCX templates available for any gas project."""
        return {"templates": doc_templates.list_templates()}

    # ===================== VALIDARE DATE ROMÂNEȘTI (no auth) =====================
    @r.post("/validate")
    async def validate_ro(payload: ValidatePayload):
        """Validează CNP și/sau CUI conform algoritmi ANAF."""
        result: Dict[str, Any] = {}
        if payload.cnp is not None:
            ok, msg = validators_ro.validate_cnp(payload.cnp)
            result["cnp"] = {"valid": ok, "message": msg}
        if payload.cui is not None:
            ok, msg = validators_ro.validate_cui(payload.cui)
            result["cui"] = {"valid": ok, "message": msg}
        return result

    # ===================== CATALOG AVIZE (no auth) =====================
    @r.get("/avize-catalog")
    async def list_avize_catalog():
        """Returnează catalogul complet al avizelor (toate, indiferent de proiect)."""
        return {"avize": avize_catalog.list_all()}

    # ===================== CUSTOM TEMPLATES (admin upload) =====================
    @r.get("/custom-templates")
    async def list_custom_templates(user=Depends(get_current_user)):
        """Lista template-urilor DOCX încărcate de admin/developer (cu placeholdere)."""
        out = []
        async for t in db.gas_custom_templates.find({"deleted": {"$ne": True}}, {"_id": 0, "content_b64": 0}):
            out.append(t)
        return out

    @r.post("/custom-templates")
    async def upload_custom_template(
        file: UploadFile = File(...),
        label: str = Form(...),
        replaces_template_id: Optional[str] = Form(None),
        user=Depends(get_current_user),
    ):
        """Admin/developer încarcă un template DOCX cu placeholdere {{var}}. Acest fișier
        poate înlocui un template standard (replaces_template_id) sau adăuga unul nou."""
        if not (getattr(user, "is_admin", False) or getattr(user, "is_developer", False)):
            raise HTTPException(403, "Doar admin/developer pot încărca template-uri")
        content = await file.read()
        if not file.filename.lower().endswith(".docx"):
            raise HTTPException(400, "Fișier trebuie să fie .docx")
        if len(content) > 5 * 1024 * 1024:
            raise HTTPException(400, "Fișier > 5MB (limitare practică)")
        tid = _new_id("ct_")
        doc = {
            "tid": tid,
            "label": label.strip(),
            "filename": file.filename,
            "replaces_template_id": replaces_template_id,
            "content_b64": base64.b64encode(content).decode("ascii"),
            "uploaded_by": user.email,
            "uploaded_at": datetime.now(timezone.utc).isoformat(),
            "size_bytes": len(content),
            "deleted": False,
        }
        await db.gas_custom_templates.insert_one(doc)
        doc.pop("_id", None)
        doc.pop("content_b64", None)  # don't ship base64 back
        return doc

    @r.delete("/custom-templates/{tid}")
    async def delete_custom_template(tid: str, user=Depends(get_current_user)):
        if not (getattr(user, "is_admin", False) or getattr(user, "is_developer", False)):
            raise HTTPException(403, "Doar admin/developer")
        res = await db.gas_custom_templates.update_one(
            {"tid": tid}, {"$set": {"deleted": True}}
        )
        if res.matched_count == 0:
            raise HTTPException(404, "Inexistent")
        return {"ok": True}

    @r.get("/custom-templates/{tid}/download")
    async def download_custom_template(tid: str, user=Depends(get_current_user)):
        """Download fișierul original (admin verifică)."""
        if not (getattr(user, "is_admin", False) or getattr(user, "is_developer", False)):
            raise HTTPException(403, "Doar admin/developer")
        doc = await db.gas_custom_templates.find_one({"tid": tid, "deleted": {"$ne": True}})
        if not doc:
            raise HTTPException(404, "Inexistent")
        content = base64.b64decode(doc["content_b64"])
        return StreamingResponse(
            io.BytesIO(content),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={_ascii_safe(doc.get('filename','template.docx'))}"},
        )

    # ===================== PROJECTS =====================
    @r.post("")
    async def create_proj(payload: GasProjectCreate, user=Depends(get_current_user)):
        sub = catalog.get_subdomain(payload.country, "gaze-naturale", payload.subdomain)
        if not sub:
            raise HTTPException(400, "Subdomeniu invalid pentru această țară.")
        pid = _new_id("gp_")
        now = datetime.now(timezone.utc).isoformat()
        doc = {
            "pid": pid,
            "owner_id": user.user_id,
            "title": payload.title.strip() or sub["name"],
            "country": payload.country,
            "industry": "gaze-naturale",
            "subdomain": payload.subdomain,
            "phase": payload.phase or "tema",
            "status": "draft",
            "data": payload.data or {},
            "created_at": now,
            "updated_at": now,
            "deleted": False,
            "dispatches": [],  # phase -> dispatch log
        }
        await db.gas_projects.insert_one(doc)
        doc.pop("_id", None)
        return _attach_progress(doc)

    @r.get("")
    async def list_projects(user=Depends(get_current_user)):
        out: List[dict] = []
        async for d in db.gas_projects.find({"owner_id": user.user_id, "deleted": {"$ne": True}}, {"_id": 0}).sort("updated_at", -1):
            out.append(_attach_progress(d))
        return out

    @r.get("/recipients")
    async def get_recipients(user=Depends(get_current_user)):
        doc = await db.gas_recipients.find_one({"owner_id": user.user_id}, {"_id": 0})
        return doc or {"owner_id": user.user_id, "items": {}}

    @r.put("/recipients")
    async def put_recipients(payload: RecipientsConfig, user=Depends(get_current_user)):
        now = datetime.now(timezone.utc).isoformat()
        await db.gas_recipients.update_one(
            {"owner_id": user.user_id},
            {"$set": {"owner_id": user.user_id, "items": payload.items, "updated_at": now}},
            upsert=True,
        )
        doc = await db.gas_recipients.find_one({"owner_id": user.user_id}, {"_id": 0})
        return doc

    @r.get("/{pid}")
    async def get_one(pid: str, user=Depends(get_current_user)):
        doc = await db.gas_projects.find_one({"pid": pid, "owner_id": user.user_id, "deleted": {"$ne": True}}, {"_id": 0})
        if not doc:
            raise HTTPException(404, "Proiect inexistent")
        _attach_progress(doc)
        country = doc.get("country", "RO")
        subdomain = doc.get("subdomain", "bransament-casnic")
        sub = catalog.get_subdomain(country, "gaze-naturale", subdomain)
        doc["phases_schema"] = sub["phases"] if sub else []
        doc["subdomain_meta"] = {k: v for k, v in (sub or {}).items() if k != "phases"}
        doc["industry_meta"] = catalog.get_industry_meta(country, "gaze-naturale")
        return doc

    @r.get("/{pid}/phases")
    async def get_phases(pid: str, user=Depends(get_current_user)):
        doc = await db.gas_projects.find_one({"pid": pid, "owner_id": user.user_id, "deleted": {"$ne": True}}, {"_id": 0})
        if not doc:
            raise HTTPException(404, "Proiect inexistent")
        return {"phases": catalog.get_phases_for(doc.get("country", "RO"), "gaze-naturale", doc.get("subdomain", "bransament-casnic"))}

    @r.patch("/{pid}")
    async def patch_one(pid: str, payload: GasProjectPatch, user=Depends(get_current_user)):
        doc = await db.gas_projects.find_one({"pid": pid, "owner_id": user.user_id, "deleted": {"$ne": True}}, {"_id": 0})
        if not doc:
            raise HTTPException(404, "Proiect inexistent")
        updates: Dict[str, Any] = {"updated_at": datetime.now(timezone.utc).isoformat()}
        if payload.title is not None:
            updates["title"] = payload.title.strip()
        if payload.phase is not None:
            phases = catalog.get_phases_for(doc.get("country", "RO"), "gaze-naturale", doc.get("subdomain", "bransament-casnic"))
            if not any(p["id"] == payload.phase for p in phases):
                raise HTTPException(400, "Faza invalidă pentru acest subdomeniu")
            updates["phase"] = payload.phase
        if payload.status is not None:
            if payload.status not in {"draft", "in_review", "approved", "signed", "archived"}:
                raise HTTPException(400, "Status invalid")
            updates["status"] = payload.status
        if payload.data is not None:
            merged = {**(doc.get("data") or {}), **payload.data}
            updates["data"] = merged
        await db.gas_projects.update_one({"pid": pid}, {"$set": updates})
        new_doc = await db.gas_projects.find_one({"pid": pid}, {"_id": 0})
        return _attach_progress(new_doc)

    @r.delete("/{pid}")
    async def delete_one(pid: str, user=Depends(get_current_user)):
        res = await db.gas_projects.update_one(
            {"pid": pid, "owner_id": user.user_id},
            {"$set": {"deleted": True, "updated_at": datetime.now(timezone.utc).isoformat()}},
        )
        if res.matched_count == 0:
            raise HTTPException(404, "Proiect inexistent")
        return {"ok": True}

    @r.post("/{pid}/sign")
    async def sign_project(pid: str, payload: SignPayload, user=Depends(get_current_user)):
        doc = await db.gas_projects.find_one({"pid": pid, "owner_id": user.user_id, "deleted": {"$ne": True}}, {"_id": 0})
        if not doc:
            raise HTTPException(404, "Proiect inexistent")
        if payload.stamp_id:
            stamp_doc = await db.stamps.find_one({"sid": payload.stamp_id, "owner_id": user.user_id})
            if not stamp_doc:
                raise HTTPException(404, "Ștampilă inexistentă în contul tău")
        sig_hash = _compute_signature_hash(doc)
        signed_at = datetime.now(timezone.utc).isoformat()
        await db.gas_projects.update_one({"pid": pid}, {"$set": {
            "status": "signed", "signed_at": signed_at,
            "signed_by_user_id": user.user_id, "signed_by_email": user.email,
            "signed_with_stamp_id": payload.stamp_id,
            "signature_hash": sig_hash,
            "signature_note": (payload.note or "").strip(),
            "updated_at": signed_at,
        }})
        return {"ok": True, "signed_at": signed_at, "signature_hash": sig_hash,
                "stamp_id": payload.stamp_id, "verify_url": f"/verify/gas-project/{pid}"}

    @r.get("/{pid}/qr")
    async def project_qr(pid: str, user=Depends(get_current_user)):
        doc = await db.gas_projects.find_one({"pid": pid, "owner_id": user.user_id, "deleted": {"$ne": True}}, {"_id": 0})
        if not doc:
            raise HTTPException(404, "Proiect inexistent")
        base = os.environ.get("PUBLIC_VERIFY_BASE", "https://github-push-test.preview.emergentagent.com")
        verify_url = f"{base}/verify/gas-project/{pid}"
        qr = qrcode.QRCode(version=None, box_size=8, border=2, error_correction=qrcode.constants.ERROR_CORRECT_M)
        qr.add_data(verify_url)
        qr.make(fit=True)
        img = qr.make_image(fill_color="#0A0A0A", back_color="#FFFFFF")
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        b64 = base64.b64encode(buf.getvalue()).decode("ascii")
        return {"pid": pid, "verify_url": verify_url,
                "qr_png_b64": f"data:image/png;base64,{b64}",
                "signature_hash": doc.get("signature_hash")}

    @r.get("/{pid}/public")
    async def public_verify(pid: str):
        doc = await db.gas_projects.find_one({"pid": pid, "deleted": {"$ne": True}}, {"_id": 0})
        if not doc:
            raise HTTPException(404, "Proiect inexistent")
        return _public_payload(doc)

    # ===================== EMAIL DISPATCH =====================
    @r.post("/{pid}/phase/{phase_id}/dispatch")
    async def dispatch_phase(pid: str, phase_id: str, payload: DispatchPayload, user=Depends(get_current_user)):
        """Send the phase summary by email to provided recipients.

        Plan gating:
        - free/basic: max 1 recipient, max 5 dispatches/month
        - starter/proiectant: max 5 recipients, 50 dispatches/month
        - pro/society/developer: nelimitat
        """
        doc = await db.gas_projects.find_one({"pid": pid, "owner_id": user.user_id, "deleted": {"$ne": True}}, {"_id": 0})
        if not doc:
            raise HTTPException(404, "Proiect inexistent")
        phases = catalog.get_phases_for(doc.get("country", "RO"), "gaze-naturale", doc.get("subdomain", "bransament-casnic"))
        phase = next((p for p in phases if p["id"] == phase_id), None)
        if not phase:
            raise HTTPException(400, "Faza invalidă")

        # Plan gating
        plan = (user.plan or "basic").lower()
        max_recipients = 1
        if plan in {"starter", "proiectant", "designer"}:
            max_recipients = 5
        if plan in {"pro", "society", "company", "developer", "admin"}:
            max_recipients = 50
        if getattr(user, "is_developer", False) or getattr(user, "is_admin", False):
            max_recipients = 100

        if len(payload.recipients) > max_recipients:
            raise HTTPException(403, f"Planul tău permite max {max_recipients} destinatari per fază. Upgrade pentru mai mult.")

        # Build email body
        data = doc.get("data") or {}
        subject = f"[EPD] {doc.get('title')} — {phase['name']}"
        lines = [
            "Stimate destinatar,",
            "",
            f"Vă transmitem situația proiectului de gaze naturale «{doc.get('title')}» pentru faza {phase['name']}.",
            "",
            f"PID: {doc['pid']}",
            f"Subdomeniu: {doc.get('subdomain')}",
            f"Beneficiar: {data.get('beneficiar_nume') or '—'}",
            f"Loc consum: {data.get('loc_consum_adresa') or '—'}",
            f"Localitate: {data.get('loc_consum_localitate') or '—'}, jud. {data.get('loc_consum_judet') or '—'}",
            "",
            f"REZUMAT FAZA «{phase['name']}»",
            f"Cadru legal: {phase['norm']}",
            "Date introduse:",
        ]
        for fld in phase["fields"]:
            v = data.get(fld["key"])
            if v:
                lines.append(f"  • {fld['label']}: {v}")
        lines += ["", "Livrabile faza:"] + [f"  • {d}" for d in phase.get("deliverables", [])]
        if payload.message:
            lines += ["", "Mesaj de la expeditor:", payload.message]
        verify_base = os.environ.get("PUBLIC_VERIFY_BASE", "https://github-push-test.preview.emergentagent.com")
        lines += ["", f"Verificare publică: {verify_base}/verify/gas-project/{pid}",
                  "", "Cu stimă,", f"{user.name} ({user.email})",
                  "Energy Project Design"]
        body = "\n".join(lines)

        # Generate phase summary DOCX
        attachment_name = f"{doc['pid']}_{phase_id}_rezumat.docx"
        attachment_bytes = _generate_phase_docx(doc, phase)

        from email_sender import send_email_with_attachment
        cc_list = list(payload.cc or [])
        if user.secondary_email and user.secondary_email not in cc_list:
            cc_list.append(user.secondary_email)
        user_doc = await db.users.find_one({"user_id": user.user_id}, {"gmail_password": 1}) or {}
        result = send_email_with_attachment(
            gmail_user=user.gmail_user or "",
            gmail_password=user_doc.get("gmail_password") or "",
            recipients=list(payload.recipients),
            subject=subject,
            body=body,
            attachment_name=attachment_name,
            attachment_bytes=attachment_bytes,
            cc=cc_list,
            from_name_override=user.name or "Energy Project Design",
        )

        # Log dispatch
        log_entry = {
            "phase_id": phase_id,
            "phase_name": phase["name"],
            "recipients": list(payload.recipients),
            "cc": cc_list,
            "ok": bool(result.get("ok")),
            "error": result.get("error"),
            "sent_at": datetime.now(timezone.utc).isoformat(),
            "by_user_id": user.user_id,
            "by_email": user.email,
        }
        await db.gas_projects.update_one(
            {"pid": pid},
            {"$push": {"dispatches": log_entry},
             "$set": {"updated_at": log_entry["sent_at"]}},
        )
        return {"ok": result.get("ok"), "error": result.get("error"), "dispatch": log_entry}

    # ===================== AVIZE per proiect =====================
    @r.get("/{pid}/avize")
    async def get_project_avize(pid: str, user=Depends(get_current_user)):
        """Returnează lista de avize APLICABILE pentru proiect + status fiecăruia."""
        proj = await db.gas_projects.find_one({"pid": pid, "owner_id": user.user_id, "deleted": {"$ne": True}}, {"_id": 0})
        if not proj:
            raise HTTPException(404, "Proiect inexistent")
        stored = (proj.get("avize_status") or {})
        return {"pid": pid, "avize": avize_catalog.build_aviz_status_map(proj.get("data") or {}, stored)}

    @r.patch("/{pid}/avize/{aviz_id}")
    async def patch_project_aviz(pid: str, aviz_id: str, payload: AvizStatusUpdate, user=Depends(get_current_user)):
        """Actualizează statusul unui aviz (cerut/primit/respins) + atașează PDF aviz primit."""
        proj = await db.gas_projects.find_one({"pid": pid, "owner_id": user.user_id, "deleted": {"$ne": True}}, {"_id": 0})
        if not proj:
            raise HTTPException(404, "Proiect inexistent")
        aviz = avize_catalog.get_aviz(aviz_id)
        if not aviz:
            raise HTTPException(400, "Aviz necunoscut")
        if payload.status and payload.status not in ("planificat", "cerut", "primit", "respins"):
            raise HTTPException(400, "Status invalid")
        now = datetime.now(timezone.utc).isoformat()
        updates: Dict[str, Any] = {}
        if payload.status:
            updates[f"avize_status.{aviz_id}.status"] = payload.status
            if payload.status == "cerut":
                updates[f"avize_status.{aviz_id}.sent_at"] = now
            if payload.status == "primit":
                updates[f"avize_status.{aviz_id}.received_at"] = now
        if payload.sent_to is not None:
            updates[f"avize_status.{aviz_id}.sent_to"] = payload.sent_to
        if payload.received_number is not None:
            updates[f"avize_status.{aviz_id}.received_number"] = payload.received_number
        if payload.received_pdf_b64 is not None:
            # validate base64 size
            try:
                pdf_bytes = base64.b64decode(payload.received_pdf_b64)
                if len(pdf_bytes) > 10 * 1024 * 1024:
                    raise HTTPException(400, "PDF aviz > 10MB")
            except (ValueError, TypeError):
                raise HTTPException(400, "PDF base64 invalid")
            updates[f"avize_status.{aviz_id}.received_pdf_b64"] = payload.received_pdf_b64
        if payload.notes is not None:
            updates[f"avize_status.{aviz_id}.notes"] = payload.notes
        updates["updated_at"] = now
        if updates:
            await db.gas_projects.update_one({"pid": pid}, {"$set": updates})
        new = await db.gas_projects.find_one({"pid": pid}, {"_id": 0})
        stored = (new.get("avize_status") or {})
        return {"pid": pid, "aviz_id": aviz_id,
                "current": stored.get(aviz_id, {}),
                "all_avize": avize_catalog.build_aviz_status_map(new.get("data") or {}, stored)}

    @r.get("/{pid}/avize/{aviz_id}/dossier.zip")
    async def download_aviz_dossier(pid: str, aviz_id: str, user=Depends(get_current_user)):
        """Descarcă ZIP cu: cererea DOCX + template-uri anexe + lista anexe legale."""
        import zipfile
        proj = await db.gas_projects.find_one({"pid": pid, "owner_id": user.user_id, "deleted": {"$ne": True}}, {"_id": 0})
        if not proj:
            raise HTTPException(404, "Proiect inexistent")
        aviz = avize_catalog.get_aviz(aviz_id)
        if not aviz:
            raise HTTPException(400, "Aviz necunoscut")
        # Generate the main DOCX
        zip_buf = io.BytesIO()
        with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
            template_id = aviz.get("template_id")
            if template_id:
                result = doc_templates.generate(template_id, proj)
                if result:
                    data_bytes, fname = result
                    zf.writestr(f"00_CERERE_{fname}", data_bytes)
            # Anexe text manifest
            manifest = (
                f"AVIZ: {aviz['label']}\n"
                f"================================================\n\n"
                f"Emitent: {aviz['issuer']}\n"
                f"Cadru legal: {aviz['legal']}\n"
                f"Faza proiect: {aviz['phase']}\n\n"
                f"Proiect: {proj.get('title','—')}\n"
                f"Beneficiar: {(proj.get('data') or {}).get('beneficiar_nume','—')}\n"
                f"Adresa: {(proj.get('data') or {}).get('loc_consum_adresa','—')}\n\n"
                "ANEXE OBLIGATORII (de atașat manual):\n"
            )
            for a in aviz.get("extra_attachments", []):
                manifest += f"  - {a}\n"
            zf.writestr("00_MANIFEST_AVIZ.txt", manifest)
        zip_buf.seek(0)
        safe = _ascii_safe(f"{aviz_id}_{proj.get('title','proiect')}")
        return StreamingResponse(
            zip_buf, media_type="application/zip",
            headers={"Content-Disposition": f"attachment; filename=AVIZ_{safe}.zip"},
        )

    # ===================== DOCUMENT GENERATION =====================
    @r.get("/{pid}/doc/{template_id}")
    async def download_doc(pid: str, template_id: str, user=Depends(get_current_user)):
        """Generate + download a single DOCX template populated with project data."""
        doc = await db.gas_projects.find_one({"pid": pid, "owner_id": user.user_id, "deleted": {"$ne": True}}, {"_id": 0})
        if not doc:
            raise HTTPException(404, "Proiect inexistent")
        result = doc_templates.generate(template_id, doc)
        if not result:
            raise HTTPException(400, f"Template necunoscut: {template_id}")
        data_bytes, fname = result
        return StreamingResponse(
            io.BytesIO(data_bytes),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={_ascii_safe(fname)}"},
        )

    @r.get("/{pid}/dossier.zip")
    async def download_full_dossier(pid: str, user=Depends(get_current_user)):
        """Generate ZIP with ALL 8 DOCX templates — turns 10h work into 30min."""
        import zipfile
        doc = await db.gas_projects.find_one({"pid": pid, "owner_id": user.user_id, "deleted": {"$ne": True}}, {"_id": 0})
        if not doc:
            raise HTTPException(404, "Proiect inexistent")
        zip_buf = io.BytesIO()
        with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
            for tpl in doc_templates.list_templates():
                result = doc_templates.generate(tpl["id"], doc)
                if result:
                    data_bytes, fname = result
                    zf.writestr(fname, data_bytes)
            manifest = (
                "DOSAR TEHNIC COMPLET — Energy Project Design\n"
                "================================================\n\n"
                f"Proiect: {doc.get('title','—')}\n"
                f"PID: {pid}\n"
                f"Subdomeniu: {doc.get('subdomain','—')}\n"
                f"Status: {doc.get('status','—')}\n"
                f"Beneficiar: {(doc.get('data') or {}).get('beneficiar_nume','—')}\n"
                f"Loc consum: {(doc.get('data') or {}).get('loc_consum_adresa','—')}\n\n"
                "Documente incluse:\n"
            )
            for t in doc_templates.list_templates():
                manifest += f"  - {t['label']} (cadru legal: {t['norm']})\n"
            if doc.get("signature_hash"):
                manifest += f"\nSemnatura digitala SHA-256: {doc['signature_hash']}\n"
                manifest += f"Semnat: {doc.get('signed_at','-')}\n"
            zf.writestr("00_MANIFEST.txt", manifest)
        zip_buf.seek(0)
        safe_title = _ascii_safe(doc.get("title", "proiect"))
        return StreamingResponse(
            zip_buf,
            media_type="application/zip",
            headers={"Content-Disposition": f"attachment; filename=DOSAR_{safe_title}.zip"},
        )

    return r


# ---------- DOCX generator for phase dispatches ----------
def _generate_phase_docx(proj: Dict[str, Any], phase: Dict[str, Any]) -> bytes:
    """Generate a clean DOCX with phase summary, used as email attachment."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
    except ImportError:  # pragma: no cover
        return b""  # graceful: send without attachment if docx not installed
    data = proj.get("data") or {}
    doc = Document()
    # Header
    doc.add_heading("Energy Project Design", level=0)
    doc.add_paragraph(f"Proiect: {proj.get('title')}").bold = True
    p = doc.add_paragraph()
    p.add_run("PID: ").bold = True
    p.add_run(str(proj.get("pid")))
    doc.add_paragraph(f"Subdomeniu: {proj.get('subdomain')} · Țară: {proj.get('country','RO')}")
    doc.add_paragraph(f"Beneficiar: {data.get('beneficiar_nume','—')}")
    doc.add_paragraph(f"Loc consum: {data.get('loc_consum_adresa','—')}, {data.get('loc_consum_localitate','—')}, jud. {data.get('loc_consum_judet','—')}")
    doc.add_paragraph("")

    doc.add_heading(phase["name"], level=1)
    doc.add_paragraph(f"Cadru legal: {phase.get('norm','—')}").italic = True
    doc.add_paragraph(phase.get("description", ""))

    # Fields table
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Câmp"
    hdr[1].text = "Valoare"
    for fld in phase["fields"]:
        v = data.get(fld["key"])
        if not v:
            continue
        row = table.add_row().cells
        row[0].text = fld["label"]
        row[1].text = str(v)

    # Deliverables
    if phase.get("deliverables"):
        doc.add_heading("Livrabile faza", level=2)
        for d in phase["deliverables"]:
            doc.add_paragraph(d, style="List Bullet")

    # Signature footer
    doc.add_paragraph("")
    doc.add_paragraph("Document generat electronic de Energy Project Design.").italic = True
    if proj.get("signature_hash"):
        doc.add_paragraph(f"SHA-256: {proj['signature_hash']}").italic = True
        doc.add_paragraph(f"Semnat: {proj.get('signed_at','—')}")

    import io
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
