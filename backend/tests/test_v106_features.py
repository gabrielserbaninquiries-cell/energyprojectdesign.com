"""V10.6 backend tests:
- Hybrid Auth (cookie + Bearer fallback)
- Stripe LIVE donation 2 RON
- OCR template-placeholders, smart-extract-llm
- OCR /fill-template (DOCX completion, error handling)
- Regression: gas-project, dossier.zip, me/plan
"""
import io
import json
import os
import re
import zipfile

import pytest
import requests
from docx import Document

BASE_URL = os.environ.get("REACT_APP_BACKEND_URL", "https://github-push-test.preview.emergentagent.com").rstrip("/")
EMAIL = "dragosserban95@gmail.com"
PASSWORD = "Nuamparola_9"
DEMO_PID = "gp_e79e2810cc64b5b4"


# ---------- Fixtures ----------
@pytest.fixture(scope="session")
def login_payload():
    """Performs raw login once. Returns dict with token + cookie value."""
    sess = requests.Session()
    r = sess.post(f"{BASE_URL}/api/auth/login", json={"email": EMAIL, "password": PASSWORD}, timeout=30)
    assert r.status_code == 200, f"Login failed: {r.status_code} {r.text[:200]}"
    body = r.json()
    cookie_val = sess.cookies.get("session_token")
    return {"token": body["token"], "cookie": cookie_val, "user": body.get("user")}


@pytest.fixture(scope="session")
def bearer_headers(login_payload):
    return {"Authorization": f"Bearer {login_payload['token']}"}


@pytest.fixture(scope="session")
def test_docx_bytes() -> bytes:
    """Build a small DOCX with mixed placeholder styles."""
    doc = Document()
    doc.add_paragraph("Beneficiar: <NUME>")
    doc.add_paragraph("Adresa: <ADRESA>")
    doc.add_paragraph("Telefon: (Completați telefonul)")
    doc.add_paragraph("CNP: ________")
    doc.add_paragraph("Email: <EMAIL>")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ---------- Hybrid Auth ----------
class TestHybridAuth:
    def test_login_returns_token_and_cookie(self, login_payload):
        assert login_payload["token"], "JSON body must include token"
        assert isinstance(login_payload["token"], str) and len(login_payload["token"]) > 20
        assert login_payload["cookie"], "Set-Cookie session_token must be present"

    def test_login_cookie_attributes(self):
        """SameSite=None, Secure, HttpOnly required for cross-site mobile Safari."""
        r = requests.post(f"{BASE_URL}/api/auth/login",
                          json={"email": EMAIL, "password": PASSWORD}, timeout=30)
        assert r.status_code == 200
        sc = r.headers.get("set-cookie", "")
        assert "session_token=" in sc
        assert "HttpOnly" in sc
        assert "Secure" in sc
        assert re.search(r"SameSite=none", sc, re.IGNORECASE)

    def test_me_with_bearer_only_no_cookie(self, bearer_headers):
        """Mobile Safari ITP simulation — Bearer only, no cookie."""
        r = requests.get(f"{BASE_URL}/api/auth/me", headers=bearer_headers, timeout=15)
        assert r.status_code == 200, r.text[:200]
        data = r.json()
        assert data.get("email") == EMAIL

    def test_me_with_cookie_only(self, login_payload):
        """Standard cookie auth (desktop browsers)."""
        cookies = {"session_token": login_payload["cookie"]}
        r = requests.get(f"{BASE_URL}/api/auth/me", cookies=cookies, timeout=15)
        assert r.status_code == 200, r.text[:200]
        assert r.json().get("email") == EMAIL

    def test_me_with_both(self, login_payload, bearer_headers):
        cookies = {"session_token": login_payload["cookie"]}
        r = requests.get(f"{BASE_URL}/api/auth/me", headers=bearer_headers, cookies=cookies, timeout=15)
        assert r.status_code == 200
        assert r.json().get("email") == EMAIL


# ---------- Stripe LIVE donation ----------
class TestStripeDonation:
    def test_donation_2_ron_live(self, bearer_headers):
        r = requests.post(
            f"{BASE_URL}/api/donations/checkout",
            headers=bearer_headers,
            json={"amount": 2, "currency": "ron"},
            timeout=30,
        )
        # Could be 200 or 201
        assert r.status_code in (200, 201), f"Stripe donation failed: {r.status_code} {r.text[:300]}"
        data = r.json()
        assert "url" in data and "session_id" in data
        assert data["session_id"].startswith("cs_live_"), f"Expected LIVE session, got {data['session_id'][:20]}"
        assert "checkout.stripe.com" in data["url"] or "stripe.com" in data["url"]


# ---------- OCR placeholders ----------
class TestOCRPlaceholders:
    def test_template_placeholders_heuristic(self, bearer_headers, test_docx_bytes):
        files = {"file": ("test_template.docx", test_docx_bytes,
                          "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        r = requests.post(f"{BASE_URL}/api/ocr/template-placeholders",
                          headers=bearer_headers, files=files, timeout=60)
        assert r.status_code == 200, r.text[:300]
        data = r.json()
        assert "placeholders" in data
        ph = data["placeholders"]
        assert isinstance(ph, list)
        assert len(ph) >= 3, f"Expected ≥3 placeholders, got {len(ph)}: {ph}"
        # Verify shape on first entry
        sample = ph[0]
        for k in ("inner", "label", "type"):
            assert k in sample, f"placeholder missing key '{k}': {sample}"

    def test_smart_extract_llm(self, bearer_headers, test_docx_bytes):
        files = {"file": ("test_template.docx", test_docx_bytes,
                          "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        r = requests.post(f"{BASE_URL}/api/ocr/smart-extract-llm",
                          headers=bearer_headers, files=files, timeout=120)
        assert r.status_code == 200, r.text[:300]
        data = r.json()
        assert "placeholders" in data
        ph = data["placeholders"]
        assert isinstance(ph, list)
        assert len(ph) >= 1, f"AI returned no placeholders: {data}"
        sample = ph[0]
        for k in ("original", "suggested_key", "suggested_label", "field_type"):
            assert k in sample, f"AI placeholder missing key '{k}': {sample}"


# ---------- Fill template ----------
class TestFillTemplate:
    def test_fill_template_basic(self, bearer_headers, test_docx_bytes):
        replacements = [
            {"original": "<NUME>", "replacement": "Ion Popescu"},
            {"original": "<ADRESA>", "replacement": "Str. Aurel Vlaicu 15"},
            {"original": "<EMAIL>", "replacement": "ion@example.com"},
        ]
        files = {"file": ("template.docx", test_docx_bytes,
                          "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        data = {"replacements": json.dumps(replacements), "filename": "completat_test"}
        r = requests.post(f"{BASE_URL}/api/ocr/fill-template",
                          headers=bearer_headers, files=files, data=data, timeout=60)
        assert r.status_code == 200, r.text[:300]
        # Content-Type DOCX
        ct = r.headers.get("content-type", "")
        assert "wordprocessingml" in ct, f"Content-Type: {ct}"
        cd = r.headers.get("content-disposition", "")
        assert "attachment" in cd
        # Header X-Replacements-Applied present and ≥3
        applied = int(r.headers.get("X-Replacements-Applied", "0"))
        assert applied >= 3, f"Expected ≥3 replacements applied, got {applied}"

        # Verify content of returned DOCX
        out_doc = Document(io.BytesIO(r.content))
        full_text = "\n".join(p.text for p in out_doc.paragraphs)
        assert "Ion Popescu" in full_text, full_text
        assert "Str. Aurel Vlaicu 15" in full_text
        assert "ion@example.com" in full_text
        assert "<NUME>" not in full_text, "Placeholder should have been replaced"

    def test_fill_template_rejects_pdf(self, bearer_headers):
        # 5-byte PDF stub
        files = {"file": ("doc.pdf", b"%PDF-1.4 stub", "application/pdf")}
        data = {"replacements": "[]"}
        r = requests.post(f"{BASE_URL}/api/ocr/fill-template",
                          headers=bearer_headers, files=files, data=data, timeout=20)
        assert r.status_code == 400
        body = r.json()
        msg = (body.get("detail") or "").lower()
        assert "docx" in msg, f"Expected DOCX-only error, got: {body}"

    def test_fill_template_invalid_json(self, bearer_headers, test_docx_bytes):
        files = {"file": ("template.docx", test_docx_bytes,
                          "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        data = {"replacements": "not-a-json"}
        r = requests.post(f"{BASE_URL}/api/ocr/fill-template",
                          headers=bearer_headers, files=files, data=data, timeout=30)
        assert r.status_code == 400

    def test_fill_template_empty_replacements(self, bearer_headers, test_docx_bytes):
        files = {"file": ("template.docx", test_docx_bytes,
                          "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        data = {"replacements": "[]"}
        r = requests.post(f"{BASE_URL}/api/ocr/fill-template",
                          headers=bearer_headers, files=files, data=data, timeout=30)
        assert r.status_code == 200
        assert r.headers.get("X-Replacements-Applied") == "0"


# ---------- Regression ----------
class TestRegression:
    def test_me_plan(self, bearer_headers):
        r = requests.get(f"{BASE_URL}/api/me/plan", headers=bearer_headers, timeout=15)
        assert r.status_code == 200
        data = r.json()
        # accept either plan or plan_id key
        plan_val = data.get("plan") or data.get("plan_id") or ""
        assert plan_val in ("developer", "society_admin"), f"Unexpected plan: {data}"

    def test_get_demo_project(self, bearer_headers):
        r = requests.get(f"{BASE_URL}/api/gas-project/{DEMO_PID}", headers=bearer_headers, timeout=20)
        assert r.status_code == 200
        data = r.json()
        # API uses 'pid' key (per iteration_18 notes)
        assert data.get("pid") == DEMO_PID or data.get("id") == DEMO_PID

    def test_dossier_zip(self, bearer_headers):
        r = requests.get(f"{BASE_URL}/api/gas-project/{DEMO_PID}/dossier.zip",
                         headers=bearer_headers, timeout=120)
        assert r.status_code == 200
        zf = zipfile.ZipFile(io.BytesIO(r.content))
        assert len(zf.namelist()) >= 10, f"Expected ≥10 files, got {len(zf.namelist())}"
