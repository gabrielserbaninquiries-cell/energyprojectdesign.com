"""DOCX processing: detect placeholders, replace text, insert stamp image."""
import io
import re
from typing import Dict, List, Optional, Tuple
from docx import Document
from docx.shared import Cm, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Match {{variable_name}} or {variable_name}
PLACEHOLDER_RE = re.compile(r"\{\{\s*([a-zA-Z0-9_\-\.\s]+?)\s*\}\}|<\s*([a-zA-Z0-9_]+?)\s*>")

# Smart IF/ELSE placeholder: {IF var<10: text X ELSE text Y}
#   - operatori suportați: < <= > >= == !=
#   - var = orice key din values dict (numeric sau string)
#   - branch-ul ELSE este opțional
SMART_IF_RE = re.compile(
    r"\{IF\s+([a-zA-Z0-9_]+)\s*(<=|>=|==|!=|<|>)\s*([^:]+?)\s*:\s*([^|}]+?)(?:\s*ELSE\s+([^|}]+?))?\s*\}",
    re.IGNORECASE,
)


def _coerce(v):
    """Try numeric, fall back to stripped string (lowercased)."""
    try:
        return float(str(v).replace(",", ".").strip())
    except (ValueError, TypeError):
        return str(v).strip().lower()


def _eval_smart_if(match, values: Dict[str, str]) -> str:
    var = match.group(1)
    op = match.group(2)
    rhs_raw = match.group(3).strip().strip('"').strip("'")
    branch_if = match.group(4).strip()
    branch_else = (match.group(5) or "").strip()
    lhs = _coerce(values.get(var, ""))
    rhs = _coerce(rhs_raw)
    # If types mismatch (one numeric, one string), coerce both to strings for ==/!=
    if isinstance(lhs, float) != isinstance(rhs, float):
        lhs, rhs = str(lhs), str(rhs)
    try:
        if op == "<":
            ok = lhs < rhs
        elif op == "<=":
            ok = lhs <= rhs
        elif op == ">":
            ok = lhs > rhs
        elif op == ">=":
            ok = lhs >= rhs
        elif op == "==":
            ok = lhs == rhs
        elif op == "!=":
            ok = lhs != rhs
        else:
            ok = False
    except TypeError:
        ok = False
    return branch_if if ok else branch_else


def _match_key(m) -> str:
    return (m.group(1) or m.group(2)).strip()


def extract_placeholders(docx_bytes: bytes) -> List[str]:
    """Detect all unique {{placeholder}} and <placeholder> tokens in the DOCX text."""
    doc = Document(io.BytesIO(docx_bytes))
    found = set()

    def scan_text(text: str):
        for m in PLACEHOLDER_RE.finditer(text):
            found.add(_match_key(m))

    for p in doc.paragraphs:
        scan_text(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    scan_text(p.text)
    for section in doc.sections:
        for hdr in (section.header, section.footer):
            for p in hdr.paragraphs:
                scan_text(p.text)
    return sorted(found)


def _replace_in_paragraph(paragraph, values: Dict[str, str]):
    """Replace placeholders in a paragraph while preserving formatting where possible.

    Because python-docx splits runs unpredictably, we join run text, do the replacement,
    then put the full result into the first run and clear the others.

    Order: (1) smart {IF ...} blocks first, (2) plain {{var}} substitutions.
    """
    full_text = "".join(run.text for run in paragraph.runs)
    if not (PLACEHOLDER_RE.search(full_text) or SMART_IF_RE.search(full_text)):
        return

    def repl(match):
        key = _match_key(match)
        return str(values.get(key, match.group(0)))

    new_text = SMART_IF_RE.sub(lambda m: _eval_smart_if(m, values), full_text)
    new_text = PLACEHOLDER_RE.sub(repl, new_text)
    if new_text == full_text:
        return

    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)


def replace_placeholders(docx_bytes: bytes, values: Dict[str, str]) -> bytes:
    doc = Document(io.BytesIO(docx_bytes))

    for p in doc.paragraphs:
        _replace_in_paragraph(p, values)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph(p, values)

    for section in doc.sections:
        for hdr in (section.header, section.footer):
            for p in hdr.paragraphs:
                _replace_in_paragraph(p, values)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def insert_stamp(
    docx_bytes: bytes,
    stamp_bytes: bytes,
    position: str = "bottom-right",
    size_cm: float = 4.0,
    x_cm: Optional[float] = None,
    y_cm: Optional[float] = None,
    page_index: int = 0,
) -> bytes:
    """Insert stamp into a docx file.

    Two placement modes:
      • Anchor presets (`position` in top-left/top-right/bottom-left/bottom-right/center)
      • Absolute coordinates (`x_cm`, `y_cm` measured from page top-left, in centimeters).
        When x_cm/y_cm are provided, the stamp is rendered as a floating image
        anchored absolutely on the chosen page (default: page 1).
    """
    from docx.oxml.ns import qn, nsmap
    from docx.oxml import OxmlElement
    import uuid as _uuid

    doc = Document(io.BytesIO(docx_bytes))

    # Absolute positioning mode
    if x_cm is not None and y_cm is not None:
        # Anchor a floating picture in a paragraph (run) with wp:anchor positioning.
        target_p = doc.add_paragraph()
        run = target_p.add_run()
        run.add_picture(io.BytesIO(stamp_bytes), width=Cm(size_cm))
        # Now wrap inline -> anchor with absolute coords
        # Locate the drawing inline element and replace it with wp:anchor
        drawing = run._r.find(qn("w:drawing"))
        if drawing is not None:
            inline = drawing.find(qn("wp:inline"))
            if inline is not None:
                # Convert wp:inline → wp:anchor
                inline.tag = qn("wp:anchor")
                anchor = inline
                # Required anchor attributes
                anchor.set("behindDoc", "0")
                anchor.set("locked", "0")
                anchor.set("layoutInCell", "1")
                anchor.set("allowOverlap", "1")
                anchor.set("relativeHeight", "251658240")
                anchor.set("simplePos", "0")
                anchor.set("distT", "0"); anchor.set("distB", "0")
                anchor.set("distL", "0"); anchor.set("distR", "0")
                # <wp:simplePos x="0" y="0"/>
                sp = OxmlElement("wp:simplePos")
                sp.set("x", "0"); sp.set("y", "0")
                anchor.insert(0, sp)
                # Positions (EMU; 1 cm = 360000 EMU)
                emu_x = int(float(x_cm) * 360000)
                emu_y = int(float(y_cm) * 360000)
                pH = OxmlElement("wp:positionH"); pH.set("relativeFrom", "page")
                pHoff = OxmlElement("wp:posOffset"); pHoff.text = str(emu_x)
                pH.append(pHoff)
                pV = OxmlElement("wp:positionV"); pV.set("relativeFrom", "page")
                pVoff = OxmlElement("wp:posOffset"); pVoff.text = str(emu_y)
                pV.append(pVoff)
                # Insert positionH/V after simplePos
                anchor.insert(1, pH); anchor.insert(2, pV)
                # Wrap mode → wrapNone (floating over text)
                wrap_none = OxmlElement("wp:wrapNone")
                # Find extent + docPr + graphic, insert wrapNone before docPr if exists
                extent = anchor.find(qn("wp:extent"))
                if extent is not None:
                    extent.addnext(wrap_none)
        out = io.BytesIO()
        doc.save(out)
        return out.getvalue()

    # Legacy preset positioning
    align_map = {
        "top-left": WD_ALIGN_PARAGRAPH.LEFT,
        "top-right": WD_ALIGN_PARAGRAPH.RIGHT,
        "bottom-left": WD_ALIGN_PARAGRAPH.LEFT,
        "bottom-right": WD_ALIGN_PARAGRAPH.RIGHT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
    }
    align = align_map.get(position, WD_ALIGN_PARAGRAPH.RIGHT)
    if position.startswith("top"):
        new_p = doc.paragraphs[0].insert_paragraph_before("") if doc.paragraphs else doc.add_paragraph()
        target_p = new_p
    else:
        target_p = doc.add_paragraph()
    target_p.alignment = align
    run = target_p.add_run()
    run.add_picture(io.BytesIO(stamp_bytes), width=Cm(size_cm))

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()
