"""Gas Natural — Template-uri DOCX SUPLIMENTARE V6.4 (legal extensions).

Acest modul adaugă 6 template-uri lipsă necesare pentru un dosar tehnic REAL
de branșament gaze, conform Legea 10/1995 + HG 273/1994 + NTPEE 2018:

1. pv_lucrari_ascunse       — PV lucrări ascunse (săpătură + acoperire șanț)
2. pv_faza_determinanta     — PV control fază determinantă (sudură, probe)
3. program_control_calitate — PCC semnat de proiectant + executant + diriginte
4. referat_verificator      — Referat Verificator Tehnic atestat (RVT VGD)
5. notificare_isc           — Notificare Inspectoratul de Stat în Construcții
6. as_built                 — Plan / Memoriu as-built (după execuție)
"""
from __future__ import annotations
from typing import Any, Dict

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

import gas_doc_templates as base

COMPANY = base.COMPANY


def _pv_la(proj: Dict[str, Any]) -> bytes:
    data = proj.get("data") or {}
    doc = Document()
    base._company_header(doc)
    base._add_heading(doc, "PROCES-VERBAL DE LUCRĂRI ASCUNSE", level=1,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
    base._add_para(doc, "(Conform Legea 10/1995 art. 23 + HG 273/1994 + NTPEE 2018)",
                   italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
    doc.add_paragraph()

    base._add_kv_table(doc, [
        ("Obiectiv", proj.get("title", "—")),
        ("Beneficiar", base._get(data, "beneficiar_nume")),
        ("Adresă", base._get(data, "loc_consum_adresa")),
        ("AC nr.", f"{base._get(data,'ac_numar','—')} / {base._get(data,'ac_data_emitere','—')}"),
        ("Executant", base._get(data, "exec_firma")),
        ("Data PV", base._get(data, "la_sapatura_data", "______________")),
    ])

    doc.add_paragraph()
    base._add_para(doc, "Lucrări verificate și acceptate spre acoperire:", bold=True)
    base._add_kv_table(doc, [
        ("Adâncime efectivă pozare (cm)", base._get(data, "la_sapatura_adancime_cm")),
        ("Grosime strat nisip (cm)", base._get(data, "la_strat_nisip_cm")),
        ("Bandă avertizoare galbenă", base._get(data, "la_banda_avertizoare", "Da")),
        ("Protecție mecanică / tubaj", base._get(data, "la_protectie_mecanica")),
        ("Compactare strat acoperire", base._get(data, "la_compactare_strat")),
        ("Material conductă", base._get(data, "sf_material_conducta")),
        ("Diametru DN", base._get(data, "sf_diametru_nominal_DN")),
        ("Lungime efectivă (m)", base._get(data, "asb_lungime_efectiva_m", base._get(data, "sf_lungime_conducta_m"))),
    ])

    doc.add_paragraph()
    base._add_para(doc, "Constatări:", bold=True)
    doc.add_paragraph("• Conducta respectă traseul aprobat din DTAC", style="List Bullet")
    doc.add_paragraph("• Adâncimea de pozare este conformă NTPEE 2018 (min. 90 cm sub trotuar / 110 cm sub carosabil)", style="List Bullet")
    doc.add_paragraph("• Stratul de nisip este uniform pe toată lungimea", style="List Bullet")
    doc.add_paragraph("• Banda avertizoare este montată la 30 cm deasupra conductei", style="List Bullet")
    doc.add_paragraph("• Toate sudurile au fost executate de sudor autorizat", style="List Bullet")

    doc.add_paragraph()
    base._add_para(doc, "Concluzie: Se aprobă acoperirea șanțului.", bold=True)
    doc.add_paragraph()
    base._add_para(doc, "Semnatari:", bold=True)
    t = doc.add_table(rows=2, cols=3)
    t.rows[0].cells[0].text = "Executant (RTE)"
    t.rows[0].cells[1].text = "Diriginte șantier"
    t.rows[0].cells[2].text = "Proiectant"
    t.rows[1].cells[0].text = f"{base._get(data,'exec_responsabil_tehnic')}\n______________"
    t.rows[1].cells[1].text = f"{base._get(data,'exec_diriginte_santier')}\nAut. MDLPA {base._get(data,'diriginte_autorizatie','—')}\n______________"
    t.rows[1].cells[2].text = f"{base._get(data,'dtac_proiectant_specialitate',COMPANY['name'])}\n______________"
    base._footer_signature(doc, proj, data)
    return base._save(doc)


def _pv_fd(proj: Dict[str, Any]) -> bytes:
    data = proj.get("data") or {}
    doc = Document()
    base._company_header(doc)
    base._add_heading(doc, "PROCES-VERBAL CONTROL FAZĂ DETERMINANTĂ", level=1,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
    base._add_para(doc, "(Conform Legea 10/1995 art. 22 + HG 1735/2006 + ISC)",
                   italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
    doc.add_paragraph()

    base._add_kv_table(doc, [
        ("Obiectiv", proj.get("title", "—")),
        ("Adresă", base._get(data, "loc_consum_adresa")),
        ("AC", f"{base._get(data,'ac_numar','—')} / {base._get(data,'ac_data_emitere','—')}"),
        ("PCC referință", base._get(data, "fd_pcc_versiune")),
        ("ISC notificat", base._get(data, "fd_isc_notificat", "Nu")),
        ("Data notificare ISC", base._get(data, "fd_isc_data_notificare")),
    ])

    doc.add_paragraph()
    base._add_para(doc, "Stadii verificate:", bold=True)
    base._add_kv_table(doc, [
        ("FD Săpătură & pozare", base._get(data, "fd_sapatura", "Programat")),
        ("FD Sudură PE/OL", base._get(data, "fd_sudura", "Programat")),
        ("FD Proba de presiune", base._get(data, "fd_proba_presiune", "Programat")),
        ("FD Acoperire șanț", base._get(data, "fd_acoperire_sant", "Programat")),
    ])

    doc.add_paragraph()
    base._add_para(doc, "Constatări la fază determinantă:", bold=True)
    base._add_para(doc, base._get(data, "proba_observatii",
                   "Lucrările respectă proiectul aprobat. Nu există abateri. Se aprobă continuarea execuției."),
                   italic=True)

    doc.add_paragraph()
    base._add_para(doc, "Semnatari (obligatorii conform Legea 10/1995):", bold=True)
    t = doc.add_table(rows=2, cols=4)
    t.rows[0].cells[0].text = "Beneficiar"
    t.rows[0].cells[1].text = "Proiectant"
    t.rows[0].cells[2].text = "Executant (RTE)"
    t.rows[0].cells[3].text = "Diriginte șantier"
    t.rows[1].cells[0].text = f"{base._get(data,'beneficiar_nume')}\n_____________"
    t.rows[1].cells[1].text = f"{base._get(data,'dtac_proiectant_specialitate',COMPANY['name'])}\n_____________"
    t.rows[1].cells[2].text = f"{base._get(data,'exec_responsabil_tehnic')}\n_____________"
    t.rows[1].cells[3].text = f"{base._get(data,'exec_diriginte_santier')}\n_____________"
    base._footer_signature(doc, proj, data)
    return base._save(doc)


def _pcc(proj: Dict[str, Any]) -> bytes:
    data = proj.get("data") or {}
    doc = Document()
    base._company_header(doc)
    base._add_heading(doc, "PROGRAM DE CONTROL AL CALITĂȚII LUCRĂRILOR (PCC)",
                      level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    base._add_para(doc, "(Conform Legea 10/1995 + Ord. MLPAT 12/N/1995 + NTPEE 2018 cap. 9)",
                   italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
    doc.add_paragraph()

    base._add_kv_table(doc, [
        ("Obiectiv", proj.get("title", "—")),
        ("Versiune PCC", base._get(data, "fd_pcc_versiune", "1.0")),
        ("Data întocmire", base._get(data, "fd_pcc_data")),
        ("Proiectant", base._get(data, "dtac_proiectant_specialitate", COMPANY["name"])),
        ("Executant", base._get(data, "exec_firma")),
        ("Diriginte", base._get(data, "exec_diriginte_santier")),
    ])

    doc.add_paragraph()
    base._add_para(doc, "Faze de control prevăzute:", bold=True)
    t = doc.add_table(rows=8, cols=4)
    t.style = "Light Grid Accent 1"
    headers = ["Faza", "Document control", "Participanți", "Tip control"]
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    rows = [
        ("Predare amplasament", "PV predare-primire", "Benef + Exec + Diri", "Verificare"),
        ("Sudură conductă", "PV faza det. + buletin sudor", "Proiectant + RTE + Diri", "Faza determinantă"),
        ("Pozare + Adâncime", "PV lucrări ascunse", "RTE + Diri", "Lucrare ascunsă"),
        ("Proba rezistență", "Buletin probe + diagramă", "RTE + Diri + OSD", "Faza determinantă"),
        ("Proba etanșeitate", "Buletin probe + diagramă", "RTE + Diri + OSD", "Faza determinantă"),
        ("Acoperire șanț", "PV lucrări ascunse", "RTE + Diri", "Lucrare ascunsă"),
        ("Recepție terminare", "PV recepție", "Comisie completă", "Recepție"),
    ]
    for i, r in enumerate(rows, 1):
        for j, val in enumerate(r):
            t.rows[i].cells[j].text = val

    doc.add_paragraph()
    base._add_para(doc, "Cerințe esențiale acoperite (Legea 10/1995):", bold=True)
    doc.add_paragraph("A — Rezistență și stabilitate: probe presiune + verificare dimensionare", style="List Bullet")
    doc.add_paragraph("B — Siguranță în exploatare: robinete de izolare + semnalizare", style="List Bullet")
    doc.add_paragraph("C — Siguranță la foc: distanțe minime + ventilare + odorizare gaz", style="List Bullet")
    doc.add_paragraph("D — Sănătate / mediu: evacuare deșeuri + protecție mediu", style="List Bullet")

    doc.add_paragraph()
    base._add_para(doc, "Semnatari PCC:", bold=True)
    sig_t = doc.add_table(rows=2, cols=3)
    sig_t.rows[0].cells[0].text = "Proiectant"
    sig_t.rows[0].cells[1].text = "Executant (RTE)"
    sig_t.rows[0].cells[2].text = "Diriginte șantier"
    sig_t.rows[1].cells[0].text = f"{base._get(data,'dtac_proiectant_specialitate',COMPANY['name'])}\n_____________"
    sig_t.rows[1].cells[1].text = f"{base._get(data,'exec_responsabil_tehnic')}\n_____________"
    sig_t.rows[1].cells[2].text = f"{base._get(data,'exec_diriginte_santier')}\n_____________"

    base._footer_signature(doc, proj, data)
    return base._save(doc)


def _rvt(proj: Dict[str, Any]) -> bytes:
    data = proj.get("data") or {}
    doc = Document()
    base._company_header(doc)
    base._add_heading(doc, "REFERAT DE VERIFICARE TEHNICĂ (RVT)", level=1,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
    base._add_para(doc, "(Conform Legea 10/1995 art. 21 + Ord. MLPAT 777/2003 — Verificator Atestat MDLPA)",
                   italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
    doc.add_paragraph()

    base._add_kv_table(doc, [
        ("Nr. RVT", base._get(data, "rvt_nr", "______")),
        ("Data RVT", base._get(data, "rvt_data")),
        ("Obiectiv", proj.get("title", "—")),
        ("Beneficiar", base._get(data, "beneficiar_nume")),
        ("Adresă", base._get(data, "loc_consum_adresa")),
        ("Proiectant", base._get(data, "dtac_proiectant_specialitate", COMPANY["name"])),
        ("Verificator VGD", base._get(data, "dtac_verificator_vgd", "______")),
        ("Domeniu verificare", "Is — Instalații pentru gaze naturale (NTPEE 2018)"),
    ])

    doc.add_paragraph()
    base._add_para(doc, "Obiectul verificării:", bold=True)
    base._add_para(doc, "Documentația tehnică (DTAC + PT) pentru:", italic=True)
    base._add_para(doc, base._get(data, "scop_lucrare"), italic=True)

    doc.add_paragraph()
    base._add_para(doc, "Norme aplicate la verificare:", bold=True)
    doc.add_paragraph("• NTPEE 2018 (Ord. ANRE 89/2018) — calcule + dimensionare + materiale", style="List Bullet")
    doc.add_paragraph("• HG 907/2016 — conținut DTAC + memoriu tehnic", style="List Bullet")
    doc.add_paragraph("• Legea 50/1991 — autorizare construire", style="List Bullet")
    doc.add_paragraph("• Legea 10/1995 — cerințe esențiale A/B/C/D", style="List Bullet")
    doc.add_paragraph("• STAS 6657-3, 6724 — rezistență + ventilație", style="List Bullet")

    doc.add_paragraph()
    base._add_para(doc, "Concluzii:", bold=True)
    base._add_para(doc, f"Documentația este: {base._get(data, 'rvt_concluzii', 'Acceptat')}", bold=True)
    base._add_para(doc, base._get(data, "rvt_observatii",
                   "Documentația respectă normele tehnice. Nu există abateri majore. Se avizează spre execuție."),
                   italic=True)

    doc.add_paragraph()
    base._add_para(doc, "Semnătură verificator atestat:", bold=True)
    base._add_para(doc, f"{base._get(data,'dtac_verificator_vgd')} — Atestat MDLPA: ______ — Domeniu: Is", italic=True)
    base._add_para(doc, "Ștampilă + semnătură: ______________", italic=True)

    base._footer_signature(doc, proj, data)
    return base._save(doc)


def _notificare_isc(proj: Dict[str, Any]) -> bytes:
    data = proj.get("data") or {}
    doc = Document()
    base._company_header(doc)
    base._add_para(doc, f"Către Inspectoratul de Stat în Construcții, {base._get(data,'isc_judet','Județean')}",
                   bold=True, size=12)
    doc.add_paragraph()
    base._add_heading(doc, "NOTIFICARE — ÎNCEPERE LUCRĂRI",
                      level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    base._add_para(doc, "(Conform Legea 50/1991 art. 7 alin. 8 + HG 1735/2006)",
                   italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
    doc.add_paragraph()

    base._add_para(doc, "Vă notificăm începerea lucrărilor de execuție pentru:", bold=True)
    base._add_kv_table(doc, [
        ("Obiectiv", proj.get("title", "—")),
        ("Adresă", base._get(data, "loc_consum_adresa")),
        ("Beneficiar", base._get(data, "beneficiar_nume")),
        ("Autorizație de Construire nr.", f"{base._get(data,'ac_numar','—')} / {base._get(data,'ac_data_emitere','—')}"),
        ("Emisă de", base._get(data, "ac_emitent")),
        ("Termen execuție AC (luni)", base._get(data, "ac_termen_executie")),
        ("Data începere lucrări", base._get(data, "exec_data_start")),
        ("Data finalizare estimată", base._get(data, "exec_data_terminare")),
        ("Executant autorizat ANRE", base._get(data, "exec_firma")),
        ("Responsabil tehnic execuție (RTE)", base._get(data, "exec_responsabil_tehnic")),
        ("Diriginte șantier (autorizație MDLPA)", base._get(data, "exec_diriginte_santier")),
        ("Aut. diriginte șantier", base._get(data, "diriginte_autorizatie")),
        ("Contact diriginte", f"{base._get(data,'diriginte_telefon','—')} · {base._get(data,'diriginte_email','—')}"),
    ])

    doc.add_paragraph()
    base._add_para(doc, "Anexe:", bold=True)
    doc.add_paragraph("• Copie Autorizație de Construire", style="List Bullet")
    doc.add_paragraph("• Copie ATR (Aviz Tehnic Racordare)", style="List Bullet")
    doc.add_paragraph("• Copie autorizație ANRE executant", style="List Bullet")
    doc.add_paragraph("• Copie autorizație diriginte șantier MDLPA", style="List Bullet")
    doc.add_paragraph("• Program de Control al Calității (PCC)", style="List Bullet")

    base._footer_signature(doc, proj, data)
    return base._save(doc)


def _as_built(proj: Dict[str, Any]) -> bytes:
    data = proj.get("data") or {}
    doc = Document()
    base._company_header(doc)
    base._add_heading(doc, "MEMORIU TEHNIC AS-BUILT (Plan executat)",
                      level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    base._add_para(doc, "(Conform HG 273/1994 + Ord. MLPAT 770/1997 — Cartea Tehnică)",
                   italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
    doc.add_paragraph()

    base._add_kv_table(doc, [
        ("Obiectiv", proj.get("title", "—")),
        ("Beneficiar", base._get(data, "beneficiar_nume")),
        ("Adresă execuție", base._get(data, "loc_consum_adresa")),
        ("AC", f"{base._get(data,'ac_numar','—')} / {base._get(data,'ac_data_emitere','—')}"),
        ("Data finalizare", base._get(data, "exec_data_terminare")),
    ])

    doc.add_paragraph()
    base._add_para(doc, "Cantități realizate (as-built):", bold=True)
    base._add_kv_table(doc, [
        ("Traseu modificat față de proiect", base._get(data, "asb_traseu_modificat", "Nu")),
        ("Lungime efectivă conductă (m)", base._get(data, "asb_lungime_efectiva_m", base._get(data, "sf_lungime_conducta_m"))),
        ("Material conductă", base._get(data, "sf_material_conducta")),
        ("Diametru DN", base._get(data, "sf_diametru_nominal_DN")),
        ("Număr sudări executate", base._get(data, "asb_numar_sudari")),
        ("Adâncime medie pozare (cm)", base._get(data, "la_sapatura_adancime_cm")),
        ("Serie robinet branșament", base._get(data, "mat_serie_robinet_br")),
        ("Marcă contor", base._get(data, "mat_marca_contor")),
    ])

    doc.add_paragraph()
    base._add_para(doc, "Coordonate GPS puncte cheie:", bold=True)
    base._add_para(doc, base._get(data, "asb_coords_gps", "(De completat în teren — start / capete / fitting-uri)"),
                   italic=True)

    doc.add_paragraph()
    base._add_para(doc, "Probe finale executate:", bold=True)
    base._add_kv_table(doc, [
        ("Proba rezistență (bar)", base._get(data, "proba_rezistenta_bar")),
        ("Durată proba rezistență (min)", base._get(data, "proba_rezistenta_durata_min")),
        ("Proba etanșeitate (bar)", base._get(data, "proba_etanseitate_bar")),
        ("Durată proba etanșeitate (ore)", base._get(data, "proba_etanseitate_durata_h")),
        ("Rezultat probe", base._get(data, "proba_rezultat", "Admis")),
    ])

    doc.add_paragraph()
    base._add_para(doc, "Acest memoriu se anexează la Cartea Tehnică a Construcției - Secțiunea C (Recepție).",
                   italic=True)

    base._footer_signature(doc, proj, data)
    return base._save(doc)


# ============================================================================
# TEMPLATE NOU 7: DTAC Lista Avize Utilități (consumă 11 avize + acord acces)
# ============================================================================
def _dtac_lista_avize(proj: Dict[str, Any]) -> bytes:
    data = proj.get("data") or {}
    doc = Document()
    base._company_header(doc)
    base._add_heading(doc, "LISTA AVIZE OBȚINUTE PENTRU DTAC",
                      level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    base._add_para(doc, "(Anexă la Documentația Tehnică pentru Autorizarea Construcției — Legea 50/1991)",
                   italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
    doc.add_paragraph()

    base._add_kv_table(doc, [
        ("Obiectiv", proj.get("title", "—")),
        ("Beneficiar", base._get(data, "beneficiar_nume")),
        ("Loc consum", base._get(data, "loc_consum_adresa")),
        ("Nr. proiect / an", base._get(data, "proiect_nr_an")),
    ])

    doc.add_paragraph()
    base._add_para(doc, "Avize utilități obținute pentru emiterea Autorizației de Construire:", bold=True)

    AVIZE_LIST = [
        ("01", "E-Distribuție Muntenia (electric)",  "aviz_edistr_nr_data"),
        ("02", "Telekom România (telecom)",          "aviz_telekom_nr_data"),
        ("03", "Apa Nova / Operator apă-canal",      "aviz_apa_nr_data"),
        ("04", "STB transport public",               "aviz_stb_nr_data"),
        ("05", "NetCity (fibră optică)",             "aviz_netcity_nr_data"),
        ("06", "Luxten (iluminat public)",           "aviz_luxten_nr_data"),
        ("07", "Direcția Străzi PMB",                "aviz_strazi_nr_data"),
        ("08", "Direcția Circulație PMB",            "aviz_circulatie_pmb_nr_data"),
        ("09", "Direcția Mediu PMB",                 "aviz_mediu_pmb_nr_data"),
        ("10", "APM — Agenția pentru Mediu",         "aviz_apm_nr_data"),
        ("11", "Acord de acces proprietate privată", "acord_acces_nr_data"),
    ]
    t = doc.add_table(rows=1, cols=4)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    hdr[0].text = "Nr."; hdr[1].text = "Aviz emitent"; hdr[2].text = "Nr. / Data"; hdr[3].text = "Status"
    for n, lbl, key in AVIZE_LIST:
        r = t.add_row().cells
        r[0].text = n
        r[1].text = lbl
        v = base._get(data, key, "—")
        r[2].text = v
        r[3].text = "Obținut" if v not in ("—", "", None) else "În curs"

    doc.add_paragraph()
    base._add_para(doc, "Notă: Toate avizele se prezintă în original la depunerea DTAC la primăria competentă "
                        "(Sector pentru București, Primăria Municipiului/Orașului pentru rest).",
                   italic=True, size=9)

    base._footer_signature(doc, proj, data)
    return base._save(doc)


# ============================================================================
# TEMPLATE NOU 8: PV Calitate (consumă 9 fields pv_calitate_*)
# ============================================================================
def _pv_calitate(proj: Dict[str, Any]) -> bytes:
    data = proj.get("data") or {}
    doc = Document()
    base._company_header(doc)
    base._add_heading(doc, "PROCES-VERBAL CONTROL CALITATE LUCRĂRI",
                      level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    base._add_para(doc, "(Conform Legea 10/1995 + HG 925/1995 + NTPEE 2018)",
                   italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
    doc.add_paragraph()

    base._add_kv_table(doc, [
        ("PV nr.",   base._get(data, "pv_calitate_pv_numar")),
        ("Data",     base._get(data, "pv_calitate_data", base._today_ro())),
        ("Obiectiv", proj.get("title", "—")),
        ("Beneficiar", base._get(data, "beneficiar_nume")),
        ("Adresă",   base._get(data, "loc_consum_adresa")),
    ])

    doc.add_paragraph()
    base._add_para(doc, "Părți participante:", bold=True)
    base._add_kv_table(doc, [
        ("Proiectant",    base._get(data, "pv_calitate_proiectant",  base._get(data, "dtac_proiectant_specialitate", COMPANY["name"]))),
        ("Executant",     base._get(data, "pv_calitate_constructor", base._get(data, "exec_firma"))),
        ("Diriginte șantier", base._get(data, "pv_calitate_diriginte", base._get(data, "exec_diriginte_santier"))),
    ])

    doc.add_paragraph()
    base._add_para(doc, "1. Documente de bază verificate:", bold=True)
    base._add_para(doc, base._get(data, "pv_calitate_documente_baza",
                                  "Proiect tehnic, caiet de sarcini, plan de control calitate (PCC), "
                                  "certificate calitate materiale, fișe sudori autorizați, jurnal de șantier."),
                   italic=True)

    doc.add_paragraph()
    base._add_para(doc, "2. Lucrări verificate:", bold=True)
    base._add_para(doc, base._get(data, "pv_calitate_lucrari",
                                  "Trasaj conductă, săpătură șanț, pat nisip, pozare conductă, "
                                  "sudură electrofuziune, acoperire șanț, marcaj banda avertizare, "
                                  "montaj robinet branșament, instalație utilizare."),
                   italic=True)

    doc.add_paragraph()
    base._add_para(doc, "3. Constatări:", bold=True)
    base._add_para(doc, base._get(data, "pv_calitate_constatari",
                                  "Lucrările au fost executate conform proiectului tehnic și a NTPEE 2018. "
                                  "Adâncimea de pozare verificată (min. 0,9 m sub trotuar / 1,0 m sub trafic). "
                                  "Banda avertizare poziționată la 30 cm deasupra conductei. "
                                  "Sudurile prezintă uniformitate vizuală și au fost executate de sudori atestați."),
                   italic=True)

    doc.add_paragraph()
    base._add_para(doc, "4. Concluzii:", bold=True)
    base._add_para(doc, base._get(data, "pv_calitate_concluzii",
                                  "Calitatea lucrărilor executate este CORESPUNZĂTOARE. "
                                  "Se autorizează continuarea lucrărilor și efectuarea probelor de presiune."),
                   bold=True)

    # Semnături tabelare
    doc.add_paragraph()
    sig_t = doc.add_table(rows=2, cols=3)
    sig_t.style = "Light Grid Accent 1"
    sig_t.rows[0].cells[0].text = "PROIECTANT"
    sig_t.rows[0].cells[1].text = "EXECUTANT"
    sig_t.rows[0].cells[2].text = "DIRIGINTE ȘANTIER"
    sig_t.rows[1].cells[0].text = f"{base._get(data,'pv_calitate_proiectant', base._get(data,'dtac_proiectant_specialitate', COMPANY['name']))}\n_______________"
    sig_t.rows[1].cells[1].text = f"{base._get(data,'pv_calitate_constructor', base._get(data,'exec_firma','—'))}\n_______________"
    sig_t.rows[1].cells[2].text = f"{base._get(data,'pv_calitate_diriginte', base._get(data,'exec_diriginte_santier','—'))}\n_______________"

    base._footer_signature(doc, proj, data)
    return base._save(doc)


# ============================================================================
# TEMPLATE NOU 9: Program Faze ISC (consumă program_faze_isc_judet + program_faze_baza_legala + program_control_model)
# ============================================================================
def _program_faze_isc(proj: Dict[str, Any]) -> bytes:
    data = proj.get("data") or {}
    doc = Document()
    base._company_header(doc)
    base._add_heading(doc, "PROGRAM DE CONTROL ÎN FAZE DETERMINANTE",
                      level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    base._add_para(doc, "(Conform Legea 10/1995 art. 22 + HG 1735/2006 + Ord. MLPAT 31/N/1995)",
                   italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
    doc.add_paragraph()

    base._add_kv_table(doc, [
        ("ISC Județean",  base._get(data, "program_faze_isc_judet", "ISC București-Ilfov")),
        ("Obiectiv",      proj.get("title", "—")),
        ("Beneficiar",    base._get(data, "beneficiar_nume")),
        ("Adresă",        base._get(data, "loc_consum_adresa")),
        ("Nr. AC",        f"{base._get(data,'ac_numar')} / {base._get(data,'ac_data_emitere')}"),
        ("Executant",     base._get(data, "exec_firma")),
        ("Cadru legal",   base._get(data, "program_faze_baza_legala",
                                      "Legea 10/1995 art. 22, HG 1735/2006, Ord. MLPAT 31/N/1995, NTPEE 2018")),
    ])

    doc.add_paragraph()
    base._add_para(doc, "Faze determinante stabilite pentru control:", bold=True)

    FAZE = [
        ("FD-01", "Predare-primire amplasament",                  "Înainte de începere lucrări",   "Proiectant + Diriginte + Executant"),
        ("FD-02", "Verificare trasaj și săpătură",                "Înainte de pozare conductă",    "Diriginte + Executant"),
        ("FD-03", "Verificare pat nisip + pozare conductă",       "Înainte de acoperire șanț",     "Diriginte + RTE + ISC (anunț)"),
        ("FD-04", "Probă de rezistență",                          "După finalizare montaj",         "Proiectant + Executant + RTE"),
        ("FD-05", "Probă de etanșeitate (24h)",                   "După probă rezistență admisă",   "Proiectant + Executant + RTE"),
        ("FD-06", "Recepție la terminarea lucrărilor (PVRTL)",    "Înainte de PIF",                 "Comisia recepție completă"),
        ("FD-07", "Punere în funcțiune (PIF)",                    "După PVRTL + cerere către OSD",  "OSD + Executant + Beneficiar"),
    ]
    t = doc.add_table(rows=1, cols=4)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    hdr[0].text = "Cod"; hdr[1].text = "Fază determinantă"; hdr[2].text = "Moment"; hdr[3].text = "Participanți"
    for c, f, m, p in FAZE:
        r = t.add_row().cells
        r[0].text = c; r[1].text = f; r[2].text = m; r[3].text = p

    doc.add_paragraph()
    base._add_para(doc, "Note:", bold=True)
    doc.add_paragraph("• ISC va fi notificat în scris cu minim 5 zile lucrătoare înainte de fiecare fază determinantă cu prezență obligatorie.", style="List Bullet")
    doc.add_paragraph("• Fiecare fază se consemnează prin Proces-Verbal de Control Fază Determinantă, semnat de toți participanții.", style="List Bullet")
    doc.add_paragraph("• Nerespectarea programului atrage sancțiuni conform Legea 10/1995 art. 31.", style="List Bullet")

    if base._truthy(data, "program_control_model"):
        doc.add_paragraph()
        base._add_para(doc, f"Model PCC aplicat: {base._get(data,'program_control_model')}", italic=True, size=9)

    base._footer_signature(doc, proj, data)
    return base._save(doc)


# ============================================================================
# REGISTRY EXTRA
# ============================================================================
EXTRA_TEMPLATES = {
    "pv_lucrari_ascunse":      {"label": "PV Lucrări Ascunse",                        "phase": "executie", "fn": _pv_la,         "norm": "Legea 10/1995 art. 23"},
    "pv_faza_determinanta":    {"label": "PV Control Fază Determinantă",              "phase": "executie", "fn": _pv_fd,         "norm": "Legea 10/1995 art. 22 + HG 1735"},
    "program_control_calitate":{"label": "Program Control Calitate (PCC)",            "phase": "pt",       "fn": _pcc,           "norm": "Legea 10/1995 + Ord. MLPAT 12/N/1995"},
    "referat_verificator":     {"label": "Referat Verificare Tehnică (RVT)",          "phase": "dtac",     "fn": _rvt,           "norm": "Legea 10/1995 art. 21 + Ord. MLPAT 777/2003"},
    "notificare_isc":          {"label": "Notificare ISC — Începere lucrări",         "phase": "executie", "fn": _notificare_isc,"norm": "Legea 50/1991 art. 7 + HG 1735/2006"},
    "as_built":                {"label": "As-Built (memoriu tehnic finalizat)",       "phase": "receptie", "fn": _as_built,      "norm": "HG 273/1994 + Ord. MLPAT 770/1997"},
    "dtac_lista_avize":        {"label": "DTAC — Listă Avize Utilități obținute",     "phase": "dtac",     "fn": _dtac_lista_avize, "norm": "Legea 50/1991 + HG 525/1996"},
    "pv_calitate":             {"label": "Proces-Verbal Control Calitate",            "phase": "executie", "fn": _pv_calitate,   "norm": "Legea 10/1995 + HG 925/1995"},
    "program_faze_isc":        {"label": "Program Control Faze Determinante (ISC)",   "phase": "executie", "fn": _program_faze_isc, "norm": "Legea 10/1995 art. 22 + HG 1735/2006"},
}


def register_into(base_module) -> None:
    """Merge EXTRA_TEMPLATES into base_module.TEMPLATES (apel la importul gas_doc_templates)."""
    for k, v in EXTRA_TEMPLATES.items():
        base_module.TEMPLATES.setdefault(k, v)
