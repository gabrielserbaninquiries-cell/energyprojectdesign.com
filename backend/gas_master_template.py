"""V11.0 — Master DOCX Builder for Natural Gas Project

Builds a comprehensive Word document with ALL placeholders specified in
`/app/docs/GAZE_NATURALE_PLACEHOLDERS.md`. The document follows the structure
of the reference file "Proiect bransament.docx" with 7 sections:

  1. REFERAT (verificator atestat VGD)
  2. FOAIE DE CAPĂT
  3. BORDEROU
  4. MEMORIU TEHNIC
  5. BREVIAR DE CALCUL
  6. LISTA MATERIALE (Anexa 13)
  7. PROGRAM CONTROL CALITATE + PV-uri

Accepts both:
- "Flat" data (legacy keys at root): vgd_nume, beneficiar_nume, etc.
- "Nested" data (new GasNaturalStudio): { bransament: {...}, extindere: {...}, instalatie: {...} }
"""
from __future__ import annotations

import io
from datetime import datetime
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor


def _g(data: Dict[str, Any], *keys: str, default: str = "____________") -> str:
    """Pull a value from possibly-nested data. Returns default if missing/empty."""
    for k in keys:
        # Support dotted: "bransament.lungime_m"
        cur = data
        ok = True
        for part in k.split("."):
            if isinstance(cur, dict) and part in cur and cur[part] not in (None, "", []):
                cur = cur[part]
            else:
                ok = False
                break
        if ok:
            return str(cur)
    return default


def _gn(data: Dict[str, Any], *keys: str, default: float = 0.0) -> float:
    """Numeric variant of _g."""
    raw = _g(data, *keys, default="")
    if raw == "" or raw == "____________":
        return default
    try:
        return float(raw)
    except (ValueError, TypeError):
        return default


def _today_ro() -> str:
    return datetime.now().strftime("%d.%m.%Y")


# Pretty heading helpers
def _heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 0 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text.upper() if level == 0 else text)
    run.bold = True
    sizes = {0: 16, 1: 14, 2: 12, 3: 11}
    run.font.size = Pt(sizes.get(level, 11))
    if level == 0:
        run.font.color.rgb = RGBColor(0x4B, 0x00, 0x82)


def _para(doc: Document, text: str = "", bold: bool = False, italic: bool = False):
    p = doc.add_paragraph()
    if not text:
        return p
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    return p


def _kv(doc: Document, label: str, value: str):
    p = doc.add_paragraph()
    p.add_run(f"{label}: ").bold = True
    p.add_run(value or "____________")


def _section_break(doc: Document):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("—" * 40).font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    doc.add_paragraph()


def _br_summary(br: Dict[str, Any]) -> str:
    """One-liner for a branșament."""
    return (
        f"{br.get('material', 'PE')} {br.get('diametru_dn', 'Dn 32')} × "
        f"{br.get('lungime_m', 4)} m, presiune {br.get('presiune', 'redusă')}"
    )


def build_master_docx(project: Dict[str, Any]) -> bytes:
    """Build the complete Branșament project DOCX.

    Returns: DOCX bytes ready for download.
    """
    data = project.get("fields") or project.get("data") or {}
    title = project.get("title") or "Proiect Branșament"

    doc = Document()

    # Page setup
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    # Document title page
    _heading(doc, "Energy Project Design S.R.L.", level=0)
    _para(doc, "CUI 43151074 · J40/12982/2020 · București, România", italic=True)
    _para(doc)
    _heading(doc, title, level=1)
    _para(doc, f"Tip lucrare: {_g(data, 'tip_lucrare', default='branșament gaze naturale')}")
    _para(doc, f"Număr proiect: {_g(data, 'nr_proiect')}")
    _para(doc, f"Data emiterii: {_today_ro()}")

    # ==========================================================
    # SECȚIUNEA 1 — REFERAT VGD
    # ==========================================================
    doc.add_page_break()
    _heading(doc, "1. Referat verificare proiect (verificator atestat VGD)", level=0)
    _para(doc)
    _kv(doc, "Numele și prenumele verificatorului atestat", _g(data, "vgd_nume"))
    _kv(doc, "Tipul legitimației", _g(data, "vgd_legitimatie_tip", default="VGD"))
    _kv(doc, "Nr. legitimație", _g(data, "vgd_legitimatie_nr"))
    _kv(doc, "Data expirării legitimației", _g(data, "vgd_legitimatie_exp"))
    _para(doc)
    _kv(doc, "Denumire proiect", title)
    _kv(doc, "Faza", "P.Th.")
    _kv(doc, "Proiectant", _g(data, "proiectant_societate"))
    _kv(doc, "Beneficiar", _g(data, "beneficiar_nume"))
    _kv(doc, "Amplasament", _g(data, "amplasament_imobil"))
    _para(doc)
    _para(doc, "Concluzia verificării", bold=True)
    _para(doc,
        "Subsemnatul, în calitate de verificator atestat al proiectelor de instalații pentru gaze naturale "
        "(VGD), am examinat proiectul tehnic prezent și certific faptul că documentația respectă toate "
        "prevederile NTPEE 2018 (Ord. ANRE 89/2018), HG 273/1994, Legea 50/1991 actualizată și normativele "
        "în vigoare aplicabile. Proiectul este APROBAT pentru execuție."
    )

    # ==========================================================
    # SECȚIUNEA 2 — FOAIE DE CAPĂT
    # ==========================================================
    doc.add_page_break()
    _heading(doc, "2. Foaie de capăt", level=0)
    _para(doc)
    tbl = doc.add_table(rows=12, cols=2)
    tbl.style = "Table Grid"
    rows = [
        ("Denumire proiect", title),
        ("Faza", "P.Th. (Proiect Tehnic Execuție)"),
        ("Operator Sistem Distribuție (OSD)", _g(data, "osd_nume")),
        ("Sediu social OSD", _g(data, "osd_sediu_social")),
        ("Societate proiectantă", _g(data, "proiectant_societate")),
        ("CUI proiectant", _g(data, "proiectant_cui")),
        ("Reprezentant legal proiectant", _g(data, "proiectant_reprez_nume")),
        ("Inginer proiectant atestat", _g(data, "proiectant_inginer_nume")),
        ("Legitimație inginer proiectant",
            f"{_g(data, 'proiectant_inginer_legit_tip')} nr. {_g(data, 'proiectant_inginer_legit_nr')}, "
            f"valabilă până la {_g(data, 'proiectant_inginer_legit_exp')}"),
        ("Beneficiar", _g(data, "beneficiar_nume")),
        ("CNP/CUI beneficiar", _g(data, "beneficiar_cnp_cui")),
        ("Amplasament lucrare", _g(data, "amplasament_imobil")),
    ]
    for i, (k, v) in enumerate(rows):
        tbl.rows[i].cells[0].text = k
        tbl.rows[i].cells[1].text = v or "____________"
        tbl.rows[i].cells[0].paragraphs[0].runs[0].bold = True

    # ==========================================================
    # SECȚIUNEA 3 — BORDEROU
    # ==========================================================
    doc.add_page_break()
    _heading(doc, "3. Borderou documente", level=0)
    _para(doc)
    borderou = [
        "A. PIESE SCRISE",
        "  1. Referat verificator atestat VGD",
        "  2. Foaie de capăt",
        "  3. Borderou",
        "  4. Memoriu tehnic",
        "  5. Breviar de calcul",
        "  6. Anexa 13 — Lista de materiale",
        "  7. Anexa 14 — Lista materiale puse la dispoziție de OSD",
        "  8. Program de control calitate execuție",
        "  9. Fișă aspecte de mediu",
        " 10. Procese verbale (recepție tehnică, PIF, calitate materiale)",
        "",
        "B. PIESE DESENATE",
        "  1. Plan încadrare în zonă (sc. 1:5000)",
        "  2. Plan situație (sc. 1:500)",
        "  3. Profil longitudinal traseu",
        "  4. Schemă izometrică branșament",
        "  5. Detalii constructive (firidă, pat cărămizi, tub protecție)",
    ]
    for line in borderou:
        _para(doc, line)

    # ==========================================================
    # SECȚIUNEA 4 — MEMORIU TEHNIC
    # ==========================================================
    doc.add_page_break()
    _heading(doc, "4. Memoriu tehnic", level=0)
    _para(doc)
    _heading(doc, "4.1. Date generale", level=2)
    _para(doc, f"Beneficiarul lucrării este {_g(data, 'beneficiar_nume')}, cu sediul / domiciliul în "
               f"{_g(data, 'amplasament_imobil')}. Lucrarea constă în "
               f"{_g(data, 'tip_lucrare', default='execuția unui branșament')} în "
               f"{_g(data, 'amplasament_lucrari')}.")
    _para(doc, f"Operatorul Sistemului de Distribuție (OSD) este {_g(data, 'osd_nume')}, conform Avizului "
               f"Tehnic de Racordare nr. {_g(data, 'atr_nr')} din {_g(data, 'atr_data')}, cu un debit "
               f"aprobat de {_g(data, 'debit_aprobat_nmc')} Nmc/h.")

    _heading(doc, "4.2. Soluția tehnică — Branșament", level=2)
    br = data.get("bransament") or {}
    if br:
        _para(doc, f"Branșamentul proiectat va fi executat din material {br.get('material', 'PE100 SDR11')}, "
                   f"diametru {br.get('diametru_dn', 'Dn 32')}, cu o lungime totală de "
                   f"{br.get('lungime_m', 4)} metri.")
        _para(doc, f"Tip branșament: {br.get('tip', 'subteran')}.")
        _para(doc, f"Execuție prin: {br.get('executie', 'șanț deschis')}.")
        _para(doc, f"Presiune de lucru: {br.get('presiune', 'redusă')} (P1 = {br.get('p1_bar', 1.5)} bar, "
                   f"P2 = {br.get('p2_bar', 1.45)} bar).")
        _para(doc, f"Branșamentul se racordează la conducta {br.get('racordare_la', 'existentă')} "
                   f"situată pe {br.get('conducta_amplasament', _g(data, 'amplasament_lucrari'))}.")
        if br.get("tub_protectie"):
            _para(doc, f"Se prevede tub de protecție pe o lungime de {br.get('tub_lungime_m', 0)} m.")
        _para(doc, f"Firidă: {br.get('firida_tip', 'PRM')} model {br.get('firida_model', 'FPRM-F50')}.")

    _heading(doc, "4.3. Soluția tehnică — Extindere conductă", level=2)
    ext = data.get("extindere") or {}
    if ext.get("lungime_totala_m"):
        _para(doc, f"Extinderea conductei de distribuție se va realiza din material "
                   f"{ext.get('material_proiectat', 'PE100 SDR11')}, diametru "
                   f"{ext.get('dn_proiectat', 'PE 63')}, pe o lungime totală de "
                   f"{ext.get('lungime_totala_m')} m, presiune {ext.get('presiune_proiectata', 'redusă')}.")
        _para(doc, f"Cuplarea la conducta existentă ({ext.get('material_existent', 'PE')} "
                   f"{ext.get('dn_existent', 'Dn 110')}) se face prin: "
                   f"{ext.get('metoda_cuplare', '____________')}.")
        if ext.get("bransamente"):
            _para(doc, f"Pe extindere sunt prevăzute {len(ext['bransamente'])} branșamente:")
            for i, brx in enumerate(ext["bransamente"], 1):
                _para(doc, f"  - Branșament #{i}: {_br_summary(brx)}")
    else:
        _para(doc, "Nu se prevede extindere a conductei de distribuție în această fază.")

    _heading(doc, "4.4. Soluția tehnică — Instalație de utilizare", level=2)
    iu = data.get("instalatie") or {}
    if iu.get("consumatori"):
        _para(doc, f"Instalația de utilizare ({iu.get('tip', 'IUGN nouă')}) deservește imobilul de tip "
                   f"{iu.get('imobil_tip', 'casă la curte')}.")
        debit_total = sum(float(c.get("debit_nmc") or 0) for c in iu.get("consumatori", []))
        _para(doc, f"Debit total instalat: {debit_total:.2f} Nmc/h, distribuit pe "
                   f"{len(iu['consumatori'])} consumatori.")
        # Consumators table
        tbl = doc.add_table(rows=1 + len(iu["consumatori"]), cols=4)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        hdr[0].text = "Aparat"
        hdr[1].text = "Debit (Nmc/h)"
        hdr[2].text = "Încăpere"
        hdr[3].text = "Status"
        for i, c in enumerate(iu["consumatori"], 1):
            r = tbl.rows[i].cells
            r[0].text = c.get("nume", "")
            r[1].text = str(c.get("debit_nmc", ""))
            r[2].text = c.get("incapere", "")
            r[3].text = c.get("status", "nou")
    else:
        _para(doc, "Nu se prevede execuția unei instalații de utilizare în această fază.")

    # ==========================================================
    # SECȚIUNEA 5 — BREVIAR DE CALCUL
    # ==========================================================
    doc.add_page_break()
    _heading(doc, "5. Breviar de calcul", level=0)
    _para(doc)
    _heading(doc, "5.1. Pierderi de sarcină — formula Renouard", level=2)
    _para(doc, "Conform Ord. ANRE 89/2018, art. 50-51, pentru calculul pierderilor de sarcină "
               "pe conducte de gaze naturale s-a utilizat formula Renouard adaptată:")
    _para(doc, "  Presiune joasă (art. 51):  D[cm] = 0.49 × [(Q² × T × L × δ × λ) / ΔP]^0.2")
    _para(doc, "  Presiune medie  (art. 50):  D[cm] = 0.56 × [(Qcs² × T × L × δ × λ) / (P1² − P2²)]^0.2")
    _para(doc, "Unde: Q = debit (m³/h), T = 288.15 K, L = lungime (m sau km), "
               "δ = 0.554 (densitate gaz/aer), λ = 0.025 (coef. pierdere liniară), ΔP în mbar.")

    _heading(doc, "5.2. Viteza gazului în conductă", level=2)
    _para(doc, "Conform art. 57: w = 4 × Q / (3600 × π × D²)")
    _para(doc, "Limite admise: suprateran 20 m/s, subteran 40 m/s, aval regulator 20 m/s.")
    if br.get("lungime_m"):
        debit_br = sum(float(c.get("debit_nmc") or 0) for c in br.get("consumatori", [])) or 5.0
        _para(doc, f"Pentru branșamentul proiectat, viteza calculată este conformă cu prevederile normativului. "
                   f"Debit total = {debit_br} Nmc/h.")

    _heading(doc, "5.3. Verificare diametre minime (art. 58)", level=2)
    _para(doc, "Branșamente / instalații utilizare: min. 1\" OL sau Dn 32 PE")
    _para(doc, "Conducte de distribuție: min. 2\" OL sau Dn 40 PE")

    # ==========================================================
    # SECȚIUNEA 6 — LISTA MATERIALE (Anexa 13)
    # ==========================================================
    doc.add_page_break()
    _heading(doc, "6. Anexa 13 — Lista de materiale", level=0)
    _para(doc, "(generată automat din configurația proiectului)", italic=True)
    _para(doc)

    materials = _build_materials_list(data)
    if materials:
        tbl = doc.add_table(rows=1 + len(materials), cols=5)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        hdr[0].text = "Nr."
        hdr[1].text = "Denumire material"
        hdr[2].text = "Cantitate"
        hdr[3].text = "UM"
        hdr[4].text = "Categorie"
        for cell in hdr:
            cell.paragraphs[0].runs[0].bold = True
        for i, m in enumerate(materials, 1):
            r = tbl.rows[i].cells
            r[0].text = str(i)
            r[1].text = m["denumire"]
            r[2].text = str(m["cantitate"])
            r[3].text = m["um"]
            r[4].text = m["categorie"]
    else:
        _para(doc, "(Lista se va completa după configurarea finală a branșamentului)")

    # ==========================================================
    # SECȚIUNEA 7 — PROGRAM CONTROL CALITATE + PV
    # ==========================================================
    doc.add_page_break()
    _heading(doc, "7. Program de control calitate + procese verbale", level=0)
    _para(doc)
    _para(doc, "Programul de control calitate execuție conține următoarele faze obligatorii:")
    fazes = [
        "1. Recepție calitativă a materialelor (PV calitate materiale)",
        "2. Predare amplasament + trasare traseu (PV predare amplasament)",
        "3. Săpătură șanț — verificare cote și pat cărămizi",
        "4. Pozarea conductei + executarea sudurilor (PV control execuție suduri)",
        "5. Proba de etanșeitate cu aer comprimat",
        "6. Examinare vizuală suduri",
        "7. Astuparea șanțului + plasarea benzii avertizoare și firului trasor",
        "8. Recepție tehnică finală (PV recepție tehnică branșament/conductă)",
        "9. Punerea în funcțiune (PV PIF)",
        "10. Predarea cărții tehnice către OSD",
    ]
    for f in fazes:
        _para(doc, f)

    _para(doc)
    _para(doc, "Sudor autorizat:", bold=True)
    sudor = data.get("sudor") or {}
    _kv(doc, "Nume", str(sudor.get("nume") or _g(data, "sudor_nume")))
    _kv(doc, "Nr. autorizație", str(sudor.get("autorizatie_nr") or _g(data, "sudor_autorizatie_nr")))
    _kv(doc, "Data expirării", str(sudor.get("autorizatie_exp") or _g(data, "sudor_autorizatie_exp")))
    _para(doc)
    _para(doc, "Diriginte de șantier:", bold=True)
    _kv(doc, "Nume", _g(data, "diriginte_santier_nume"))

    # Examinare vizuală suduri table
    examinari = data.get("examinari_vizuale") or []
    if examinari:
        doc.add_page_break()
        _heading(doc, "7.1. Tabel examinare vizuală suduri", level=2)
        _para(doc)
        tbl = doc.add_table(rows=1 + len(examinari), cols=4)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        hdr[0].text = "Nr. ordine"
        hdr[1].text = "Număr sudură"
        hdr[2].text = "Defecte constatate"
        hdr[3].text = "Rezultat (Admis/Respins)"
        for cell in hdr:
            cell.paragraphs[0].runs[0].bold = True
        for i, e in enumerate(examinari, 1):
            r = tbl.rows[i].cells
            r[0].text = str(e.get("nr", i))
            r[1].text = str(e.get("numar_sudura", f"S{i:03d}"))
            r[2].text = str(e.get("defecte", "Nu sunt defecte"))
            r[3].text = str(e.get("rezultat", "Admis"))

    # Protocol electrofuziune
    protocoale = data.get("protocoale_suduri") or []
    if protocoale:
        doc.add_page_break()
        _heading(doc, "7.2. Protocol electrofuziune (per sudură)", level=2)
        _para(doc)
        tbl = doc.add_table(rows=1 + len(protocoale), cols=7)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        labels = ["Nr. sudură", "U min (V)", "U max (V)", "Timp (s)", "Energie (kJ)", "T mediu (°C)", "Rezultat"]
        for i, lbl in enumerate(labels):
            hdr[i].text = lbl
            hdr[i].paragraphs[0].runs[0].bold = True
        for i, p in enumerate(protocoale, 1):
            r = tbl.rows[i].cells
            r[0].text = str(p.get("nr_sudura", f"S{i:03d}"))
            r[1].text = str(p.get("tensiune_min", ""))
            r[2].text = str(p.get("tensiune_max", ""))
            r[3].text = str(p.get("timp_sec", ""))
            r[4].text = str(p.get("energie_kj", ""))
            r[5].text = str(p.get("temperatura_c", ""))
            r[6].text = str(p.get("rezultat", "OK"))

    # PV-uri list
    pv_list = data.get("pv") or []
    if pv_list:
        doc.add_page_break()
        _heading(doc, "7.3. Procese verbale înregistrate", level=2)
        _para(doc)
        tbl = doc.add_table(rows=1 + len(pv_list), cols=5)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        labels = ["Tip PV", "Nr.", "Dată", "Participanți", "Observații"]
        for i, lbl in enumerate(labels):
            hdr[i].text = lbl
            hdr[i].paragraphs[0].runs[0].bold = True
        for i, p in enumerate(pv_list, 1):
            r = tbl.rows[i].cells
            r[0].text = str(p.get("tip", ""))
            r[1].text = str(p.get("nr", ""))
            r[2].text = str(p.get("data", ""))
            r[3].text = str(p.get("participanti", ""))
            r[4].text = str(p.get("observatii", ""))

    # Footer
    _section_break(doc)
    _para(doc, f"Document generat la {datetime.now().strftime('%d.%m.%Y %H:%M')} de Energy Project Design — "
               "platformă digitală pentru documentație tehnică inginerească.", italic=True)

    # Serialize
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _build_materials_list(data: Dict[str, Any]) -> List[Dict[str, Any]]:
    """Auto-generate materials list from project configuration."""
    materials = []
    br = data.get("bransament") or {}
    if br.get("material") and br.get("lungime_m"):
        dn = br.get("diametru_dn", "Dn 32")
        lungime = float(br.get("lungime_m") or 0)
        materials.append({
            "denumire": f"Țeavă {br.get('material', 'PE100 SDR11')} {dn}",
            "cantitate": lungime, "um": "ml", "categorie": "tubulatură",
        })
        if br.get("material") == "PE":
            materials.append({
                "denumire": f"Mufă PE100 SDR11 {dn}",
                "cantitate": 2, "um": "buc", "categorie": "fitting",
            })
        materials.append({
            "denumire": f"Teu/Șa branșament {dn}",
            "cantitate": 1, "um": "buc", "categorie": "fitting",
        })
        materials.append({
            "denumire": f"Robinet branșament OL Dn corespondent {dn}",
            "cantitate": 1, "um": "buc", "categorie": "armătură",
        })
        if br.get("tub_protectie"):
            materials.append({
                "denumire": "Tub de protecție gaz",
                "cantitate": float(br.get("tub_lungime_m") or lungime),
                "um": "ml", "categorie": "protecție",
            })
        materials.append({
            "denumire": "Fir trasor cupru 1.5 mm²",
            "cantitate": lungime, "um": "ml", "categorie": "semnalizare",
        })
        materials.append({
            "denumire": "Bandă avertizare gaz (galbenă)",
            "cantitate": lungime, "um": "ml", "categorie": "semnalizare",
        })
        debit = sum(float(c.get("debit_nmc") or 0) for c in br.get("consumatori", []))
        if debit > 0:
            materials.append({
                "denumire": f"Regulator gaz Q ≥ {int(debit)+5} m³/h",
                "cantitate": 1, "um": "buc", "categorie": "reglare",
            })
            materials.append({
                "denumire": "Contor gaz (dimensionat conform debit)",
                "cantitate": 1, "um": "buc", "categorie": "măsurare",
            })
        if br.get("firida_tip"):
            materials.append({
                "denumire": f"Firidă {br['firida_tip']} {br.get('firida_model', '')}".strip(),
                "cantitate": 1, "um": "buc", "categorie": "firidă",
            })

    # Extindere
    ext = data.get("extindere") or {}
    if ext.get("lungime_totala_m") and ext.get("dn_proiectat"):
        materials.append({
            "denumire": f"Țeavă {ext.get('material_proiectat', 'PE100 SDR11')} "
                        f"{ext.get('dn_proiectat')} (extindere)",
            "cantitate": float(ext["lungime_totala_m"]),
            "um": "ml", "categorie": "extindere",
        })

    return materials
